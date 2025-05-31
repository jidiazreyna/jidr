#!/usr/bin/env python3
"""
Generador GUI – Sobreseimiento por Prescripción
==============================================
Versión revisada (bug-fix 1)
----------------------------
• Corrige `AttributeError: HechoWidget` → ahora guarda `self.idx` y lo usa en `texto()`.
• Resto idéntico a la versión anterior (fecha automática, tipo de tribunal, sexo M/F, n-imputados/hechos, zoom, copiar, exportar DOCX).

Dependencias rápidas
--------------------
    pip install PySide6 python-docx

Ejecución
---------
    python prescripcion_gui.py
"""
from __future__ import annotations

import sys, re
from typing import List, Dict
from collections import UserDict
from pathlib import Path
from datetime import datetime
from PySide6.QtCore import Qt
from PySide6.QtGui import QFont
from PySide6.QtWidgets import (
    QApplication, QWidget, QLabel, QLineEdit, QTextEdit, QTextBrowser,
    QSpinBox, QComboBox, QRadioButton, QButtonGroup, QPushButton,
    QVBoxLayout, QHBoxLayout, QGridLayout, QScrollArea, QFileDialog,
    QMessageBox, QSizePolicy
)

# ─────────────────────────────────────────────────────────── Helpers plantilla ──

def _pronombres_imputados(sexos: List[str]) -> Dict[str, str]:
    n = len(sexos)
    cant_m = sexos.count("M"); cant_f = n - cant_m
    if n == 1:
        art = "el imputado" if sexos[0] == "M" else "la imputada"
        asistido = "asistido" if sexos[0] == "M" else "asistida"
        acusado = "acusado" if sexos[0] == "M" else "acusada"
        le_les = "le"
    else:
        if cant_f == n:
            art, asistido, acusado = "las imputadas", "asistidas", "acusadas"
        elif cant_m == n:
            art, asistido, acusado = "los imputados", "asistidos", "acusados"
        else:
            art, asistido, acusado = "los imputados", "asistidos", "acusados"
        le_les = "les"
    return {
        "imputado_articulo": art,
        "le_les": le_les,
        "asistido_label": asistido,
        "acusado_label": acusado,
    }

UNIDADES = (
    'cero', 'uno', 'dos', 'tres', 'cuatro', 'cinco',
    'seis', 'siete', 'ocho', 'nueve', 'diez', 'once',
    'doce', 'trece', 'catorce', 'quince', 'dieciséis',
    'diecisiete', 'dieciocho', 'diecinueve', 'veinte',
    'veintiuno', 'veintidós', 'veintitrés', 'veinticuatro',
    'veinticinco', 'veintiséis', 'veintisiete', 'veintiocho',
    'veintinueve'
)
DECENAS = (
    'treinta', 'cuarenta', 'cincuenta', 'sesenta',
    'setenta', 'ochenta', 'noventa'
)
CENTENAS = (
    'cien', 'doscientos', 'trescientos', 'cuatrocientos',
    'quinientos', 'seiscientos', 'setecientos', 'ochocientos',
    'novecientos'
)

def numero_a_letras(num: int) -> str:
    if num < 0:
        return "menos " + numero_a_letras(abs(num))
    if num <= 29:
        return UNIDADES[num]
    if num < 100:
        dec = (num // 10) - 3
        uni = num % 10
        letra_decena = DECENAS[dec]
        return letra_decena if uni == 0 else f"{letra_decena} y {UNIDADES[uni]}"
    if num < 1000:
        cent = (num // 100) - 1
        resto = num % 100
        if num == 100:
            return "cien"
        return CENTENAS[cent] if resto == 0 else f"{CENTENAS[cent]} {numero_a_letras(resto)}"
    if num < 10000:
        mil = num // 1000
        resto = num % 1000
        prefix = "mil" if mil == 1 else f"{numero_a_letras(mil)} mil"
        return prefix if resto == 0 else f"{prefix} {numero_a_letras(resto)}"
    return str(num)

def obtener_fecha_en_letras():
    fecha_actual = datetime.now()
    dia = fecha_actual.day
    mes_numero = fecha_actual.month
    anio = fecha_actual.year
    dia_letras = numero_a_letras(dia)
    anio_letras = numero_a_letras(anio)
    meses = {
        1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
        5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
        9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
    }
    mes_str = meses.get(mes_numero, '')
    return f"{dia_letras} de {mes_str} de {anio_letras}"

class SafeDict(UserDict):
    """Dict que devuelve el marcador sin reemplazar ante claves faltantes."""
    def __missing__(self, key):
        return "{" + key + "}"

# Ordinales para numerar los hechos cuando hay más de uno
ORDINALES_HECHOS = [
    "Primer",
    "Segundo",
    "Tercer",
    "Cuarto",
    "Quinto",
    "Sexto",
    "Séptimo",
    "Octavo",
    "Noveno",
    "Décimo",
    "Undécimo",
    "Duodécimo",
    "Decimotercero",
    "Decimocuarto",
    "Decimoquinto",
]

TEMPLATE = """
<p align='justify'>Córdoba, {fecha_letras}.</p>
<p align='justify'>VISTA: la presente causa caratulada {caratula}, venida a {este_esta} {tribunal} a los efectos de resolver la situación procesal de {nombre_apellido}</p>
<p align='justify'>DE LA QUE RESULTA: Que {imputado_articulo} {nombre_apellido} se {le_les} atribuye {hechos_label}:</p>
{hechos_html}
<p align='justify'><b>Y CONSIDERANDO:</b></p>
<p align='justify'>I. Que durante la instrucción se colectaron los siguientes elementos probatorios: {prueba}</p>
<p align='justify'>II. Que {fiscal_titulo} {fiscal} requiere el sobreseimiento {tipo_sobreseimiento} en la presente causa respecto de {nombre_apellido}, por {hechos_mencionados} supra, {encuadrado_bajo} bajo la calificación legal de {delitos}, en virtud de lo dispuesto por los arts. 348 y 350 inc. 4º del CPP, en función del art. 59 inc. 3º del CP, brindando los siguientes argumentos: {argumentos_fiscal}</p>
<p align='justify'><b>III. Conclusiones</b></p>
<p align='justify'>Analizada la cuestión traída a estudio, se advierte que {hechos_atribuidos} a {nombre_apellido} {encuadra_n} efectivamente bajo la calificación legal de {delitos}, cuya pena máxima conminada en abstracto es de {penamaxima} de prisión. En este sentido, cabe aclarar que a los fines de computar el término para la prescripción del hecho imputado a {nombre_apellido} en los presentes autos se debe tener en cuenta {interrupcion}, conforme surge de la planilla prontuarial, del Registro Nacional de Reincidencia y del Sistema de Administración de Causas. En efecto, {fundamentacion}</p>
<p align='justify'>Así, teniendo en cuenta los términos referidos, entiendo que corresponde desvincular de la presente causa al imputado {nombre_apellido} por la causal de procedencia descripta en el art. 350 inc. 4º del CPP. Ello así, porque, tal como lo manifestó {fiscal_articulo}, a la fecha, ha transcurrido con exceso el término establecido por el art. 62 inc. 2° del CP ({penamaxima} en este caso), el que desde la fecha {fechas_prescripcion} no fue interrumpido por la comisión de nuevos delitos, conforme surge de la planilla prontuarial y del informe del Registro Nacional de Reincidencia incorporados digitalmente, y no procede ninguna de las causales contempladas por el art. 67 del CP, motivo por el cual ha de tenerse a la prescripción como causal de previo y especial pronunciamiento. Así lo establece el alto tribunal de esta provincia: “…Esta Sala, compartiendo la posición ya asumida por otra integración y por mayoría (A. nº 76, 29/6/93, &quot;Cappa&quot;; A. nº 60, 14/6/94, &quot;Vivian&quot;), ha sostenido que habida cuenta de la naturaleza sustancial de las distintas causales de sobreseimiento, las extintivas de la acción deben ser de previa consideración (T.S.J., Sala Penal, A. n° 26, 19/2/99, &quot;Rivarola&quot;; &quot;Pérez&quot;, cit.). Por ello, la sola presencia de una causal extintiva de la acción -en el caso, la prescripción- debe ser estimada independientemente cualquiera sea la oportunidad de su producción y de su conocimiento por el Tribunal, toda vez que -en términos procesales- significa un impedimento para continuar ejerciendo los poderes de acción y de jurisdicción en procura de un pronunciamiento sobre el fondo (TSJ, Sala Penal, “CARUNCHIO, Oscar Rubén p.s.a. Homicidio Culposo -Recurso de Casación-” -Expte. &quot;C&quot;, 36/03-, S. n.° 104 de fecha 16/9/2005).</p>
<p align='justify'>IV. En consecuencia, y de conformidad a lo normado por los arts. 59 inc. 3° y 62 inc. 2° del CP y 350 del CPP, corresponde declarar prescripta la pretensión punitiva penal emergente {hechos_configurativos} de {delitos} que se le {atribuia_n} a {nombre_apellido}.</p>
<p align='justify'>V. Finalmente, deberá oficiarse a la Policía de la Provincia de Córdoba y al Registro Nacional de Reincidencia a fin de informar lo aquí resuelto.</p>
<p align='justify'>Por lo expresado y disposiciones legales citadas; <b>RESUELVO:</b></p>
<p align='justify'>I. Sobreseer {sobreseimiento_tipo}, respecto {hecho_plural} de {fechas} {fechasdeloshechos}, a {nombre_apellido}, de condiciones personales ya relacionadas, por {hecho_calificado} como {delitos}, de conformidad con lo establecido por los arts. 348 y 350 inc. 4º del CPP, en función de los arts. 59 inc. 3º, 62 inc. 2º y 67 del CP.</p>
<p align='justify'>II. Ofíciese a la Policía de la Provincia de Córdoba y al Registro Nacional de Reincidencia, a sus efectos.</p>
<p align='justify'>PROTOCOLÍCESE Y NOTIFÍQUESE.</p>
"""

def render_prescripcion(*, sexos_imputados: List[str], nombres: List[str], hechos: List[str], **campos) -> str:
    auto = _pronombres_imputados(sexos_imputados)
    auto["hechos_label"] = "el siguiente hecho" if len(hechos) == 1 else "los siguientes hechos"

    imp_lines = []
    for idx, (nom, sx) in enumerate(zip(nombres, sexos_imputados), start=1):
        art = "el imputado" if sx == "M" else "la imputada"
        imp_lines.append(f"{idx}. {art} {nom}")
    if len(hechos) == 1:
        hechos_html = f"<p align='justify'><i>{hechos[0]}</i></p>"
    else:
        hechos_html = "\n".join(
            f"<p align='justify'><b>{ORDINALES_HECHOS[i] if i < len(ORDINALES_HECHOS) else f'{i+1}°'} hecho:</b> <i>{txt}</i></p>"
            for i, txt in enumerate(hechos)
        )

    auto["imputados_html"] = "\n".join(
        f"<p align='justify'>{l}</p>" for l in imp_lines
    ) or "[imputados]"
    auto["hechos_html"] = hechos_html or "[hechos]"

    data = SafeDict(auto)
    data.update(campos)
    return TEMPLATE.format_map(data)

# ───────────────────────────────────────────────────────────── Widgets ──
class ImputadoWidget(QWidget):
    def __init__(self, idx: int):
        super().__init__(); self.idx = idx
        lay = QHBoxLayout(self)
        self.edit_nombre = QLineEdit();
        self.edit_nombre.setPlaceholderText(f"Nombre imputado #{idx+1}")
        self.edit_datos = QLineEdit();
        self.edit_datos.setPlaceholderText(f"Datos personales #{idx+1}")
        self.rb_m = QRadioButton("M"); self.rb_f = QRadioButton("F");
        self.rb_m.setChecked(True)
        grp = QButtonGroup(self); grp.addButton(self.rb_m); grp.addButton(self.rb_f)
        for w in (self.edit_nombre, self.edit_datos, self.rb_m, self.rb_f):
            lay.addWidget(w)
        lay.addStretch()

    def nombre(self) -> str:
        return self.edit_nombre.text().strip() or f"Imputado#{self.idx+1}"

    def datos(self) -> str:
        return self.edit_datos.text().strip()

    def sexo(self) -> str:
        return "M" if self.rb_m.isChecked() else "F"

class HechoWidget(QWidget):
    def __init__(self, idx:int):
        super().__init__(); self.idx = idx
        lay = QVBoxLayout(self)
        self.txt = QTextEdit(); self.txt.setPlaceholderText(f"Descripción hecho #{idx+1}"); self.txt.setFixedHeight(50)
        lay.addWidget(self.txt)
    def texto(self)->str: return self.txt.toPlainText().strip() or f"[hecho {self.idx+1}]"

class ZoomBrowser(QTextBrowser):
    def wheelEvent(self,ev):
        if ev.modifiers() & Qt.ControlModifier:
            self.zoomIn(1) if ev.angleDelta().y()>0 else self.zoomOut(1); ev.accept()
        else: super().wheelEvent(ev)

class PrescripcionGUI(QWidget):
    def __init__(self):
        super().__init__(); self.setWindowTitle("Prescripción – Generador rápido"); self.resize(1100,600)
        main = QHBoxLayout(self)
        # formulario con scroll
        scroll=QScrollArea(); scroll.setWidgetResizable(True); form_holder=QWidget(); self.form_layout=QVBoxLayout(form_holder); self.form_layout.setAlignment(Qt.AlignTop); scroll.setWidget(form_holder)
        main.addWidget(scroll,2)
        # preview
        self.preview=ZoomBrowser(); self.preview.document().setDefaultFont(QFont("Times New Roman",12)); main.addWidget(self.preview,3)
        # campos fijos
        grid=QGridLayout(); row=0; self.form_layout.addLayout(grid)
        grid.addWidget(QLabel("Carátula:"),row,0); self.ed_caratula=QLineEdit(); grid.addWidget(self.ed_caratula,row,1); row+=1
        grid.addWidget(QLabel("Tribunal:"),row,0); self.ed_tribunal=QLineEdit(); grid.addWidget(self.ed_tribunal,row,1); row+=1
        grid.addWidget(QLabel("Tipo tribunal:"),row,0); self.cb_tipo=QComboBox(); self.cb_tipo.addItems(["Juzgado","Cámara"]); grid.addWidget(self.cb_tipo,row,1); row+=1
        grid.addWidget(QLabel("Prueba:"),row,0); self.ed_prueba=QLineEdit(); grid.addWidget(self.ed_prueba,row,1); row+=1
        grid.addWidget(QLabel("Fiscal:"),row,0); self.ed_fiscal=QLineEdit(); grid.addWidget(self.ed_fiscal,row,1); row+=1
        grid.addWidget(QLabel("N° imputados:"),row,0); self.spin_imp=QSpinBox(minimum=1,maximum=8,value=1); grid.addWidget(self.spin_imp,row,1); row+=1
        grid.addWidget(QLabel("N° hechos:"),row,0); self.spin_hec=QSpinBox(minimum=1,maximum=8,value=1); grid.addWidget(self.spin_hec,row,1); row+=1
        # contenedores dinámicos
        self.box_imputados=QVBoxLayout(); self.form_layout.addLayout(self.box_imputados)
        self.box_hechos=QVBoxLayout(); self.form_layout.addLayout(self.box_hechos)
        # botones
        btn_row=QHBoxLayout(); self.form_layout.addLayout(btn_row)
        btn_copy=QPushButton("Copiar"); btn_docx=QPushButton("Exportar DOCX"); btn_row.addWidget(btn_copy); btn_row.addWidget(btn_docx); btn_row.addStretch()
        # señales
        for w in (self.ed_caratula,self.ed_tribunal,self.ed_prueba,self.ed_fiscal): w.textChanged.connect(self.update_preview)
        self.cb_tipo.currentTextChanged.connect(self.update_preview)
        self.spin_imp.valueChanged.connect(self.refresh_imputados)
        self.spin_hec.valueChanged.connect(self.refresh_hechos)
        btn_copy.clicked.connect(self.copy_clip)
        btn_docx.clicked.connect(self.save_docx)
        # dinámicos
        self.imputados_widgets:List[ImputadoWidget]=[]; self.hechos_widgets:List[HechoWidget]=[]
        self.refresh_imputados(); self.refresh_hechos(); self.update_preview()

    def refresh_imputados(self):
        target=self.spin_imp.value()
        while len(self.imputados_widgets)<target:
            w = ImputadoWidget(len(self.imputados_widgets))
            self.imputados_widgets.append(w)
            self.box_imputados.addWidget(w)
            w.edit_nombre.textChanged.connect(self.update_preview)
            w.edit_datos.textChanged.connect(self.update_preview)
            w.rb_m.toggled.connect(self.update_preview)
            w.rb_f.toggled.connect(self.update_preview)
        while len(self.imputados_widgets)>target:
            w=self.imputados_widgets.pop(); w.setParent(None)
        self.update_preview()

    def refresh_hechos(self):
        target=self.spin_hec.value()
        while len(self.hechos_widgets)<target:
            w=HechoWidget(len(self.hechos_widgets)); self.hechos_widgets.append(w); self.box_hechos.addWidget(w); w.txt.textChanged.connect(self.update_preview)
        while len(self.hechos_widgets)>target:
            w=self.hechos_widgets.pop(); w.setParent(None)
        self.update_preview()

    def update_preview(self):
        sexos = [w.sexo() for w in self.imputados_widgets]
        nombres = [w.nombre() for w in self.imputados_widgets]
        datos = [w.datos() for w in self.imputados_widgets]
        nombres_datos = [
            f"{n}, {d}" if d else n
            for n, d in zip(nombres, datos)
        ]
        hechos = [w.texto() for w in self.hechos_widgets]
        campos = {
            "fecha_letras": obtener_fecha_en_letras(),
            "caratula": self.ed_caratula.text().strip() or "[carátula]",
            "este_esta": "este" if self.cb_tipo.currentText() == "Juzgado" else "esta",
            "tribunal": self.ed_tribunal.text().strip() or "[tribunal]",
            "nombre_apellido": ", ".join(nombres_datos) or "[imputados]",
            "prueba": self.ed_prueba.text().strip() or "[prueba]",
            "fiscal": self.ed_fiscal.text().strip() or "[fiscal]",
        }
        html = render_prescripcion(
            sexos_imputados=sexos,
            nombres=nombres_datos,
            hechos=hechos,
            **campos,
        )
        self.preview.setHtml(html)

    def copy_clip(self):
        QApplication.clipboard().setText(self.preview.toPlainText())
        QMessageBox.information(self,"Copiado","Texto copiado al portapapeles")

    def save_docx(self):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            QMessageBox.warning(self,"python-docx faltante","Instale con: pip install python-docx"); return
        path,_=QFileDialog.getSaveFileName(self,"Guardar DOCX","","Word (*.docx)")
        if not path: return
        doc=Document(); doc._body.clear_content()
        for par in self.preview.toPlainText().split('\n'):
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            run=p.add_run(par); run.font.name="Times New Roman"; run.font.size=Pt(12)
        doc.save(path); QMessageBox.information(self,"Guardado",f"Archivo guardado en:\n{path}")

# ───────────────────────────────────────────────────────── main ──
if __name__=="__main__":
    app=QApplication(sys.argv)
    gui=PrescripcionGUI(); gui.show()
    sys.exit(app.exec())
