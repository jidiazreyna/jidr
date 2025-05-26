#!/usr/bin/env python
# -*- coding: utf-8 -*-

import sys, os, re
from datetime import datetime
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PySide6.QtCore import Qt, QEvent
from PySide6.QtWidgets import QFileDialog
from PySide6.QtGui import QFont, QPainter
from functools import partial
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QLabel, QLineEdit, QTextEdit,
    QComboBox, QSpinBox, QRadioButton, QButtonGroup, QPushButton,
    QGridLayout, QVBoxLayout, QHBoxLayout, QScrollArea, QDialog,
    QDialogButtonBox, QSizePolicy, QToolButton
)
from PySide6.QtCore import Qt, QTimer
from PySide6.QtGui import QIcon
import ctypes
from PySide6.QtCore import Qt, Signal 
from core_data import CausaData
from PySide6.QtWidgets import QDialog, QVBoxLayout     
from PySide6.QtWidgets import QMessageBox
import html  
from html import unescape
from PySide6.QtGui import QFont, QPainter, QTextCharFormat, QAction, QTextDocument 


myappid = 'com.miempresa.miproducto.1.0'  # Identificador único
ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)

###############################################################################
# Funciones auxiliares
###############################################################################
class ZoomableTextEdit(QTextEdit):
    zoomChanged = Signal(int)           # porcentaje (int)

    def __init__(self, *a, **kw):
        super().__init__(*a, **kw)
        self._steps = 0                 # cantidad de “clics” de zoom
        self._factor_per_step = 1.10    # 10 % por paso
        self.setLineWrapMode(QTextEdit.WidgetWidth)

    def wheelEvent(self, ev):
        if ev.modifiers() & Qt.ControlModifier:
            delta = ev.angleDelta().y() or ev.pixelDelta().y()
            if delta == 0:
                return

            step = 1 if delta > 0 else -1
            self._steps += step

            # usar zoomIn/zoomOut para que el texto se re‑envuelva
            if step > 0:
                self.zoomIn(1)
            else:
                self.zoomOut(1)

            # nuevo porcentaje
            pct = round((self._factor_per_step ** self._steps) * 100)
            self.zoomChanged.emit(pct)
            ev.accept()
        else:
            super().wheelEvent(ev)

def get_resource_path(relative_path):
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

UNIDADES = (
    'cero', 'uno', 'dos', 'tres', 'cuatro', 'cinco',
    'seis', 'siete', 'ocho', 'nueve', 'diez', 'once',
    'doce', 'trece', 'catorce', 'quince', 'dieciséis',
    'diecisiete', 'dieciocho', 'diecinueve', 'veinte',
    'veintiuno', 'veintidós', 'veintitrés', 'veinticuatro',
    'veinticinco', 'veintiséis', 'veintisiete', 'veintiocho',
    'veintinueve'
)
DECENAS = (
    'treinta', 'cuarenta', 'cincuenta', 'sesenta',
    'setenta', 'ochenta', 'noventa'
)
CENTENAS = (
    'cien', 'doscientos', 'trescientos', 'cuatrocientos',
    'quinientos', 'seiscientos', 'setecientos', 'ochocientos',
    'novecientos'
)

    #  helpers internos de RTF – pégalos encima de copy_to_clipboard
_rx_tag      = re.compile(r'<(/?)(b|strong|i|em|u|p)(?:\s+[^>]*)?>', re.I)
_rx_p_align  = re.compile(r'text-align\s*:\s*(left|right|center|justify)', re.I)

def _html_to_rtf_fragment(html: str) -> str:
    """
    Convierte un HTML muy sencillo (p, b/strong, i/em, u)
    a la secuencia RTF equivalente.
    """
    rtf = []
    stack = []              # para llevar el estado <b>, <i>, <u>

    pos = 0
    for m in _rx_tag.finditer(html):
        text = html[pos:m.start()]
            # --- texto normal (escapar) --------------------------------------
        text = (text.replace('\\', r'\\')
                    .replace('{',  r'\{')
                    .replace('}',  r'\}')
                    .replace('&nbsp;', ' '))
        rtf.append(text)
        pos = m.end()
        closing, tag = m.group(1), m.group(2).lower()
        if tag == 'p':                                     # <p …>
            if closing:                    # </p>
                rtf.append(r'\par ')
            else:                          # <p …>
                style = m.group(0)
                alg   = 'justify'          # por defecto
                ma = _rx_p_align.search(style)
                if ma:
                    alg = ma.group(1).lower()
                rtf.append(r'\pard')
                rtf.append({'left': r'\ql ',
                            'right': r'\qr ',
                            'center': r'\qc ',
                            'justify': r'\qj '}[alg])
        elif tag in ('b', 'strong'):
            rtf.append(r'\b0 ' if closing else r'\b ')
        elif tag in ('i', 'em'):
            rtf.append(r'\i0 ' if closing else r'\i ')
        elif tag == 'u':
            rtf.append(r'\ulnone ' if closing else r'\ul ')

        # resto del texto tras la última etiqueta
    tail = html[pos:]
    tail = (tail.replace('\\', r'\\')
                .replace('{',  r'\{')
                .replace('}',  r'\}')
                .replace('&nbsp;', ' '))
    rtf.append(tail)

    return ''.join(rtf)

def numero_a_letras(num: int) -> str:
    if num < 0:
        return "menos " + numero_a_letras(abs(num))
    if num <= 29:
        return UNIDADES[num]
    if num < 100:
        dec = (num // 10) - 3
        uni = num % 10
        letra_decena = DECENAS[dec]
        return letra_decena if uni == 0 else f"{letra_decena} y {UNIDADES[uni]}"
    if num < 1000:
        cent = (num // 100) - 1
        resto = num % 100
        if num == 100:
            return "cien"
        return CENTENAS[cent] if resto == 0 else f"{CENTENAS[cent]} {numero_a_letras(resto)}"
    if num < 10000:
        mil = num // 1000
        resto = num % 1000
        prefix = "mil" if mil == 1 else f"{numero_a_letras(mil)} mil"
        return prefix if resto == 0 else f"{prefix} {numero_a_letras(resto)}"
    return str(num)

def _sanitize_html_italic_only(html_raw: str) -> str:
    """
    Limpia el HTML e IMPIDE negrita/subrayado.
    Si no hay cursiva explícita, envuelve todo en <i>…</i>.
    """
    import re, html

    # Nos quedamos sólo con el <body> (igual que antes)
    m = re.search(r'<body[^>]*>(.*?)</body>', html_raw, flags=re.I | re.S)
    if m: html_raw = m.group(1)

    # 1) fuera <b>, </b>, <strong>, </strong>, <u>, </u>, y spans con font-weight
    html_raw = re.sub(r'</?(b|strong|u)[^>]*>', '', html_raw, flags=re.I)
    html_raw = re.sub(
        r'<span[^>]*style="[^"]*font-weight\s*:\s*(?:bold|700)[^"]*"[^>]*>',
        '',
        html_raw,
        flags=re.I
    )
    html_raw = re.sub(r'</span>', '', html_raw, flags=re.I)

    # 2) Nos quedamos *solo* con <i>/<em> –los demás tags fuera–
    #    (primero <em>→<i> para unificar)
    html_raw = re.sub(r'</?em>', lambda m:'<i>' if m.group(0)[1]!='/' else '</i>', html_raw, flags=re.I)
    html_raw = re.sub(r'</?(?!i\b)[a-z][^>]*>', '', html_raw)  # quita todo salvo <i>

    # 3) compactamos espacios/entidades raras
    html_raw = re.sub(r'(\r\n|\r|\n|&nbsp;|\u2028|\u2029)', ' ', html_raw)
    html_raw = re.sub(r'\s+', ' ', html_raw).strip()

    # 4) Si NO quedó ningún <i>…</i>, lo rodeamos completo
    if '<i>' not in html_raw.lower():
        html_raw = f"<i>{html.escape(html_raw)}</i>"

    return html_raw

def _sanitize_html(html_raw: str) -> str:
    """
    Devuelve SOLO el fragmento que estaba dentro de <body>,
    manteniendo <b>, <i>, <u> y quitando todo estilo / saltos raros.
    """
    import re, html

    # A)  ───── EXTRAEMOS SOLO <body> … </body> ─────
    m = re.search(r'<body[^>]*>(.*?)</body>', html_raw, flags=re.I | re.S)
    if m:
        html_raw = m.group(1)
    # (si por algún motivo no hay <body>, seguimos con lo que venga)

    # B)  ───── A partir de aquí van los pasos que ya tenías ─────
    # a) <strong>/<em> → <b>/<i>
    html_raw = re.sub(r'</?strong>', lambda m: '<b>' if m.group(0)[1] != '/' else '</b>', html_raw, flags=re.I)
    html_raw = re.sub(r'</?em>',     lambda m: '<i>' if m.group(0)[1] != '/' else '</i>', html_raw, flags=re.I)

    # b) <span style="font-weight:...">…</span> → <b>…</b>
    html_raw = re.sub(
        r'<span[^>]*style="[^"]*font-weight\s*:\s*(?:bold|700)[^"]*"[^>]*>(.*?)</span>',
        r'<b>\1</b>',
        html_raw,
        flags=re.I | re.S
    )

    # c) quitamos atributos style, class, dir, lang…
    html_raw = re.sub(r'\s*(style|class|dir|lang)="[^"]*"', '', html_raw, flags=re.I)

    # d) quitamos cualquier <span> remanente
    html_raw = re.sub(r'</?span[^>]*>', '', html_raw, flags=re.I)

    # d-bis) fuera <br>
    html_raw = re.sub(r'(?i)<br\s*/?>', ' ', html_raw)

    # d-ter) fuera párrafos vacíos de Qt
    html_raw = re.sub(
        r'<p[^>]*-qt-paragraph-type:empty[^>]*>\s*(<br\s*/?>)?\s*</p>',
        ' ',
        html_raw,
        flags=re.I
    )

    # e) limpia saltos y nbsp
    html_raw = re.sub(
        r'(\r\n|\r|\n|&#10;|&#13;|\u2028|\u2029|&nbsp;)',
        ' ',
        html_raw
    )

    # f) compacta espacios
    html_raw = re.sub(r'\s+', ' ', html_raw).strip()

    # g) si quedó un único <p> que envuelve todo, lo quitamos
    if html_raw.lower().startswith('<p') and html_raw.lower().endswith('</p>'):
        html_raw = re.sub(r'^<p[^>]*>|</p>$', '', html_raw, flags=re.I).strip()

    return html.unescape(html_raw)

def obtener_fecha_en_letras():
    fecha_actual = datetime.now()
    dia = fecha_actual.day
    mes_numero = fecha_actual.month
    anio = fecha_actual.year
    dia_letras = numero_a_letras(dia)
    anio_letras = numero_a_letras(anio)
    meses = {
        1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
        5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
        9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
    }
    mes_str = meses.get(mes_numero, '')
    return f"{dia_letras} de {mes_str} de {anio_letras}"

def format_list_for_sentence(items):
    """Separa con comas y añade ' y ' antes del último elemento."""
    items = [i for i in items if i.strip()]
    if len(items) == 0:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} y {items[1]}"
    return f"{', '.join(items[:-1])} y {items[-1]}"

def format_list_with_semicolons(items):
    """Separa con ';' y añade '; y ' antes del último elemento."""
    items = [i.strip() for i in items if i.strip()]
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]}; y {items[1]}"
    return "; ".join(items[:-1]) + f"; y {items[-1]}"

def numero_romano(n: int) -> str:
    romanos = [
        "1", "2", "3", "4", "5", "6", "7", "8", "9", "10",
        "11", "12", "13", "14", "15", "16", "17", "18", "19", "20"
    ]
    return romanos[n-1] if 1 <= n <= len(romanos) else str(n)

ORDINALES_HECHOS = [
    "Primer",
    "Segundo",
    "Tercer",
    "Cuarto",
    "Quinto",
    "Sexto",
    "Séptimo",
    "Octavo",
    "Noveno",
    "Décimo",
    "Undécimo",
    "Duodécimo",
    "Decimotercero",
    "Decimocuarto",
    "Decimoquinto"
]


class CollapsibleGroup(QWidget):
    """Widget con un botón tipo sección desplegable."""

    def __init__(self, title: str, parent=None):
        super().__init__(parent)

        self._title = title
        self.toggle_button = QToolButton(checkable=True)
        self.toggle_button.setChecked(False)
        self.toggle_button.setToolButtonStyle(Qt.ToolButtonTextOnly)
        self.toggle_button.setArrowType(Qt.NoArrow)
        self.toggle_button.setStyleSheet(
            "QToolButton { border: none; padding:4px; font-weight:bold; }"
        )
        self.toggle_button.clicked.connect(self._on_toggled)

        self.content_area = QWidget()
        self.content_area.setVisible(False)
        self.content_area.setStyleSheet("padding:4px;")

        lay = QVBoxLayout(self)
        lay.setSpacing(0)
        lay.setContentsMargins(0, 0, 0, 0)
        lay.addWidget(self.toggle_button)
        lay.addWidget(self.content_area)

        # Estilo general para resaltar cada bloque
        self.setStyleSheet(
            "CollapsibleGroup {"
            "border:1px solid #ccc;"
            "border-radius:5px;"
            "background-color:#f7f7f7;"
            "margin-top:4px;"
            "}"
        )

        self._update_button_text()

    def _update_button_text(self):
        arrow = "▼" if self.toggle_button.isChecked() else "▶"
        self.toggle_button.setText(f"{arrow} {self._title}")

    def _on_toggled(self, checked: bool):
        self._update_button_text()
        self.content_area.setVisible(checked)


class SentenciaWidget(QWidget):
    def __init__(self, data: CausaData, parent=None):
        super().__init__(parent)

        # listas que irán guardando los widgets dinámicos
        self.imputados: list = []
        self.hechos: list = []

        # para resaltar cambios en la plantilla
        self._prev_plain = ""
        # widgets que resaltarán secciones de la plantilla al enfocarse
        self._focus_highlight_map = {}

        self.data = data
        # ───────────────────────────────────────────────
        # Localidad
        self.var_localidad = QLineEdit()
        self.var_localidad.setText(self.data.localidad)
        self.var_localidad.textChanged.connect(
            lambda t: setattr(self.data, "localidad", t.strip())
        )

        # Carátula
        self.var_caratula = QLineEdit()
        self.var_caratula.setText(self.data.caratula)
        self.var_caratula.textChanged.connect(
            lambda t: setattr(self.data, "caratula", t.strip())
        )

        # Tribunal  (Combo editable)
        self.var_tribunal = QComboBox()
        self.var_tribunal.setEditable(True)
        lista_tribunales = [  
            "Cámara en lo Criminal y Correccional de Primera Nominación", 
            "Cámara en lo Criminal y Correccional de Segunda Nominación", 
            "Cámara en lo Criminal y Correccional de Tercera Nominación", 
            "Cámara en lo Criminal y Correccional de Cuarta Nominación", 
            "Cámara en lo Criminal y Correccional de Quinta Nominación", 
            "Cámara en lo Criminal y Correccional de Sexta Nominación", 
            "Cámara en lo Criminal y Correccional de Séptima Nominación", 
            "Cámara en lo Criminal y Correccional de Octava Nominación", 
            "Cámara en lo Criminal y Correccional de Novena Nominación", 
            "Cámara en lo Criminal y Correccional de Décima Nominación", 
            "Cámara en lo Criminal y Correccional de Onceava Nominación", 
            "Cámara en lo Criminal y Correccional de Doceava Nominación", 
            "Juzgado de Control en lo Penal Económico", 
            "Juzgado de Control y Faltas N° 2", 
            "Juzgado de Control y Faltas N° 3", 
            "Juzgado de Control y Faltas N° 4", 
            "Juzgado de Control y Faltas N° 5", 
            "Juzgado de Control en Violencia de Género y Familiar N° 1", 
            "Juzgado de Control en Violencia de Género y Familiar N° 2", 
            "Juzgado de Control y Faltas N° 7", 
            "Juzgado de Control y Faltas N° 8", 
            "Juzgado de Control y Faltas N° 9", 
            "Juzgado de Control y Faltas N° 10",
            "Juzgado de Control y Faltas N° 11", 
            "Juzgado de Control de Lucha contra el Narcotráfico"
        ]
        self.var_tribunal.addItems(lista_tribunales)
        if self.data.tribunal and self.data.tribunal not in lista_tribunales:
            self.var_tribunal.addItem(self.data.tribunal)
        self.var_tribunal.setCurrentText(self.data.tribunal)
        self.var_tribunal.currentTextChanged.connect(
            lambda t: setattr(self.data, "tribunal", t.strip())
        )
        self.var_tribunal.currentTextChanged.connect(self.actualizar_plantilla)
        self.install_focus_highlight(self.var_tribunal, lambda: self.var_tribunal.currentText())
        if self.var_tribunal.lineEdit():
            self.install_focus_highlight(self.var_tribunal.lineEdit(), lambda: self.var_tribunal.currentText())

        # Sala
        self.var_sala = QComboBox()
        self.var_sala.setEditable(True)
        salas_opciones = [
            "Sala OGA 1 del MOPLO", "Sala OGA 2 del MOPLO", "Sala OGA 3 del MOPLO",
            "Sala OGA 4 del MOPLO", "Sala OGA 5 del MOPLO", "Sala OGA 6 del MOPLO",
            "Sala OGA 7 del MOPLO", "Sala OGA 8 del MOPLO", "Sala OGA 9 del MOPLO",
            "Sala OGA 10 del MOPLO",
            "Sala de audiencias de la Cámara en lo Criminal y Correccional"
        ]
        self.var_sala.addItems(salas_opciones)
        if self.data.sala and self.data.sala not in salas_opciones:
            self.var_sala.addItem(self.data.sala)
        self.var_sala.setCurrentText(self.data.sala)
        self.var_sala.currentTextChanged.connect(
            lambda t: setattr(self.data, "sala", t.strip())
        )
        self.install_focus_highlight(self.var_sala, lambda: self.var_sala.currentText())
        if self.var_sala.lineEdit():
            self.install_focus_highlight(self.var_sala.lineEdit(), lambda: self.var_sala.currentText())

        # ───────────────────────────────────────────────
        # 2) INTERVINIENTES
        # ───────────────────────────────────────────────
        # Juez + sexo
        self.var_juez = QLineEdit()
        self.var_juez.setText(self.data.juez_nombre)
        self.var_juez.textChanged.connect(
            lambda t: setattr(self.data, "juez_nombre", t.strip())
        )
        self.rb_juez_m = QRadioButton("M")
        self.rb_juez_f = QRadioButton("F")
        if self.data.juez_sexo == "F":
            self.rb_juez_f.setChecked(True)
        else:
            self.rb_juez_m.setChecked(True)
        self.rb_juez_m.toggled.connect(
            lambda chk: chk and setattr(self.data, "juez_sexo", "M")
        )
        self.rb_juez_f.toggled.connect(
            lambda chk: chk and setattr(self.data, "juez_sexo", "F")
        )
        self.cargo_juez = self.data.juez_cargo  # “juez” o “vocal”

        # Fiscal + sexo
        self.var_fiscal = QLineEdit()
        self.var_fiscal.setText(self.data.fiscal_nombre)
        self.var_fiscal.textChanged.connect(
            lambda t: setattr(self.data, "fiscal_nombre", t.strip())
        )
        self.rb_fiscal_m = QRadioButton("M")
        self.rb_fiscal_f = QRadioButton("F")
        if self.data.fiscal_sexo == "F":
            self.rb_fiscal_f.setChecked(True)
        else:
            self.rb_fiscal_m.setChecked(True)
        self.rb_fiscal_m.toggled.connect(
            lambda chk: chk and setattr(self.data, "fiscal_sexo", "M")
        )
        self.rb_fiscal_f.toggled.connect(
            lambda chk: chk and setattr(self.data, "fiscal_sexo", "F")
        )

        # Día de audiencia
        self.var_dia_audiencia = QLineEdit()
        self.var_dia_audiencia.setText(self.data.fecha_audiencia)
        self.var_dia_audiencia.textChanged.connect(
            lambda t: setattr(self.data, "fecha_audiencia", t.strip())
        )

        # N° imputados
        self.var_num_imputados = QSpinBox()
        self.var_num_imputados.setRange(1, 10)
        self.var_num_imputados.setValue(self.data.n_imputados)
        self.var_num_imputados.valueChanged.connect(
            lambda v: setattr(self.data, "n_imputados", v)
        )
        # Caso de violencia familiar / género
        self.var_caso_vf = QComboBox()
        self.var_caso_vf.addItems(
            ["No", "violencia de género", "violencia de género doméstica", "violencia familiar"]
        )

        self.var_caso_vf.currentTextChanged.connect(
            lambda t: setattr(self.data, "caso_vf", t.strip())
        )
        # ───────────────────────────────────────────────
        # 3) DATOS ADICIONALES (los nuevos)
        # ───────────────────────────────────────────────
        self.var_sujeto_eventual = QLineEdit()
        self.var_sujeto_eventual.setText(self.data.sujeto_eventual)
        self.var_sujeto_eventual.textChanged.connect(
            lambda t: setattr(self.data, "sujeto_eventual", t.strip())
        )

        self.var_manifestacion = QLineEdit()
        self.var_manifestacion.setText(self.data.manifestacion_sujeto)
        self.var_manifestacion.textChanged.connect(
            lambda t: setattr(self.data, "manifestacion_sujeto", t.strip())
        )

        self.var_victima = QLineEdit()
        self.var_victima.setText(self.data.victima)
        self.var_victima.textChanged.connect(
            lambda t: setattr(self.data, "victima", t.strip())
        )

        self.var_victima_plural = QComboBox()
        self.var_victima_plural.addItems(["Una", "Más"])
        self.var_victima_plural.setCurrentIndex(1 if self.data.victima_plural else 0)
        self.var_victima_plural.currentIndexChanged.connect(
            lambda i: setattr(self.data, "victima_plural", i == 1)
        )

        self.var_victima_manifestacion = QLineEdit()
        self.var_victima_manifestacion.setText(self.data.manifestacion_victima)
        self.var_victima_manifestacion.textChanged.connect(
            lambda t: setattr(self.data, "manifestacion_victima", t.strip())
        )
        # en __init__, junto al resto de los var_…
        self.var_prueba = ""                       # aquí guardaremos el HTML plano  
        self.var_pruebas_importantes = ""          # idem

        
        self.btn_prueba = QPushButton("Redactar pruebas")
        self.btn_prueba.clicked.connect(self.abrir_ventana_prueba)

        self.btn_pruebas_importantes = QPushButton("Redactar pruebas relevantes")
        self.btn_pruebas_importantes.clicked.connect(self.abrir_ventana_pruebas_importantes)

        # Alegatos (se guardan al cerrar el diálogo)
        self.var_alegato_fiscal = self.data.alegato_fiscal
        self.btn_alegato_fiscal = QPushButton("Redactar alegato fiscal")

        self.var_alegato_defensa = self.data.alegato_defensa
        self.btn_alegato_defensa = QPushButton("Redactar alegato de la defensa")

        # Calificación legal
        self.var_calificacion_legal = QComboBox()
        self.var_calificacion_legal.addItems(["Correcta", "Incorrecta"])
        self.var_calificacion_legal.setCurrentText(self.data.calif_legal)
        self.var_calificacion_legal.currentTextChanged.connect(
            lambda t: setattr(self.data, "calif_legal", t)
        )

        self.var_correccion_calif = QLineEdit()
        self.var_correccion_calif.setText(self.data.calif_correccion)
        self.var_correccion_calif.textChanged.connect(
            lambda t: setattr(self.data, "calif_correccion", t.strip())
        )

        # Términos potenciales
        self.var_uso_terminos_potenciales = QComboBox()
        self.var_uso_terminos_potenciales.addItems(["No", "Sí"])
        self.var_uso_terminos_potenciales.setCurrentIndex(
            1 if self.data.usa_potenciales else 0
        )
        self.var_uso_terminos_potenciales.currentIndexChanged.connect(
            lambda i: setattr(self.data, "usa_potenciales", i == 1)
        )

        # Decomiso
        self.var_decomiso_option = QComboBox()
        self.var_decomiso_option.addItems(["No", "Sí"])
        self.var_decomiso_option.setCurrentIndex(1 if self.data.decomiso_si else 0)
        self.var_decomiso_option.currentIndexChanged.connect(
            lambda i: setattr(self.data, "decomiso_si", i == 1)
        )

        self.var_decomiso_text = QLineEdit()
        self.var_decomiso_text.setText(self.data.decomiso_texto)
        self.var_decomiso_text.textChanged.connect(
            lambda t: setattr(self.data, "decomiso_texto", t.strip())
        )

        # Restricción de acercamiento
        self.var_restriccion_option = QComboBox()
        self.var_restriccion_option.addItems(["No", "Sí"])
        self.var_restriccion_option.setCurrentIndex(1 if self.data.restriccion_si else 0)
        self.var_restriccion_option.currentIndexChanged.connect(
            lambda i: setattr(self.data, "restriccion_si", i == 1)
        )

        self.var_restriccion_text = QLineEdit()
        self.var_restriccion_text.setText(self.data.restriccion_texto)
        self.var_restriccion_text.textChanged.connect(
            lambda t: setattr(self.data, "restriccion_texto", t.strip())
        )

        self.var_num_hechos = QSpinBox()
        self.var_num_hechos.setRange(1, 15)
        self.var_num_hechos.setValue(len(self.data.hechos) or 1)
        self.var_num_hechos.valueChanged.connect(
            lambda v: setattr(self.data, "num_hechos", v)
        )

        # ───────────────────────────────────────────────
        # RESTO DEL CÓDIGO (setup_ui, conexiones, etc.)
        # ───────────────────────────────────────────────
        self.setup_ui()
        self.setup_connections()

        self.data.apply_to_sentencia(self)     # carga modelo
        self.update_imputados_section()        # crea pestañas imputados
        self.update_hechos_section()           # crea pestañas hechos
        self.actualizar_plantilla()            # ya existen ambas listas

    def _update_zoom_label(self, percent: int):
        self.lbl_zoom.setText(f"ZOOM {percent}%")
        self.lbl_zoom.setVisible(True)           # ¡mostrar!
        self._hide_zoom_timer.start(1000)        # se ocultará en 1 s

    def _highlight_diff(self, old_text: str, new_text: str) -> None:
        """Resalta en amarillo los fragmentos modificados."""
        from difflib import SequenceMatcher
        from PySide6.QtGui import QTextCursor, QTextCharFormat, QBrush

        # Cursor independiente para no mover la posición del usuario
        cursor = QTextCursor(self.texto_plantilla.document())

        # Limpia resaltados previos
        fmt_clear = QTextCharFormat()
        fmt_clear.setBackground(QBrush(Qt.transparent))
        cursor.select(QTextCursor.Document)
        cursor.mergeCharFormat(fmt_clear)
        self._clear_highlight_timer.stop()

        matcher = SequenceMatcher(None, old_text, new_text)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal' or j1 == j2:
                continue
            cursor.setPosition(j1)
            cursor.setPosition(j2, QTextCursor.KeepAnchor)
            fmt = QTextCharFormat()
            fmt.setBackground(QBrush(Qt.yellow))
            cursor.mergeCharFormat(fmt)

        # El resaltado se limpiará automáticamente tras 3 segundos
        self._clear_highlight_timer.start(3000)

    def _highlight_section_text(self, text: str) -> None:
        """Resalta todas las apariciones de ``text`` en la plantilla."""
        from PySide6.QtGui import QTextCursor, QTextCharFormat, QBrush

        cursor = QTextCursor(self.texto_plantilla.document())

        fmt_clear = QTextCharFormat()
        fmt_clear.setBackground(QBrush(Qt.transparent))
        cursor.select(QTextCursor.Document)
        cursor.mergeCharFormat(fmt_clear)
        self._clear_highlight_timer.stop()

        if not text:
            return

        plain = self.texto_plantilla.toPlainText()
        text_lower = text.lower()
        pos = plain.lower().find(text_lower)
        fmt = QTextCharFormat()
        fmt.setBackground(QBrush(Qt.yellow))
        while pos != -1:
            cursor.setPosition(pos)
            cursor.setPosition(pos + len(text), QTextCursor.KeepAnchor)
            cursor.mergeCharFormat(fmt)
            pos = plain.lower().find(text_lower, pos + len(text))

        self._clear_highlight_timer.start(3000)

    def _clear_highlight(self) -> None:
        """Quita el resaltado preservando la posición de scroll."""
        from PySide6.QtGui import QTextCursor, QTextCharFormat, QBrush

        sb = self.texto_plantilla.verticalScrollBar()
        pos = sb.value()

        cursor = QTextCursor(self.texto_plantilla.document())
        fmt_clear = QTextCharFormat()
        fmt_clear.setBackground(QBrush(Qt.transparent))
        cursor.select(QTextCursor.Document)
        cursor.mergeCharFormat(fmt_clear)

        sb.setValue(pos)

    def toggle_cargo_juez(self):
        if self.cargo_juez == "juez":
            self.cargo_juez = "vocal"
        else:
            self.cargo_juez = "juez"
        self.boton_cargo_juez.setText(self.cargo_juez)
        self.actualizar_plantilla()

    def _toggle_bold(self, editor: QTextEdit):
        cursor = editor.textCursor()
        if not cursor.hasSelection():
            return                                # nada seleccionado
        fmt = QTextCharFormat()
        bold_now = cursor.charFormat().fontWeight() > QFont.Normal
        fmt.setFontWeight(QFont.Normal if bold_now else QFont.Bold)
        cursor.mergeCharFormat(fmt)


    @staticmethod
    def html_a_plano(html: str, mantener_saltos: bool = True) -> str:
        if not html:
            return ""

        doc = QTextDocument(); doc.setHtml(html)
        texto = doc.toPlainText()

        # → equivale a &nbsp; y &nbsp; finos (202F)
        texto = texto.replace("\u00A0", " ").replace("\u202F", " ")

        if not mantener_saltos:
            texto = texto.replace("\n", " ")

        return texto.strip()

    def install_focus_highlight(self, widget, text_getter):
        """Destaca la sección correspondiente al obtener foco."""
        widget.installEventFilter(self)
        self._focus_highlight_map[widget] = text_getter

    def eventFilter(self, obj, event):
        if event.type() == QEvent.FocusIn and obj in self._focus_highlight_map:
            try:
                text = self._focus_highlight_map[obj]()
            except Exception:
                text = ""
            self._highlight_section_text(text)
        return super().eventFilter(obj, event)
    
    def _rich_text_dialog_italic_only(self, title: str, initial_html: str, on_accept):
        dlg = QDialog(self); dlg.setWindowTitle(title); dlg.resize(650, 420)
        lay = QVBoxLayout(dlg)
        editor = QTextEdit()
        editor.setAcceptRichText(True)
        base_font = QFont("Times New Roman", 12)
        editor.setFont(base_font); editor.document().setDefaultFont(base_font)
        editor.setHtml(initial_html or "")
        lay.addWidget(editor)

        # SIN barra de negrita  ni atajos extras

        btn_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        lay.addWidget(btn_box)

        def _on_ok():
            raw = editor.toHtml()
            clean = _sanitize_html_italic_only(raw)
            on_accept(clean)
            clean = html.unescape(clean)   
            dlg.accept()

        btn_box.accepted.connect(_on_ok)
        btn_box.rejected.connect(dlg.reject)
        dlg.exec()


    def _rich_text_dialog(self, title: str, initial_html: str, on_accept_callback):
        """
        Abre un diálogo con un QTextEdit que:
        • Acepta pegar desde Word/web conservando <b>/<i>/<u>.
        • Muestra Times New Roman 12 pt mientras editás.
        • Tiene botón/atajo Ctrl+B para alternar negrita.
        • Al aceptar ➜ limpia HTML con _sanitize_html(...) y lo
            entrega al callback.
        """
        dlg = QDialog(self)
        dlg.setWindowTitle(title)
        dlg.resize(650, 420)

        lay_top   = QVBoxLayout(dlg)
        toolbar   = QHBoxLayout()
        editor    = QTextEdit()

        # --- apariencia del editor -----------------------------------
        base_font = QFont("Times New Roman", 12)
        editor.setFont(base_font)
        editor.document().setDefaultFont(base_font)
        editor.setAcceptRichText(True)
        editor.setHtml(initial_html or "")          # muestra lo que ya tenías

        # --- botón “B” (negrita) --------------------------------------
        btn_bold = QPushButton("B")
        btn_bold.setCheckable(True)
        btn_bold.setFixedWidth(32)
        btn_bold.setStyleSheet("font-weight:bold;")
        btn_bold.clicked.connect(lambda: self._toggle_bold(editor))
        toolbar.addWidget(btn_bold)
        toolbar.addStretch()
        lay_top.addLayout(toolbar)

        # --- atajo Ctrl+B ---------------------------------------------
        editor.addAction(QAction(self, shortcut="Ctrl+B",
                                triggered=lambda: self._toggle_bold(editor)))

        lay_top.addWidget(editor)

        # --- OK / Cancel ----------------------------------------------
        btn_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel,
                                parent=dlg)
        lay_top.addWidget(btn_box)

        # Cuando aprietan OK…
        def _on_ok():
            raw_html = editor.toHtml()
            clean    = _sanitize_html(raw_html)
            clean = html.unescape(clean)   
            # devolvemos HTML *limpio* al método que llamó
            on_accept_callback(clean)
            dlg.accept()

        btn_box.accepted.connect(_on_ok)
        btn_box.rejected.connect(dlg.reject)
        dlg.exec()

    # ──────────────────────────────────────────────────────────────
    def _abrir_editor_rich_sobre_lineedit(self, qle: QLineEdit, titulo: str):
        """Abre el rich-text dialog, guarda HTML limpio en qle.property('html')
        y muestra un resumen plano (máx. 200 c) en el QLineEdit."""
        html_inicial = qle.property("html") or qle.text()
        self._rich_text_dialog(
            titulo,
            html_inicial,
            lambda html_limpio: (
                qle.setProperty("html", html_limpio),
                qle.setText(QTextDocument(html_limpio).toPlainText()[:200]),
                self.actualizar_plantilla()
            )
        )

    def _rich_text_dialog_no_bold(self, title: str, initial_html: str, on_accept):
        """Editor rico SIN negrita (ni botón ni atajo)."""
        from PySide6.QtWidgets import QDialog, QVBoxLayout, QTextEdit, QDialogButtonBox
        import html, re

        dlg = QDialog(self); dlg.setWindowTitle(title); dlg.resize(650, 420)
        lay = QVBoxLayout(dlg)

        editor = QTextEdit()
        base_font = QFont("Times New Roman", 12)
        editor.setFont(base_font); editor.document().setDefaultFont(base_font)
        editor.setAcceptRichText(True)
        editor.setHtml(initial_html or "")
        lay.addWidget(editor)

        btn_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        lay.addWidget(btn_box)

        # ── helper local: quita toda negrita ──────────────────────────────
        def _remove_bold(raw_html: str) -> str:
            # 1) elimina <b>/<strong>
            raw_html = re.sub(r'</?(b|strong)[^>]*>', '', raw_html, flags=re.I)
            # 2) elimina spans con font-weight:bold/700
            raw_html = re.sub(
                r'<span[^>]*style="[^"]*font-weight\s*:\s*(?:bold|700)[^"]*"[^>]*>(.*?)</span>',
                r'\1', flags=re.I | re.S, string=raw_html)
            return html.unescape(raw_html).strip()

        def _on_ok():
            clean_html = _sanitize_html(editor.toHtml())          # tu sanitizador normal
            clean_html = _remove_bold(clean_html)                 # …pero sin negrita
            on_accept(clean_html)
            dlg.accept()

        btn_box.accepted.connect(_on_ok)
        btn_box.rejected.connect(dlg.reject)
        dlg.exec()


    # ──────────────────────────────────────────────────────────────
    def copiar_sentencia(self, te: QTextEdit) -> None:
        """
        Copia **texto plano**, **RTF** (Times 12 + estilos) y **HTML**.
        Word elegirá el HTML (con negritas/alineaciones), pero el RTF
        queda por si lo necesita otro programa.
        """
        from PySide6.QtCore    import QMimeData
        from PySide6.QtWidgets import QApplication
        from PySide6.QtGui     import QClipboard

        # ---------- 1) texto sin formato --------------------------------------
        plain_text = te.toPlainText().strip()

        # ---------- 2) HTML limpio + CSS Times 12 -----------------------------
        # ---------- 2) HTML limpio + CSS Times 12, alineado a la izquierda -----------------
        basic_html = te.toHtml()
        # quitamos tamaños en línea para aplicar el nuestro
        basic_html = re.sub(r'font-size\s*:[^;"]+;?', '', basic_html, flags=re.I)
        # forzamos que cada párrafo venga con align="left"
        basic_html = re.sub(r'<p\b', '<p align="justify"', basic_html, flags=re.I)

        html_full = (
            "<!DOCTYPE html><html><head><meta charset='UTF-8'>"
            "<style>"
            "body,p{font-family:'Times New Roman',serif;"
            "font-size:12pt;line-height:1.0;margin:0;"
            "text-align:justify;}"
            "</style></head><body><!--StartFragment-->"
            + basic_html +
            "<!--EndFragment--></body></html>"
        )

        # ---------- 3) RTF (Times 12, alineación + b/i/u) ----------------------
        rtf_paragraphs = []
        for para_html in re.findall(r'<p[^>]*>.*?</p>', basic_html, flags=re.S|re.I):
            rtf_paragraphs.append(_html_to_rtf_fragment(para_html))

        rtf_content = (
            r"{\rtf1\ansi\deff0"
            r"{\fonttbl{\f0 Times New Roman;}}"
                        r"\fs24 " + ''.join(rtf_paragraphs) + "}"
        )
        # sustituimos cualquier \qj (justificado) por \ql (alineado a la izquierda)
        rtf_content = rtf_content.replace(r'\ql ', r'\qj ')
        

        # ---------- 4) al portapapeles (HTML lo dejamos el último) ------------
        mime = QMimeData()
        mime.setText(plain_text)
        mime.setData("text/rtf", rtf_content.encode("utf-8"))
        mime.setHtml(html_full)
        QApplication.clipboard().setMimeData(mime, QClipboard.Clipboard)

    def _flatten_inline(self, html_raw: str) -> str:
        """
        Convierte un blo-HTML en inline-HTML:

        • quita <p>, </p>, <div>, </div>, <br>
        • elimina separadores U+2028/U+2029, NBSP, \\r, \\n
        • colapsa espacios consecutivos
        """
        import re, html
        print(repr(html_raw))
        # A) fuera <p>, </p>, <div>, </div>
        h = re.sub(r'</?p[^>]*>',  ' ', html_raw, flags=re.I)
        h = re.sub(r'</?div[^>]*>', ' ', h,        flags=re.I)

        # B) fuera <br> y variantes
        h = re.sub(r'(?i)<br\s*/?>', ' ', h)

        # C) fuera saltos ocultos y nbsp
        h = re.sub(r'(\r\n|\r|\n|&#10;|&#13;|\u2028|\u2029|&nbsp;)', ' ', h)

        # D) compactar espacios
        h = re.sub(r'\s+', ' ', h).strip()

        return html.unescape(h)
    
    def abrir_ventana_alegato_fiscal(self):
        self._rich_text_dialog(
            "Escribir alegato fiscal",
            self.var_alegato_fiscal,
            lambda h: (
                setattr(self, "var_alegato_fiscal", self._flatten_inline(h)),
                self.actualizar_plantilla()
            )
        )

    def _guardar_alegato_fiscal(self, texto, dlg):
        self.var_alegato_fiscal = texto.strip()
        dlg.accept()
        self.actualizar_plantilla()

    def abrir_ventana_alegato_defensa(self):
        self._rich_text_dialog(
            "Editar el alegato de la defensa",
            self.var_alegato_defensa,
            lambda h: (
                setattr(self, "var_alegato_defensa", self._flatten_inline(h)),
                self.actualizar_plantilla()
            )
        )

    def _guardar_alegato_defensa(self, texto, dlg):
        self.var_alegato_defensa = texto.strip()
        dlg.accept()
        self.actualizar_plantilla()
    
    def abrir_ventana_prueba(self):
        self._rich_text_dialog(
            "Agregar prueba",
            self.var_prueba,
            lambda h: (
                setattr(self, "var_prueba", self._flatten_inline(h)),     
                self.actualizar_plantilla()
            )
        )

    def _guardar_prueba(self, texto, dlg):
        self.var_prueba = texto.strip()
        dlg.accept()
        self.actualizar_plantilla()
    
    def abrir_ventana_pruebas_importantes(self):
        self._rich_text_dialog(
            "Agregar pruebas relevantes",
            self.var_pruebas_importantes,
            lambda h: (
                setattr(self, "var_pruebas_importantes", self._flatten_inline(h)),
                self.actualizar_plantilla()
            )
        )

    def guardar_pruebas_importantes(self, texto, dlg):
        self.var_pruebas_importantes = texto.strip()
        dlg.accept()
        self.actualizar_plantilla()

    def _guardar_html_lineedit(self, qlineedit, html):
        """Guarda `html` aplanado en property('html') y preview plano en el LineEdit."""
        h = self._flatten_inline(html)
        qlineedit.setProperty("html", h)
        qlineedit.setText(QTextDocument(h).toPlainText()[:200])
        self.actualizar_plantilla()


    
    def abrir_ventana_descripcion(self, idx):
        """Abre el editor rich-text para la descripción del hecho #idx."""
        qle = self.hechos[idx]["descripcion"]
        html_inicial = qle.property("html") or qle.text()
        self._rich_text_dialog(
            f"Editar descripción del suceso #{idx+1}",
            html_inicial,
            lambda h: (
                self._guardar_html_lineedit(qle, h),
                self.actualizar_plantilla()
            )
        )


    def abrir_ventana_datos(self, idx):
        """Edita datos personales (#idx) usando el diálogo sin negrita."""
        qle = self.imputados[idx]["datos"]
        html_inicial = qle.property("html") or qle.text()

        self._rich_text_dialog_no_bold(                        # ← uso del nuevo diálogo
            f"Editar datos personales – imputado #{idx+1}",
            html_inicial,
            lambda h: (
                self._guardar_html_lineedit(qle, h),           # ⟵ se guarda igual que antes
                self.actualizar_plantilla()
            )
        )

    def abrir_ventana_condiciones(self, idx):
        """Abre el editor rich-text para datos personales agregados del imputado #idx."""
        qle = self.imputados[idx]["condiciones"]
        html_inicial = qle.property("html") or qle.text()
        self._rich_text_dialog(
            f"Editar datos personales agregados – imputado #{idx+1}",
            html_inicial,
            lambda h: (
                self._guardar_html_lineedit(qle, h),
                self.actualizar_plantilla()
            )
        )
    
    def abrir_ventana_pautas(self, idx):
        """Abre el editor rich-text para datos personales agregados del imputado #idx."""
        qle = self.imputados[idx]["pautas"]
        html_inicial = qle.property("html") or qle.text()
        self._rich_text_dialog(
            f"Editar pautas de conducta – imputado #{idx+1}",
            html_inicial,
            lambda h: (
                self._guardar_html_lineedit(qle, h),
                self.actualizar_plantilla()
            )
        )

    def abrir_ventana_decomiso(self):
        self._rich_text_dialog(
            "Editar texto de Decomiso",
            self.var_decomiso_text.property("html") or self.TEXTO_DECOMISO_DEFECTO,
            self._guardar_decomiso
        )

    def abrir_ventana_restriccion(self):
        self._rich_text_dialog(
            "Editar texto de Restricción de contacto",
            self.var_restriccion_text.property("html") or self.TEXTO_RESTRICCION_DEFECTO,
            self._guardar_restriccion
        )
    from PySide6.QtGui import QTextDocument

    def _guardar_decomiso(self, html_limpio: str):
        clean = html_limpio.strip()
        # 1) guardo HTML completo
        self.var_decomiso_text.setProperty("html", clean)
        # 2) genero preview plano en el QLineEdit
        doc = QTextDocument()
        doc.setHtml(clean)
        preview = doc.toPlainText().replace("\n", " ")[:200]
        self.var_decomiso_text.setText(preview)
        # 3) refresco la plantilla
        self.actualizar_plantilla()

    def _guardar_restriccion(self, html_limpio: str):
        clean = html_limpio.strip()
        self.var_restriccion_text.setProperty("html", clean)
        doc = QTextDocument()
        doc.setHtml(clean)
        preview = doc.toPlainText().replace("\n", " ")[:200]
        self.var_restriccion_text.setText(preview)
        self.actualizar_plantilla()


    TEXTO_RESTRICCION_DEFECTO = (
        "dadas las características y el contexto de la victimización acreditada en los presentes, considero adecuado imponer a XXX la prohibición de "
        "establecer cualquier clase de contacto o comunicación (verbal, telefónica, personalmente o por interpósita persona o por cualquier medio electrónico o informático, etc.) con XXX, "
        "hasta que la presente sentencia, luego de que quede firme, sea comunicada al Tribunal de Gestión Asociada del Fuero de Niñez, Adolescencia, Violencia Familiar y de Género de esta "
        "ciudad / a la Oficina Única de Violencia Familiar y de Genero de la ciudad de XXX, para que allí se adopten las medidas que pudieren corresponder al respecto (arts. 16 –inc. e– de "
        "la Ley Nacional 26485; y 3, 4, 9, 20, 21 –inc. e– y cc de la Ley Provincial 9283). Ello obedece a que, de conformidad a lo prescripto por el art. 16 –inc. e– de la Ley Nacional n° "
        "26485, los organismos del Estado, en cualquier procedimiento judicial, deben garantizar a las mujeres el derecho a recibir protección judicial urgente y preventiva cuando se encuentren "
        "amenazados o vulnerados cualquiera de los derechos enunciados en el artículo 3º de la misma ley, entre ellos la integridad psicológica, que podría verse afectada si el encartado "
        "procura lograr alguna clase de comunicación con ella desde su lugar de encierro. Sin embargo, el sistema jurídico no ofrece norma alguna que autorice a este tribunal a imponer una "
        "limitación de esa naturaleza en esta clase de condenas (privativas de la libertad de cumplimiento efectivo) para ser aplicada durante toda la extensión de la pena, lo que además "
        "tendría consecuencias no solo para el imputado, sino también para el libre albedrío de la damnificada. En ese sentido, entiendo que los Juzgados de Niñez, Adolescencia, Violencia "
        "Familiar y de Género son los únicos órganos jurisdiccionales con competencia en esta Provincia para tomar tales medidas de esa clase y que excedan lo meramente urgente. A su vez, "
        "aparece como razonable el pedido del acusado de retomar el contacto con su hijo menor de edad, también hijo de la nombrada, en función de lo previsto el art. 168 de la ley 24660; "
        "sin embargo, dado que ese niño ha sido testigo de los hechos sufridos por su madre a manos del encartado, resulta aconsejable que sea ese mismo juez especializado en las temáticas "
        "de la niñez, la violencia familiar y la violencia de género el que examine la conveniencia o no de que se materialice el pedido del imputado y, eventualmente, la modalidad con que se "
        "retome ese contacto. Por todo ello, se debe remitir copia de la presente sentencia al órgano judicial que preintervino en este conflicto para que, a partir de su recepción, adopte "
        "las medidas que pudieren corresponder a partir de ese momento, a dichos fines.\n\n"
    )

    TEXTO_DECOMISO_DEFECTO = (
        "corresponde ordenar el decomiso de XXX, "
        "en razón de que se trata de un instrumento/provecho/producto "
        "del delito, debido a que XXX (art. 23 del CP)."
    )

    def abrir_ventana_resuelvo(self):
        """
        Abre el diálogo de edición de Resuelvo, cargando —si existe—
        el HTML completo que guardamos en la property “html”.
        """
        html_actual = self.var_resuelvo.property("html")  # HTML íntegro
        if not html_actual:                               # primera vez
            html_actual = self.var_resuelvo.text()        # lo que haya

        self._rich_text_dialog(
            "Editar texto de Resuelvo",
            html_actual,                  # ← usa la variable, no vuelvas a leer
            self._guardar_resuelvo_html
        )


    def _guardar_resuelvo_html(self, html_limpio: str) -> None:
        clean = html_limpio.strip()

        # 1) guardo el HTML en la property oculta (para re-editar)
        self.var_resuelvo.setProperty("html", clean)

        # 2) genero un preview de 0‒200 c (solo texto)
        from PySide6.QtGui import QTextDocument
        doc = QTextDocument(); doc.setHtml(clean)
        preview = doc.toPlainText().replace("\n", " ")[:200]
        self.var_resuelvo.setText(preview)

        # 3) ***ACTUALIZO el modelo compartido***
        self.data.resuelvo_html = clean           # HTML íntegro
        self.data.resuelvo      = preview         # texto plano corto

        # 2-b  si la ventana de Trámites existe, actualiza su preview
        if getattr(self, "main_win", None):
            self.main_win.entry_resuelvo.setProperty("html", clean)
            self.main_win.entry_resuelvo.setText(preview)

        # 4) refresco plantilla / ventanas que dependan
        self.actualizar_plantilla()

    # ───────────────────── Fin bloque «Resuelvo» ─────────────────────
    def cargo_juez_en_mayusculas(self):
        cargo = self.boton_cargo_juez.text().lower()  # "juez" o "vocal"
        if cargo == "juez":
            if self.rb_juez_m.isChecked():
                return "EL JUEZ"
            else:
                return "LA JUEZA"
        else:
            if self.rb_juez_m.isChecked():
                return "EL VOCAL"
            else:
                return "LA VOCAL"

    def setup_ui(self):
        # -------------------------------
        main_layout = QHBoxLayout(self)
        # ------------------------------------------------------------------
        self.left_scroll = QScrollArea()
        self.left_scroll.setWidgetResizable(True)

        self.left_container = QWidget()
        self.left_layout    = QVBoxLayout(self.left_container)
        # Compactamos el espacio vertical en el panel izquierdo
        self.left_layout.setSpacing(2)
        # Mantenemos los campos pegados al inicio para evitar que se
        # distribuyan por toda la altura disponible
        self.left_layout.setAlignment(Qt.AlignTop)
        # Grupos colapsables al estilo del explorador de archivos
        self.general_group = CollapsibleGroup("Datos generales")
        self.imputados_group = CollapsibleGroup("Imputados")
        self.hechos_group = CollapsibleGroup("Hechos")
        self.extra_group = CollapsibleGroup("Otras opciones")

        for grp in (self.general_group, self.imputados_group,
                    self.hechos_group, self.extra_group):
            self.left_layout.addWidget(grp)

        self.left_scroll.setWidget(self.left_container)
        # Reducimos el ancho mínimo para que la interfaz sea visible incluso
        # cuando la ventana no está en pantalla completa
        self.left_container.setMinimumWidth(450)

        main_layout.addWidget(self.left_scroll, 2)

        # ------------------------------------------------------------------
        self.btn_generar_docx = QPushButton("Generar Word")
        self.btn_copiar       = QPushButton("Copiar sentencia")
        self.btn_ver_tramites = QPushButton("▶ Ver trámites")

        for b in (self.btn_generar_docx, self.btn_copiar, self.btn_ver_tramites):
            b.setFixedSize(180, 40)

        self.btn_ver_tramites.clicked.connect(self.abrir_tramites)
        self.btn_generar_docx.clicked.connect(self.generar_docx_con_html)
        self.btn_copiar.clicked.connect(lambda checked:
            self.copiar_sentencia(self.texto_plantilla)
        )

        # ------------------------------------------------------------------
        #  PANEL DERECHO : editor con zoom + barra de botones inferior
        # ------------------------------------------------------------------
        # 1) El QTextEdit que amplía/reduce con Ctrl+rueda
        self.texto_plantilla = ZoomableTextEdit()
        font = QFont("Times New Roman", 12)
        self.texto_plantilla.setFont(font)
        self.texto_plantilla.document().setDefaultFont(font)
        self.texto_plantilla.setReadOnly(True)
        self.texto_plantilla.setAlignment(Qt.AlignJustify)
        self.texto_plantilla.setStyleSheet("font-family:'Times New Roman';")
        # Un ancho más reducido permite que toda la aplicación se ajuste a
        # resoluciones menores sin necesidad de maximizar la ventana
        self.texto_plantilla.setMinimumWidth(500)
        opt = self.texto_plantilla.document().defaultTextOption()
        opt.setAlignment(Qt.AlignJustify)
        self.texto_plantilla.document().setDefaultTextOption(opt)
        self.texto_plantilla.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        right_layout = QVBoxLayout()

        # 2) Indicador de zoom (se muestra al hacer Ctrl+rueda y se oculta tras 1 s)
        self.lbl_zoom = QLabel("")
        self.lbl_zoom.setAlignment(Qt.AlignRight)
        self.lbl_zoom.setVisible(False)
        right_layout.addWidget(self.lbl_zoom)

        self._hide_zoom_timer = QTimer(self)
        self._hide_zoom_timer.setSingleShot(True)
        self._hide_zoom_timer.timeout.connect(lambda: self.lbl_zoom.setVisible(False))

        # Temporizador para eliminar el resaltado después de unos segundos
        self._clear_highlight_timer = QTimer(self)
        self._clear_highlight_timer.setSingleShot(True)
        self._clear_highlight_timer.timeout.connect(self._clear_highlight)

        # 3) El editor
        right_layout.addWidget(self.texto_plantilla)
        self.texto_plantilla.zoomChanged.connect(self._update_zoom_label)

        # 4) Botones inferiores (alineados a la derecha)
        bottom_buttons = QHBoxLayout()
        bottom_buttons.addStretch()
        bottom_buttons.addWidget(self.btn_generar_docx)
        bottom_buttons.addWidget(self.btn_copiar)
        bottom_buttons.addWidget(self.btn_ver_tramites)
        right_layout.addLayout(bottom_buttons)

        right_widget = QWidget()
        right_widget.setLayout(right_layout)
        main_layout.addWidget(right_widget, 3)

        # ------------------------------------------------------------------
        #  Resto de tu configuración del formulario (idéntica a la tuya)
        # ------------------------------------------------------------------
        general_page = self.general_group.content_area
        general_layout = QGridLayout(general_page)
        general_layout.setColumnStretch(0, 1)
        general_layout.setColumnStretch(1, 3)
        general_layout.setColumnStretch(2, 1)
        # Menor espacio vertical entre filas para compactar el formulario
        general_layout.setVerticalSpacing(1)

        row = 0
        lbl_loc = QLabel("Localidad:")
        general_layout.addWidget(lbl_loc, row, 0)
        general_layout.addWidget(self.var_localidad, row, 1)
        row += 1

        lbl_car = QLabel("Carátula:")
        general_layout.addWidget(lbl_car, row, 0)
        general_layout.addWidget(self.var_caratula, row, 1)
        row += 1

        lbl_trib = QLabel("Tribunal:")
        general_layout.addWidget(lbl_trib, row, 0)
        hbox_trib = QHBoxLayout()
        hbox_trib.addWidget(self.var_tribunal)
        general_layout.addLayout(hbox_trib, row, 1)
        row += 1

        lbl_sala = QLabel("Sala:")
        general_layout.addWidget(lbl_sala, row, 0)
        general_layout.addWidget(self.var_sala, row, 1)
        row += 1

        lbl_juez = QLabel("Juez:")
        general_layout.addWidget(lbl_juez, row, 0)
        hbox_juez = QHBoxLayout()
        hbox_juez.addWidget(self.var_juez)
        hbox_juez.addWidget(self.rb_juez_m)
        hbox_juez.addWidget(self.rb_juez_f)
        self.boton_cargo_juez = QPushButton("juez")
        self.boton_cargo_juez.setCheckable(True)
        self.boton_cargo_juez.clicked.connect(self.toggle_cargo_juez)
        hbox_juez.addWidget(self.boton_cargo_juez)
        general_layout.addLayout(hbox_juez, row, 1)
        row += 1

        lbl_fisc = QLabel("Fiscal:")
        general_layout.addWidget(lbl_fisc, row, 0)
        hbox_fisc = QHBoxLayout()
        hbox_fisc.addWidget(self.var_fiscal)
        hbox_fisc.addWidget(self.rb_fiscal_m)
        hbox_fisc.addWidget(self.rb_fiscal_f)
        general_layout.addLayout(hbox_fisc, row, 1)
        row += 1

        lbl_dia = QLabel("Día de audiencia:")
        general_layout.addWidget(lbl_dia, row, 0)
        general_layout.addWidget(self.var_dia_audiencia, row, 1)
        row += 1

        # Contenido del grupo "Datos generales"
        self.general_group.content_area.setLayout(general_layout)

        imputados_page = self.imputados_group.content_area
        imp_layout = QGridLayout(imputados_page)
        imp_layout.setColumnStretch(0, 1)
        imp_layout.setColumnStretch(1, 3)
        imp_layout.setVerticalSpacing(1)

        row = 0
        lbl_numimp = QLabel("Número de imputados:")
        imp_layout.addWidget(lbl_numimp, row, 0)
        imp_layout.addWidget(self.var_num_imputados, row, 1)
        row += 1

        lbl_imp = QLabel("Imputados:")
        imp_layout.addWidget(lbl_imp, row, 0, 1, 2)
        row += 1
        self.imputados_container = QWidget()
        self.imputados_layout = QVBoxLayout(self.imputados_container)
        imp_layout.addWidget(self.imputados_container, row, 0, 1, 2)
        row += 1

        self.imputados_group.content_area.setLayout(imp_layout)

        hechos_page = self.hechos_group.content_area
        hechos_layout = QGridLayout(hechos_page)
        hechos_layout.setColumnStretch(0, 1)
        hechos_layout.setColumnStretch(1, 3)
        hechos_layout.setVerticalSpacing(1)

        row = 0
        lbl_numhec = QLabel("Número de hechos:")
        hechos_layout.addWidget(lbl_numhec, row, 0)
        hechos_layout.addWidget(self.var_num_hechos, row, 1)
        row += 1

        lbl_hec = QLabel("Hechos:")
        hechos_layout.addWidget(lbl_hec, row, 0, 1, 2)
        row += 1
        self.hechos_container = QWidget()
        self.hechos_layout = QVBoxLayout(self.hechos_container)
        hechos_layout.addWidget(self.hechos_container, row, 0, 1, 2)
        row += 1

        self.hechos_group.content_area.setLayout(hechos_layout)

        extra_page = self.extra_group.content_area
        extra_layout = QGridLayout(extra_page)
        extra_layout.setColumnStretch(0, 1)
        extra_layout.setColumnStretch(1, 3)
        extra_layout.setColumnStretch(2, 1)
        extra_layout.setVerticalSpacing(1)

        row = 0
        lbl_sujev = QLabel("Sujeto eventual:")
        extra_layout.addWidget(lbl_sujev, row, 0)
        extra_layout.addWidget(self.var_sujeto_eventual, row, 1)
        row += 1

        lbl_manif = QLabel("Manifestaciones (del sujeto):")
        extra_layout.addWidget(lbl_manif, row, 0)
        extra_layout.addWidget(self.var_manifestacion, row, 1)
        row += 1

        lbl_vic = QLabel("Víctima:")
        extra_layout.addWidget(lbl_vic, row, 0)
        h_box_victima = QHBoxLayout()
        h_box_victima.addWidget(self.var_victima)
        h_box_victima.addWidget(self.var_victima_plural)
        extra_layout.addLayout(h_box_victima, row, 1)
        row += 1

        lbl_vicmani = QLabel("Manifestación (víctima):")
        extra_layout.addWidget(lbl_vicmani, row, 0)
        extra_layout.addWidget(self.var_victima_manifestacion, row, 1)
        row += 1

        lbl_prueba = QLabel("Pruebas:")
        extra_layout.addWidget(lbl_prueba, row, 0)
        extra_layout.addWidget(self.btn_prueba, row, 1)
        row += 1

        lbl_pruebrel = QLabel("Pruebas relevantes:")
        extra_layout.addWidget(lbl_pruebrel, row, 0)
        extra_layout.addWidget(self.btn_pruebas_importantes, row, 1)
        row += 1

        lbl_alegfis = QLabel("Alegato fiscal:")
        extra_layout.addWidget(lbl_alegfis, row, 0)
        extra_layout.addWidget(self.btn_alegato_fiscal, row, 1)
        row += 1

        lbl_alegdef = QLabel("Alegato de la defensa:")
        extra_layout.addWidget(lbl_alegdef, row, 0)
        extra_layout.addWidget(self.btn_alegato_defensa, row, 1)
        row += 1

        lbl_calif = QLabel("Calificación legal:")
        extra_layout.addWidget(lbl_calif, row, 0)
        extra_layout.addWidget(self.var_calificacion_legal, row, 1)
        row += 1

        lbl_corr = QLabel("Correcciones de calificación:")
        extra_layout.addWidget(lbl_corr, row, 0)
        extra_layout.addWidget(self.var_correccion_calif, row, 1)
        row += 1

        lbl_casovf = QLabel("¿Es un caso de VF o G?")
        extra_layout.addWidget(lbl_casovf, row, 0)
        extra_layout.addWidget(self.var_caso_vf, row, 1)
        row += 1

        lbl_uso_pot = QLabel("¿Se usó términos potenciales?")
        extra_layout.addWidget(lbl_uso_pot, row, 0)
        extra_layout.addWidget(self.var_uso_terminos_potenciales, row, 1)
        row += 1

        lbl_deco = QLabel("¿Decomiso?")
        extra_layout.addWidget(lbl_deco, row, 0)
        extra_layout.addWidget(self.var_decomiso_option, row, 1)
        btn_decomiso = QPushButton("Editar texto de Decomiso")
        extra_layout.addWidget(btn_decomiso, row, 2)
        btn_decomiso.clicked.connect(self.abrir_ventana_decomiso)
        row += 1

        lbl_restr = QLabel("¿Restricción de contacto?")
        extra_layout.addWidget(lbl_restr, row, 0)
        extra_layout.addWidget(self.var_restriccion_option, row, 1)
        btn_restriccion = QPushButton("Editar texto de Restricción")
        extra_layout.addWidget(btn_restriccion, row, 2)
        btn_restriccion.clicked.connect(self.abrir_ventana_restriccion)
        row += 1

        lbl_res = QLabel("Resuelvo:")
        extra_layout.addWidget(lbl_res, row, 0)
        self.var_resuelvo = QLineEdit()
        self.var_resuelvo.setVisible(False)
        self.btn_resuelvo = QPushButton("Editar resuelvo")
        extra_layout.addWidget(self.btn_resuelvo, row, 1)
        row += 1

        self.extra_group.content_area.setLayout(extra_layout)

        self.data.apply_to_sentencia(self)

    def generar_docx_con_html(self):
        """Genera un archivo DOCX respetando <p>, <b>, <i>... (parser básico)"""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        from PySide6.QtWidgets import QFileDialog
        import re
        import html

        # 1) Recogemos el HTML
        raw_html = self.texto_plantilla.toHtml()

        # -- A) Reemplazar <span style="font-weight:700;"> -> <b> y </span> -> </b>
        pattern_bold_open = re.compile(
            r'<span[^>]*style="[^"]*font-weight\s*:\s*(?:bold|700)[^"]*"[^>]*>',
            flags=re.IGNORECASE
        )
        raw_html = pattern_bold_open.sub("<b>", raw_html)
        #   1) Cursiva  →  <i>…</i>
        raw_html = re.sub(
            r'<span[^>]*font-style\s*:\s*italic[^>]*>(.*?)</span>',
            r'<i>\1</i>',
            raw_html,
            flags=re.IGNORECASE | re.DOTALL
        )

        #   2) Negrita  →  <b>…</b>
        raw_html = re.sub(r'</span>', '</b>', raw_html, flags=re.IGNORECASE)
        # -- B) Convertir <br> en espacio (o en salto de línea)
        raw_html = re.sub(r'(?i)<br\s*/?>', ' ', raw_html)
        # Opcional: si te queda un <br /> al final, también lo elimina:
        raw_html = raw_html.strip()

        # 2) Comenzamos un nuevo document
        document = Document()

        # 3) Encontramos los párrafos <p>...</p>
        paragraph_regex = re.compile(r'(?is)<p.*?>(.*?)</p>')
        paragraphs = paragraph_regex.findall(raw_html)
        if not paragraphs:
            paragraphs = [raw_html]

        for par_html in paragraphs:
            p = document.add_paragraph("")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # 4) Separamos tokens con <b>, </b>, <i>, </i>
            tokens = re.split(r'(<b>|</b>|<i>|</i>)', par_html)

            bold_active = False
            italic_active = False

            for token in tokens:
                # En vez de .strip(), usamos .rstrip() para solo eliminar

                t = html.unescape(token.replace('\n', ' '))  # Reemplaza saltos de línea por espacio
                if not t:
                    continue

                # Detectar etiquetas
                lt = t.lower()
                if lt == '<b>':
                    bold_active = True
                    continue
                elif lt == '</b>':
                    bold_active = False
                    continue
                elif lt == '<i>':
                    italic_active = True
                    continue
                elif lt == '</i>':
                    italic_active = False
                    continue

                # Si es texto normal, añadimos un espacio final para evitar
                # que se "junten" palabras:
                run = p.add_run(t + " ")
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.bold = bold_active
                run.italic = italic_active

        ruta, _ = QFileDialog.getSaveFileName(
            self, "Guardar DOCX", "", "Documentos de Word (*.docx)"
        )
        if ruta:
            document.save(ruta)
            print("Documento guardado en:", ruta)

    def setup_connections(self):
        # Conexiones
        self.var_localidad.textChanged.connect(self.actualizar_plantilla)
        self.var_caratula.textChanged.connect(self.actualizar_plantilla)
        self.var_tribunal.currentTextChanged.connect(self.actualizar_plantilla)
        self.var_sala.currentTextChanged.connect(self.actualizar_plantilla)
        self.var_juez.textChanged.connect(self.actualizar_plantilla)
        self.rb_juez_m.toggled.connect(self.actualizar_plantilla)
        self.rb_juez_f.toggled.connect(self.actualizar_plantilla)
        self.var_fiscal.textChanged.connect(self.actualizar_plantilla)
        self.rb_fiscal_m.toggled.connect(self.actualizar_plantilla)
        self.rb_fiscal_f.toggled.connect(self.actualizar_plantilla)
        self.var_dia_audiencia.textChanged.connect(self.actualizar_plantilla)
        self.var_num_imputados.valueChanged.connect(lambda: (self.update_imputados_section(), self.actualizar_plantilla()))
        self.var_caso_vf.currentTextChanged.connect(self.actualizar_plantilla)
        self.var_num_hechos.valueChanged.connect(self.update_hechos_section)
        self.var_num_hechos.valueChanged.connect(self.actualizar_plantilla)
        self.var_sujeto_eventual.textChanged.connect(self.actualizar_plantilla)
        self.var_manifestacion.textChanged.connect(self.actualizar_plantilla)
        self.var_victima.textChanged.connect(self.actualizar_plantilla)
        self.var_victima_plural.currentTextChanged.connect(self.actualizar_plantilla)
        self.var_victima_manifestacion.textChanged.connect(self.actualizar_plantilla)

        self.btn_alegato_fiscal.clicked.connect(self.abrir_ventana_alegato_fiscal)
        self.btn_alegato_defensa.clicked.connect(self.abrir_ventana_alegato_defensa)
        self.var_calificacion_legal.currentTextChanged.connect(self.actualizar_plantilla)
        self.var_calificacion_legal.currentTextChanged.connect(self.update_correccion_state)
        self.var_correccion_calif.textChanged.connect(self.actualizar_plantilla)
        self.var_uso_terminos_potenciales.currentTextChanged.connect(self.actualizar_plantilla)
        self.var_decomiso_option.currentTextChanged.connect(self.actualizar_plantilla)
        self.var_restriccion_option.currentTextChanged.connect(self.actualizar_plantilla)
        self.btn_resuelvo.clicked.connect(self.abrir_ventana_resuelvo)

        self.var_fiscal.textChanged.connect(
            lambda t: setattr(self.data, "fiscal_nombre", t.strip()))
        self.var_num_imputados.valueChanged.connect(
            lambda v: setattr(self.data, "n_imputados", v))


    def update_correccion_state(self, *_):
        if self.var_calificacion_legal.currentText() == "Incorrecta":
            self.var_correccion_calif.setEnabled(True)
        else:
            self.var_correccion_calif.setEnabled(False)
            self.var_correccion_calif.clear()

    def add_row(self, row, label_text, widget):
        lbl = QLabel(label_text)
        self.left_layout.addWidget(lbl, row, 0)
        self.left_layout.addWidget(widget, row, 1)

    from PySide6.QtWidgets import QMainWindow, QMessageBox

    # tramsent.py  (método del widget)
    def abrir_tramites(self):
        self.data.from_sentencia(self)     # 1) vuelca todo al modelo

        # 2) refrescamos inmediatamente la ventana principal
        main = self.parent()               # <- depende de cómo la instancies;
        while main and not hasattr(main, "rebuild_imputados"):
            main = main.parent()           # subimos hasta MainWindow
        if main:
            self.data.apply_to_main(main)  # ←★ aquí se copian los nombres BBB/CCC

        win = self.window()                # SentenciaWindow
        win.skip_confirm = True
        win.close()        

    def showEvent(self, e):
        super().showEvent(e)
        self.data.apply_to_sentencia(self)

    def update_imputados_section(self):
        n = self.var_num_imputados.value()
        while len(self.imputados) > n:
            w = self.imputados.pop()
            w["container"].deleteLater()
        while len(self.imputados) < n:
            idx = len(self.imputados) + 1
            container = QWidget()
            layout = QGridLayout(container)
            layout.setVerticalSpacing(1)
            layout.setVerticalSpacing(1)
            lbl_nombre = QLabel(f"Imputado/a #{idx} - Nombre:")
            le_nombre = QLineEdit()
            layout.addWidget(lbl_nombre, 0, 0)
            layout.addWidget(le_nombre, 0, 1, 1, 3)
            le_nombre.textChanged.connect(
                lambda txt, i=idx-1: self._sync_imp(i, "nombre", txt)
            )
            lbl_sexo = QLabel("Sexo:")
            rb_m = QRadioButton("M")
            rb_f = QRadioButton("F")
            rb_m.setChecked(True)
            grupo_sexo = QButtonGroup(container)
            grupo_sexo.addButton(rb_m)
            grupo_sexo.addButton(rb_f)
            layout.addWidget(lbl_sexo, 1, 0)
            layout.addWidget(rb_m, 1, 1)
            layout.addWidget(rb_f, 1, 2)
            lbl_datos = QLabel("Datos personales:")
            le_datos = QLineEdit()
            btn_datos = QPushButton("Editar datos personales")
            btn_datos.clicked.connect(partial(self.abrir_ventana_datos, idx-1))
            layout.addWidget(lbl_datos, 2, 0)
            layout.addWidget(btn_datos, 2, 1, 1, 3)
            le_datos.textChanged.connect(
                lambda txt, i=idx-1: self._sync_imp(i, "datos", txt)
            )
            lbl_defensor = QLabel("Defensor (nombre):")
            le_defensor = QLineEdit()
            layout.addWidget(lbl_defensor, 3, 0)
            layout.addWidget(le_defensor, 3, 1, 1, 3)
            le_defensor.textChanged.connect(
                lambda txt, i=idx-1: self._sync_imp(i, "defensa", txt)
            )
            lbl_tipo_def = QLabel("Tipo de Defensor:")
            cb_tipo_def = QComboBox()
            cb_tipo_def.addItems(["Público", "Privado"])
            layout.addWidget(lbl_tipo_def, 4, 0)
            layout.addWidget(cb_tipo_def, 4, 1)

            lbl_delitos = QLabel("Delitos (con sus artículos):")
            le_delitos = QLineEdit()
            layout.addWidget(lbl_delitos, 6, 0)
            layout.addWidget(le_delitos, 6, 1, 1, 3)
            le_delitos.textChanged.connect(
                lambda txt, i=idx-1: self._sync_imp(i, "delitos", txt)
            )
            lbl_condena = QLabel("Condena:")
            le_condena = QLineEdit()
            layout.addWidget(lbl_condena, 7, 0)
            layout.addWidget(le_condena, 7, 1, 1, 3)
            le_condena.textChanged.connect(
                lambda txt, i=idx-1: self._sync_imp(i, "condena", txt)
            )
            lbl_cond = QLabel("Datos personales agregados:")
            btn_cond = QPushButton("Editar datos agregados")
            btn_cond.clicked.connect(partial(self.abrir_ventana_condiciones, idx-1))
            le_cond = QLineEdit()
            layout.addWidget(lbl_cond, 8, 0)
            # botón en columna 1, colspan=2
            layout.addWidget(btn_cond, 8, 1, 1, 3)
            lbl_ant = QLabel("¿Antecedentes penales?")
            rb_ant_no = QRadioButton("No registra")
            rb_ant_si = QRadioButton("Registra")
            rb_ant_no.setChecked(True)
            grupo_ant = QButtonGroup(container)
            grupo_ant.addButton(rb_ant_no)
            grupo_ant.addButton(rb_ant_si)
            layout.addWidget(lbl_ant, 9, 0)
            layout.addWidget(rb_ant_no, 9, 1)
            layout.addWidget(rb_ant_si, 9, 2)
            lbl_ant_text = QLabel("Antecedentes:")
            le_ant = QLineEdit()
            le_ant.setEnabled(False)
            rb_ant_si.toggled.connect(lambda checked, w=le_ant: w.setEnabled(checked))
            layout.addWidget(lbl_ant_text, 10, 0)
            layout.addWidget(le_ant, 10, 1, 1, 3)
            lbl_confesion = QLabel("Confesión:")
            le_confesion = QLineEdit()
            layout.addWidget(lbl_confesion, 11, 0)
            layout.addWidget(le_confesion, 11, 1, 1, 3)
            lbl_ultima = QLabel("Última palabra:")
            le_ultima = QLineEdit()
            layout.addWidget(lbl_ultima, 12, 0)
            layout.addWidget(le_ultima, 12, 1, 1, 3)
            lbl_pautas = QLabel("Pautas de mensuración:")
            le_pautas = QLineEdit()
            btn_pautas = QPushButton("Añadir pautas de mensuración")

            # Hacemos que abra el diálogo rico sobre el QLineEdit de pautas:
            btn_pautas.clicked.connect(partial(self.abrir_ventana_pautas, idx-1))

            # Cuando el usuario edite el QLineEdit, refrescamos la plantilla:
            le_pautas.textChanged.connect(self.actualizar_plantilla)

            # Lo agregamos al layout, igual que 'datos' y 'condiciones':
            layout.addWidget(lbl_pautas, 13, 0)
            layout.addWidget(btn_pautas, 13, 1, 1, 3)

            if idx-1 < len(self.data.imputados):
                dprev = self.data.imputados[idx-1]
                le_nombre.setText( dprev.get("nombre",  "") )
                le_datos .setText( dprev.get("datos",   "") )
                le_cond.setText( dprev.get("condiciones", "") )
                le_defensor.setText(dprev.get("defensa", "") )
                le_condena   .setText( dprev.get("condena",    "") )
                le_delitos.setText( dprev.get("delitos", "") )
                le_pautas.setText( dprev.get("pautas", "") )
            
            container.setLayout(layout)
            self.imputados_layout.addWidget(container)
            for w in [le_nombre, le_datos, le_defensor, le_delitos, le_condena, le_cond, le_ant, le_confesion, le_ultima, le_pautas]:
                w.textChanged.connect(self.actualizar_plantilla)
            for w in [rb_m, rb_f, rb_ant_no, rb_ant_si]:
                w.toggled.connect(self.actualizar_plantilla)
            for w in [cb_tipo_def]:
                w.currentTextChanged.connect(self.actualizar_plantilla)
            
            self.imputados.append({
                "container": container,
                "nombre": le_nombre,
                "sexo_rb": (rb_m, rb_f),
                "datos": le_datos,
                "defensor": le_defensor,
                "tipo_def": cb_tipo_def,
                "delitos": le_delitos,
                "condena": le_condena,
                "condiciones": le_cond,
                "antecedentes_opcion": (rb_ant_no, rb_ant_si),
                "antecedentes": le_ant,
                "confesion": le_confesion,
                "ultima": le_ultima,
                "pautas": le_pautas,
            })
        if len(self.hechos) >= self.var_num_hechos.value():   # ya está lista la sección Hechos
            self.actualizar_plantilla()

    def update_hechos_section(self):
        n = self.var_num_hechos.value()
        while len(self.hechos) > n:
            w = self.hechos.pop()
            w["container"].deleteLater()
        # ——— Aquí creas los nuevos hechos ———
        while len(self.hechos) < n:
            idx = len(self.hechos)              # índice 0-based
            container = QWidget()
            layout = QGridLayout(container)

            lbl_desc = QLabel(f"Descripción del suceso #{idx+1}:")
            layout.addWidget(lbl_desc, 0, 0)
            le_desc = QLineEdit()
            # quitamos el le_desc completamente

            btn_desc = QPushButton("Redactar el hecho")
            btn_desc.clicked.connect(partial(self.abrir_ventana_descripcion, idx))
            # Para que expanda en horizontal hasta llenar el espacio
            btn_desc.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            # Ahora abarca desde la columna 1 y ocupa 3 columnas (igual que tus QLineEdit con colspan=3)
            layout.addWidget(btn_desc, 0, 1, 1, 1)

            lbl_aclar = QLabel(f"Aclaraciones hecho #{idx}:")
            le_aclar = QLineEdit()
            layout.addWidget(lbl_aclar, 1, 0)
            layout.addWidget(le_aclar, 1, 1, 1, 1)
            next_row = 2
            lbl_ofi = QLabel("Oficina que elevó:")
            le_ofi = QLineEdit()
            rb_j = QRadioButton("Juzgado")
            rb_f = QRadioButton("Fiscalía")
            rb_j.setChecked(True)
            grupo_ofi = QButtonGroup(container)
            grupo_ofi.addButton(rb_j)
            grupo_ofi.addButton(rb_f)
            layout.addWidget(lbl_ofi, next_row, 0)
            layout.addWidget(le_ofi, next_row, 1)
            layout.addWidget(rb_j, next_row, 2)
            layout.addWidget(rb_f, next_row, 3)
            next_row += 1
            lbl_auto = QLabel("N° del auto:")
            le_auto = QLineEdit()
            layout.addWidget(lbl_auto, next_row, 0)
            layout.addWidget(le_auto, next_row, 1, 1, 1)

            next_row += 1
            lbl_fec = QLabel("Fecha de elevación:")
            le_fec = QLineEdit()
            layout.addWidget(lbl_fec, next_row, 0)
            layout.addWidget(le_fec, next_row, 1, 1, 1)

            container.setLayout(layout)
            self.hechos_layout.addWidget(container)
            for w in [le_desc, le_aclar, le_ofi, le_auto, le_fec]:
                w.textChanged.connect(self.actualizar_plantilla)
            for w in [rb_j, rb_f]:
                w.toggled.connect(self.actualizar_plantilla)
            
            self.hechos.append({
                "container": container,
                "descripcion": le_desc,
                "aclaraciones": le_aclar,
                "oficina": le_ofi,
                "rb_j": rb_j,
                "rb_f": rb_f,
                "num_auto": le_auto,
                "fecha_elev": le_fec,
            })
        self.actualizar_plantilla()

    def get_sexos_imputados(self):
        sexos = []
        for imp in self.imputados:
            if imp["sexo_rb"][0].isChecked():
                sexos.append("M")
            else:
                sexos.append("F")
        return sexos

    def actualizar_plantilla(self):
        sb = self.texto_plantilla.verticalScrollBar()
        pos = sb.value()
        
        if not self.imputados:
            return
        
        n_imp = self.var_num_imputados.value()
        n_hec = min(self.var_num_hechos.value(), len(self.hechos))

        # 1) Localidad
        localidad = self.var_localidad.text().strip()
        if not localidad:
            localidad = "Córdoba"  # fallback

        # 2) Fecha en letras
        fecha_letras = self.var_dia_audiencia.text().strip()

        # 3) Causa/caratula
        caratula = self.var_caratula.text().strip()

        # 4) Tribunal
        tribunal = self.var_tribunal.currentText()

        # 5) Sala
        sala = self.var_sala.currentText().strip()

        # 6) Juez
        juez_nombre = self.var_juez.text().strip()
        if self.rb_juez_m.isChecked():
            juez_intro = "del juez"
        else:
            juez_intro = "de la jueza"

        juez_cargo = self.boton_cargo_juez.text().lower()  # "juez" o "vocal"

        # Si eligió "vocal", entonces reemplazamos "juez"/"jueza" por "vocal"
        if juez_cargo == "vocal":
            juez_intro = "del vocal" if self.rb_juez_m.isChecked() else "de la vocal"

        texto_juez = f"{juez_intro} {juez_nombre}"

        # 7) Fiscal
        fiscal_nombre = self.var_fiscal.text().strip()
        fiscal_articulo = "el" if self.rb_fiscal_m.isChecked() else "la"

        # 8) Imputados => para “el/la/las/los imputado/a/as/os”, 
        n_imp = self.var_num_imputados.value()
        sexos = self.get_sexos_imputados()
        cant_masc = sum(1 for s in sexos if s == "M")
        cant_fem = n_imp - cant_masc

        #   “el/la/los/las” ...
        if n_imp == 1:
            # Solo uno
            if sexos[0] == "M":
                imput_label = "el imputado"
                asistido_label = "asistido"
                acusado_label = "acusado"
            else:
                imput_label = "la imputada"
                asistido_label = "asistida"
                acusado_label = "acusada"
        else:
            # Varios
            if cant_fem == n_imp:
                # Todas mujeres
                imput_label = "las imputadas"
                asistido_label = "asistidas"
                accused_label = "acusadas"
            elif cant_masc == n_imp:
                imput_label = "los imputados"
                asistido_label = "asistidos"
                accused_label = "acusados"
            else:
                # Mixto
                imput_label = "los imputados"
                asistido_label = "asistidos"
                accused_label = "acusados"


        if n_imp == 1:
            acusado_label = acusado_label  # ya definimos en el if
        else:
            # Varios
            if cant_fem == n_imp:
                acusado_label = "acusadas"
            elif cant_masc == n_imp:
                acusado_label = "acusados"
            else:
                acusado_label = "acusados"

        # 9) Nombre y Apellido => con conjunción
        names_list = []
        for i, imp in enumerate(self.imputados):
            nm = imp["nombre"].text().strip()
            if not nm:
                nm = f"Imputado#{i+1}"
            names_list.append(nm)
        nombres_conj = format_list_for_sentence(names_list)

        #    Recopilar
        defenders_list = [imp["defensor"].text().strip() for imp in self.imputados]
        def_dict = {}
        for i, d in enumerate(defenders_list):
            if d:
                def_dict.setdefault(d, []).append(names_list[i])
        # Al final, unimos con conjunción

        defensores_unicos = list(def_dict.keys())  # no repetidos
        defensa_final = format_list_for_sentence(defensores_unicos)

        # 12) “{fue/ron} {acusado/a/as/os}”:
        #   => “fue” si 1, “fueron” si >1
        fue_ron = "fue" if n_imp == 1 else "fueron"

        datos_personales_list = []
        for i, imp in enumerate(self.imputados):
            nm = imp["nombre"].text().strip()
            if not nm:
                nm = f"Imputado#{i+1}"
            d = imp["datos"].text().strip()
            # Armamos => "<b>Nombre</b>, datospersonales"
            comb = f"<b>{nm}</b>"
            if d:
                comb += f", {d}"
            datos_personales_list.append(comb)
        datos_personales_str = format_list_with_semicolons(datos_personales_list)
        art_tribunal = "el" if self.boton_cargo_juez.text().lower() == "juez" else "la"
        primer_parrafo = (
            f"En la ciudad de {localidad}, el {fecha_letras}, se dan a conocer "
            f"los fundamentos de la sentencia dictada en la causa <b>{caratula}</b>, "
            f"juzgada por {art_tribunal} {tribunal}, en la sala {sala} a cargo {texto_juez}."
        )

        segundo_parrafo = (
            f"En el debate intervinieron {fiscal_articulo} {fiscal_nombre}, "
            f"y {imput_label} {nombres_conj}, {asistido_label} por {defensa_final}."
        )

        tercer_parrafo = (
            f"En esta causa {fue_ron} {acusado_label} {datos_personales_str}."
        )

        # Construct final
        nuevo_inicio = (
            f"<p align='justify'>{primer_parrafo}</p>"
            f"<p align='justify'>{segundo_parrafo}</p>"
            f"<p align='justify'>{tercer_parrafo}</p>"
        )

        plantilla = nuevo_inicio

        acusaciones_parciales = []
        for h in self.hechos:
            oficina_rb_val = "Juzgado" if h["rb_j"].isChecked() else "Fiscalía"
            oficina_txt = h["oficina"].text().strip()
            num_auto = h["num_auto"].text().strip()
            fecha_elev = h["fecha_elev"].text().strip()
            aclaraciones = h["aclaraciones"].text().strip()


            if oficina_rb_val == "Juzgado":
                base = "El auto de elevación a juicio"
                if num_auto and fecha_elev:
                    texto = f"{base} n° {num_auto} de fecha {fecha_elev}"
                elif num_auto:
                    texto = f"{base} n° {num_auto}"
                elif fecha_elev:
                    texto = f"{base} de fecha {fecha_elev}"
                else:
                    texto = base
            else:
                base = "El requerimiento de citación a juicio"
                texto = f"{base} de fecha {fecha_elev}" if fecha_elev else base

            if oficina_txt:
                texto += f", dictado por {oficina_txt},"
            texto_aclar = f" ({aclaraciones})" if aclaraciones else ""

            acusaciones_parciales.append(f"{texto}")

        # Dedupl
        unique_acusaciones = []
        seen = set()
        for item in acusaciones_parciales:
            norm = " ".join(item.strip().lower().split())
            if norm not in seen:
                seen.add(norm)
                unique_acusaciones.append(item)

        base_texts = {"el auto de elevación a juicio", "el requerimiento de citación a juicio"}
        non_base_exists = any(" ".join(u.strip().lower().split()) not in base_texts for u in unique_acusaciones)
        if non_base_exists and len(unique_acusaciones) > 1:
            unique_acusaciones = [u for u in unique_acusaciones if " ".join(u.strip().lower().split()) not in base_texts]

        for i, uacc in enumerate(unique_acusaciones):
            uacc = uacc.strip()
            # Elimina comas, puntos, o punto y coma finales
            while len(uacc) > 0 and uacc[-1] in [',', ';', '.']:
                uacc = uacc[:-1]

            # Si no es el primer ítem, forzalo a comenzar en minúscula
            # (por ej. para que "El requerimiento..." se transforme en "el requerimiento...")
            if i > 0 and uacc:
                uacc = uacc[0].lower() + uacc[1:]

            unique_acusaciones[i] = uacc


        acus_unificado = format_list_with_semicolons(unique_acusaciones)

        if len(unique_acusaciones) > 1:
            verbo_atribuir = "atribuyeron"
        else:
            verbo_atribuir = "atribuyó"

        n_hec = self.var_num_hechos.value()

        sexos = self.get_sexos_imputados()
        n_imp = self.var_num_imputados.value()
        cant_masc = sum(1 for s in sexos if s == "M")
        cant_fem = n_imp - cant_masc

        acus_final = f" en {acus_unificado}" if acus_unificado else ""
        sexos = self.get_sexos_imputados()
        cant_masc = sum(1 for s in sexos if s == "M")
        cant_fem = n_imp - cant_masc
        if n_imp == 1:
            if sexos[0] == "M":
                al_imput_label = "al imputado"
            else:
                al_imput_label = "a la imputada"
        else:
            if cant_fem == n_imp:
                al_imput_label = "a las imputadas"
            else:
                al_imput_label = "a los imputados"

        # Y definimos "hechos_label" (o puedes reusar tu "frase_hechos"):
        if self.var_num_hechos.value() == 1:
            hechos_label = "el siguiente hecho"
        else:
            hechos_label = "los siguientes hechos"
        
        plantilla += (
    f"<p align='justify'>"
    f"{acus_unificado} {verbo_atribuir} {al_imput_label} {hechos_label}:"
    f"</p>"
)

        # Listado de hechos
        for i in range(n_hec):
            desc_str = self.hechos[i]["descripcion"].text().strip()
            aclar_str = self.hechos[i]["aclaraciones"].text().strip()
            if n_hec == 1:
                if aclar_str:
                    plantilla += f"<p align='justify'><i>{desc_str}</i> ({aclar_str})</p>"
                else:
                    plantilla += f"<p align='justify'><i>{desc_str}</i></p>"
            else:
                ordinal = ORDINALES_HECHOS[i] if i < len(ORDINALES_HECHOS) else f"{i+1}°"
                if aclar_str:
                    plantilla += f"<p align='justify'><b>{ordinal} hecho ({aclar_str})</b>: <i>{desc_str}</i></p>"
                else:
                    plantilla += f"<p align='justify'><b>{ordinal} hecho:</b> <i>{desc_str}</i></p>"

        # Determinamos si decimos "la existencia del hecho" o "la existencia de los hechos"
        if n_hec == 1:
            exist_label = "la existencia del hecho"
        else:
            exist_label = "la existencia de los hechos"

        if n_imp == 1:
            # Solo un imputado
            if sexos[0] == "F":
                resp_label = "de la acusada"
            else:
                resp_label = "del acusado"
        else:
            # Varios imputados
            if cant_fem == n_imp:
                resp_label = "de las acusadas"
            else:
                resp_label = "de los acusados"

        primera_cuestion = f"¿Están probadas {exist_label} y la participación responsable {resp_label}?"
        plantilla += (
            f"<p align='justify'>El tribunal se planteó las siguientes cuestiones a resolver:</p>"
            f"<p align='justify'>&nbsp;&nbsp;&nbsp;&nbsp;<b>PRIMERA CUESTIÓN:</b> {primera_cuestion}</p>"
            f"<p align='justify'>&nbsp;&nbsp;&nbsp;&nbsp;<b>SEGUNDA CUESTIÓN:</b> en su caso, ¿qué calificación legal es aplicable?</p>"
            f"<p align='justify'>&nbsp;&nbsp;&nbsp;&nbsp;<b>TERCERA CUESTIÓN:</b> ¿qué pronunciamiento corresponde dictar?</p>"
            f"<p align='justify'><b>A LA PRIMERA CUESTIÓN PLANTEADA, {self.cargo_juez_en_mayusculas()} {juez_nombre.upper()} DIJO:</b></p>"
        )

        acus_unificado_minus = acus_unificado

        import re
        acus_unificado_minus = re.sub(r'^(El|La|Los|Las)\b', lambda m: m.group(1).lower(), acus_unificado_minus.strip())

        if n_hec == 1:
            hecho_label2 = "del hecho contenido"
        else:
            hecho_label2 = "de los hechos contenidos"

        # Ahora armás el nuevo párrafo, usando la versión con minúscula:
        plantilla += (
            f"<p align='justify'><b>1. Acusación:</b> la exigencia impuesta en el artículo 408, inc. 1º del CPP "
            f"se encuentra satisfecha con la enunciación al comienzo de la sentencia {hecho_label2} "
            f"en {acus_unificado_minus}, a donde me remito para ser breve.</p>"
        )
        accusations = []
        for i, imp in enumerate(self.imputados):
            nm = imp["nombre"].text().strip()
            if not nm:
                nm = f"Imputado#{i+1}"
            delit = imp["delitos"].text().strip()
            accusations.append(f"{nm} bajo la calificación legal de {delit}")

        if n_imp == 1 and n_hec == 1:
            acusacion_prefix = "Por tal conducta se acusa"
        else:
            acusacion_prefix = "Por tales conductas se acusa"

        delitos_dict = {}
        for i, imp in enumerate(self.imputados):
            nm = imp["nombre"].text().strip()
            if not nm:
                nm = f"Imputado#{i+1}"
            delit = imp["delitos"].text().strip()
            if delit not in delitos_dict:
                delitos_dict[delit] = []
            delitos_dict[delit].append(nm)

        accusations_grouped = []
        for delito, lista_nombres in delitos_dict.items():
            if len(lista_nombres) == 1:
                unica_persona = lista_nombres[0]
                fragmento = f"{unica_persona} bajo la calificación legal de {delito}"
            else:
                nombres_unidos = format_list_for_sentence(lista_nombres)
                fragmento = f"{nombres_unidos} bajo la calificación legal de {delito}"
            accusations_grouped.append(fragmento)

        if n_imp == 1 and n_hec == 1:
            acusacion_prefix = "Por tal conducta se acusa"
        else:
            acusacion_prefix = "Por tales conductas se acusa"

        if not accusations_grouped:
            pass
        elif len(accusations_grouped) == 1:
            plantilla += f"<p align='justify'>{acusacion_prefix} a {accusations_grouped[0]}.</p>"
        else:
            accusations_with_a = [f"a {x}" for x in accusations_grouped]
            last = accusations_with_a.pop()
            joined = "; ".join(accusations_with_a)
            plantilla += f"<p align='justify'>{acusacion_prefix} {joined}; y {last}.</p>"

        # Siguientes secciones “II. Trámite de juicio abreviado...”, etc.
        # (Las copio sin tocar)

        if n_imp == 1:
            defense_text = "la defensa"
            agreement_text = "del acuerdo alcanzado"
        else:
            unique_defenders = {imp["defensor"].text().strip() for imp in self.imputados if imp["defensor"].text().strip()}
            defense_text = "las defensas" if len(unique_defenders) > 1 else "la defensa"
            agreement_text = "de los acuerdos alcanzados"

        plantilla += (
            f"<p align='justify'><b>2. Trámite de juicio abreviado (art. 415 CPP):</b></p>"
            f"<p align='justify'><b>a) Acuerdo:</b> {defense_text} y la fiscalía hicieron conocer los términos {agreement_text} para la realización de un juicio abreviado que, en cuanto a la pena, "
        )

        if n_imp == 1:
            condena_unica = self.imputados[0]["condena"].text().strip()
            plantilla += f"determinó la de {condena_unica}.</p>"
        else:
            frag_penas = []
            for i in range(n_imp):
                nombre_tmp = self.imputados[i]["nombre"].text().strip()
                if not nombre_tmp:
                    nombre_tmp = f"Imputado#{i+1}"
                frag_penas.append(f"para {nombre_tmp}, la de {self.imputados[i]['condena'].text().strip()}")
            acuerdo_str = format_list_with_semicolons(frag_penas)
            plantilla += f"determinó {acuerdo_str}.</p>"

        sujeto_str = self.var_sujeto_eventual.text().strip()
        mani_str = self.var_manifestacion.text().strip()
        if sujeto_str or mani_str:
            plantilla += (
                f"<p align='justify'>Se le concedió la palabra a {sujeto_str} "
                f"para que exprese su opinión acerca del acuerdo informado, y manifestó: {mani_str}.</p>"
            )

        if n_imp == 1:
            if sexos[0] == "M":
                acus_label = "al acusado"
                verb_comp = "comprendía"
                verb_con = "conocía"
            else:
                acus_label = "a la acusada"
                verb_comp = "comprendía"
                verb_con = "conocía"
        else:
            if cant_fem == n_imp:
                acus_label = "a las acusadas"
            else:
                acus_label = "a los acusados"
            verb_comp = "comprendían"
            verb_con = "conocían"

        plantilla += (
            f"<p align='justify'>Las características de esta modalidad de juzgamiento y del acuerdo mencionado fueron explicados por el tribunal {acus_label}, y se verificó así que {verb_comp} su contenido y sus consecuencias, que {verb_con} su derecho a exigir un juicio oral, y que su conformidad era libre y voluntaria.</p>"
        )

        victim = self.var_victima.text().strip() if self.var_victima.text().strip() else "la víctima"
        manifest_victim = self.var_victima_manifestacion.text().strip()
        victim_plural_mode = (self.var_victima_plural.currentText().strip().lower() == "más")
        if manifest_victim:
            if victim_plural_mode:
                plantilla += f"<p align='justify'>Además, el fiscal hizo saber que {victim} fueron previamente informadas acerca de dichos aspectos y que manifestaron {manifest_victim}.</p>"
            else:
                plantilla += f"<p align='justify'>Además, el fiscal hizo saber que {victim} fue previamente informada acerca de dichos aspectos y que manifestó {manifest_victim}.</p>"

        # (b) Declaración del imputado
        if n_imp == 1:
            if sexos[0] == "M":
                header = "<p align='justify'><b>b) Declaración del imputado:</b></p>"
            else:
                header = "<p align='justify'><b>b) Declaración de la imputada:</b></p>"
        else:
            if all(s == "F" for s in sexos):
                header = "<p align='justify'><b>b) Declaración de las imputadas:</b></p>"
            else:
                header = "<p align='justify'><b>b) Declaración de los imputados:</b></p>"

        plantilla += header

        if n_imp == 1:
            if sexos[0] == "M":
                interrogado = "al ser interrogado"
            else:
                interrogado = "al ser interrogada"
        else:
            if all(s == "F" for s in sexos):
                interrogado = "al ser interrogadas"
            elif all(s == "M" for s in sexos):
                interrogado = "al ser interrogados"
            else:
                interrogado = "al ser interrogados"

        plantilla += (
            f"<p align='justify'><b>Condiciones personales:</b> {interrogado} por el tribunal y las partes, además de los datos consignados al comienzo de esta resolución, "
        )

        prefixes = ["A su vez, ", "Por su parte, ", "A su turno, ", "También, "]
        verbs = ["agregó", "dijo", "mencionó", "añadió"]

        final_names_list = []
        for i, imp in enumerate(self.imputados):
            nm = imp["nombre"].text().strip()
            if not nm:
                nm = f"Imputado#{i+1}"
            final_names_list.append(nm)

        for i, imp in enumerate(self.imputados):
            name_i = f"<b>{final_names_list[i]}</b>"
            condiciones = imp["condiciones"].text().strip()
            verb = verbs[i % len(verbs)]

            if i == 0:
                plantilla += f"{name_i} {verb} que {condiciones}."
            else:
                prefix = prefixes[(i-1) % len(prefixes)]
                plantilla += f" {prefix}{name_i} {verb} que {condiciones}."

        plantilla += "</p>"

        no_registra_list = []
        si_registra_list = []
        for i, imp in enumerate(self.imputados):
            nm = final_names_list[i]
            name_i = f"<b>{nm}</b>"
            no_registra = imp["antecedentes_opcion"][0].isChecked()
            ant_text = imp["antecedentes"].text().strip()
            if no_registra:
                no_registra_list.append(name_i)
            else:
                if ant_text:
                    si_registra_list.append((name_i, ant_text))
                else:
                    si_registra_list.append((name_i, "registra antecedentes penales (sin detalle)."))

        mentions = []
        if no_registra_list:
            if len(no_registra_list) == 1:
                mention_sin = f"{no_registra_list[0]} no registra condenas computables."
            else:
                conj = ", ".join(no_registra_list[:-1]) + " y " + no_registra_list[-1]
                mention_sin = f"{conj} no registran condenas computables."
            mentions.append(mention_sin)

        for (nombre_en_negrita, texto_antecedentes) in si_registra_list:
            mentions.append(f"{nombre_en_negrita} registra los siguientes antecedentes: {texto_antecedentes}")

        if not mentions:
            plantilla += (
                "<p align='justify'>En cuanto a sus antecedentes penales, por Secretaría no se cuenta con "
                "información alguna o no hubo datos cargados.</p>"
            )
        else:
            texto_antecedentes = "<p align='justify'>En cuanto a sus antecedentes penales, por Secretaría se informó que "
            total_m = len(mentions)
            prefixes_cycle = ["A su vez,", "Separadamente,", "Asimismo,"]

            for i, mention in enumerate(mentions):
                if i == 0:
                    texto_antecedentes += mention
                else:
                    es_ultima = (i == total_m - 1)
                    if i == 1 and no_registra_list and si_registra_list:
                        prefix = "Por su parte,"
                    else:
                        if es_ultima:
                            prefix = "Finalmente,"
                        else:
                            prefix = prefixes_cycle[(i - 1) % len(prefixes_cycle)]
                    texto_antecedentes += f" {prefix} {mention}"

            texto_antecedentes += "</p>"
            plantilla += texto_antecedentes

        plantilla += "<p align='justify'><b>Confesión:</b> "

        if n_imp == 1:
            nm = final_names_list[0]
            if sexos[0] == "M":
                info_text = "fue informado"
                atrib_text = "se le atribuye" if n_hec == 1 else "se le atribuyen"
                facto_text = "del hecho" if n_hec == 1 else "de los hechos"
                plantilla += (
                    f"A fin de ratificar la voluntad manifestada en el acuerdo previo para la realización del juicio abreviado, "
                    f"el imputado {info_text} detalladamente {facto_text} que {atrib_text}, "
                    f"de las pruebas existentes en su contra y de la facultad que le acuerda la ley de abstenerse de prestar declaración "
                    f"sin que su silencio implique una presunción de culpabilidad (arts. 385 y 259 CPP) sino la sola consecuencia "
                    f"de impedir el trámite del art. 415 CPP."
                )
            else:
                info_text = "fue informada"
                atrib_text = "se le atribuye" if n_hec == 1 else "se le atribuyen"
                facto_text = "del hecho" if n_hec == 1 else "de los hechos"
                plantilla += (
                    f"A fin de ratificar la voluntad manifestada en el acuerdo previo para la realización del juicio abreviado, "
                    f"la imputada {info_text} detalladamente {facto_text} que {atrib_text}, "
                    f"de las pruebas existentes en su contra y de la facultad que le acuerda la ley de abstenerse de prestar declaración "
                    f"sin que su silencio implique una presunción de culpabilidad (arts. 385 y 259 CPP) sino la sola consecuencia "
                    f"de impedir el trámite del art. 415 CPP."
                )
            
            conf_text = self.imputados[0]["confesion"].text().strip()
            if conf_text:
                plantilla += f" Ante ello, {nm} dijo: “{conf_text}”.</p>"
            else:
                plantilla += f" Ante ello, {nm} .</p>"
        else:
            # Caso de varios imputados: se imprime una parte colectiva y luego las confesiones individuales.
            if all(s == "F" for s in sexos):
                collective = "las imputadas fueron informadas"
                atrib_text = "se les atribuye" if n_hec == 1 else "se les atribuyen"
            else:
                collective = "los imputados fueron informados"
                atrib_text = "se les atribuye" if n_hec == 1 else "se les atribuyen"
            facto_text = "del hecho" if n_hec == 1 else "de los hechos"
            plantilla += (
                f"A fin de ratificar la voluntad manifestada en el acuerdo previo para la realización del juicio abreviado, "
                f"{collective} detalladamente {facto_text} que {atrib_text}, "
                f"de las pruebas existentes en su contra y de la facultad que la ley les acuerda de abstenerse de prestar declaración "
                f"sin que su silencio implique una presunción de culpabilidad (arts. 385 y 259 CPP) sino la sola consecuencia "
                f"de impedir el trámite del art. 415 CPP.</p>"
            )
            # Ciclo para agregar las confesiones individuales usando prefijos y verbos cíclicos
            prefixes_cycle = ["Ante ello,", "A su turno,", "Luego,", "Después,"]
            verbs_cycle = ["expresó", "manifestó", "refirió", "declaró", "afirmó"]
            for i, imp in enumerate(self.imputados):
                nm = final_names_list[i]
                conf_text = imp["confesion"].text().strip()
                prefix = prefixes_cycle[i % len(prefixes_cycle)]
                verb = verbs_cycle[i % len(verbs_cycle)]
                if conf_text:
                    plantilla += f"<p align='justify'>{prefix} {nm} {verb}: “{conf_text}”.</p>"
                else:
                    plantilla += f"<p align='justify'>{prefix} {nm} .</p>"

        # c) Aceptación
        if n_imp == 1:
            if sexos[0] == "M":
                suj_label = "el imputado"
                ha_sido = "ha sido"
                informado = "informado"
                han_expresado = "ha expresado"
                han_reconocido = "ha reconocido"
            else:
                suj_label = "la imputada"
                ha_sido = "ha sido"
                informado = "informada"
                han_expresado = "ha expresado"
                han_reconocido = "ha reconocido"
        else:
            if cant_fem == n_imp:
                suj_label = "las acusadas"
            else:
                suj_label = "los acusados"
            ha_sido = "han sido"
            informado = "informados"
            han_expresado = "han expresado"
            han_reconocido = "han reconocido"

        plantilla += (
            f"<p align='justify'><b>c) Aceptación del Tribunal:</b> de la reseña que precede surge que se han cumplimentado los requisitos de ley, "
            f"pues se ha corroborado que {suj_label} {ha_sido} acabadamente {informado} de los términos del acuerdo y que {han_expresado} su conformidad "
            f"de manera libre y voluntaria. Asimismo, {han_reconocido} lisa y llanamente su responsabilidad en los mismos términos en que se les ha sido "
            f"atribuida por la acusación.</p>"
        )

        calif_es_correcta = (self.var_calificacion_legal.currentText() == "Correcta")

        n_hec = self.var_num_hechos.value()
        # (Tramo idéntico al original)
        if calif_es_correcta:
            if n_imp == 1 and n_hec == 1:
                calif_text = (
                    "La calificación legal asignada por la fiscalía es correcta "
                    "para el hecho que se le achaca y la pena pactada se encuentra dentro "
                    "de la escala penal prevista para el delito endilgado (art. 415 CPP)."
                )
            elif n_imp > 1 and n_hec == 1:
                calif_text = (
                    "La calificación legal asignada por la fiscalía es correcta "
                    "para el hecho que se les achaca y las penas pactadas se encuentran dentro "
                    "de la escala penal prevista para los delitos endilgados (art. 415 CPP)."
                )
            elif n_imp == 1 and n_hec > 1:
                calif_text = (
                    "La calificación legal asignada por la fiscalía es correcta "
                    "para los hechos que se le achacan y la pena pactada se encuentra dentro "
                    "de la escala penal prevista para los delitos endilgados (art. 415 CPP)."
                )
            else:
                calif_text = (
                    "La calificación legal asignada por la fiscalía es correcta "
                    "para los hechos que se les achacan y las penas pactadas se encuentran dentro "
                    "de la escala penal prevista para los delitos endilgados (art. 415 CPP)."
                )
        else:
            if n_imp == 1 and n_hec == 1:
                calif_text = (
                    "La calificación legal amerita cierta corrección que se expondrá "
                    "luego en la segunda cuestión, pero que no afecta el monto punitivo "
                    "acordado porque este se encuentra dentro de la escala penal "
                    "prevista para el delito aplicable (art. 415 CPP)."
                )
            elif n_imp > 1 and n_hec == 1:
                calif_text = (
                    "La calificación legal amerita cierta corrección que se expondrá "
                    "luego en la segunda cuestión, pero que no afecta los montos punitivos "
                    "acordados porque estos se encuentran dentro de la escala penal "
                    "prevista para los delitos aplicables (art. 415 CPP)."
                )
            elif n_imp == 1 and n_hec > 1:
                calif_text = (
                    "La calificación legal amerita cierta corrección que se expondrá "
                    "luego en la segunda cuestión, pero que no afecta el monto punitivo "
                    "acordado porque este se encuentra dentro de la escala penal "
                    "prevista para los delitos aplicables (art. 415 CPP)."
                )
            else:
                calif_text = (
                    "La calificación legal amerita cierta corrección que se expondrá "
                    "luego en la segunda cuestión, pero que no afecta los montos punitivos "
                    "acordados porque estos se encuentran dentro de la escala penal "
                    "prevista para los delitos aplicables (art. 415 CPP)."
                )

        plantilla += f"<p align='justify'>{calif_text}</p>"
                # Placeholder para “{la/s solicitud/es formulada/s}”
        if n_imp == 1:
            solicitudes_str = "la solicitud formulada"
        else:
            solicitudes_str = "las solicitudes formuladas"

        defenders_list = [imp["defensor"].text().strip() for imp in self.imputados if imp["defensor"].text().strip()]
        def_dict = {}
        for i, d in enumerate(defenders_list):
            # Agrupar por defensor (sin repeticiones)
            def_dict.setdefault(d, []).append(i)
        # Obtener la lista de defensores únicos
        defensores_unicos = list(def_dict.keys())
        # También ya calculas 'defensa_final' para otra parte:
        defensa_final = format_list_for_sentence(defensores_unicos)

        # Ahora, para el placeholder {su/s defensa/s}:
        if not defensores_unicos:
            # Si ninguno ingresó defensor, usamos la forma singular por defecto
            defensa_str = "la defensa"
        elif len(defensores_unicos) == 1:
            defensa_str = "su defensa"
        else:
            defensa_str = "sus defensas"


        # Ahora construyes la cadena final reemplazando las partes entre llaves:
        plantilla += (
            f'Tales constataciones son las únicas habilitadas por la ley al Tribunal en el marco del juicio abreviado '
            f'(TSJ, Sala Penal, S. n° 124, 19/04/2017, "Cabrera", entre otros; Jaime, Marcelo Nicolás, "El juicio abreviado", '
            f'en AAVV, Comentarios a la reforma del Código Procesal Penal, dir. Maximiliano Hairabedián, Advocatus, 2017, págs. 161/162; '
            f'Cafferata Nores –Tarditti, cit., T. 2, pág. 314), y por ello corresponde hacer lugar a {solicitudes_str} por el Ministerio Público Fiscal, '
            f'{imput_label} y {defensa_str}.'
        )

        plantilla += (
            f"<p align='justify'><b>3. Enumeración de la prueba:</b> según lo dispuesto por el artículo 415 CPP y a pedido de las partes, se incorporó la prueba recolectada durante la investigación penal preparatoria y la investigación preliminar: {self.var_prueba}</p>"
            f"<p align='justify'><b>4. Discusión final:</b> finalmente, las partes emitieron sus conclusiones de acuerdo con sus respectivos intereses. Así, la Fiscalía manifestó {self.var_alegato_fiscal}. Por su parte, la defensa expuso {self.var_alegato_defensa}.</p>"
        )

# ======================================
# BLOQUE PARA MOSTRAR "ÚLTIMA PALABRA"
# ======================================
        speakers_2 = []
        non_speakers_2 = []

        for i, imp in enumerate(self.imputados):
            ultima_str = imp["ultima"].text().strip()   # <-- la "Última palabra" ingresada
            if ultima_str:
                speakers_2.append((i, ultima_str))      # este imputado sí habló
            else:
                non_speakers_2.append(i)                # este imputado NO habló

        def nombre(idx):
            return final_names_list[idx]  # asumes que arriba tenés la lista final_names_list

        # Si NADIE dijo nada y NADIE existe, no hacemos nada
        if not speakers_2 and not non_speakers_2:
            pass

        # Si NADIE dijo nada, pero sí hay imputados (non_speakers_2 no vacío)
        elif not speakers_2 and non_speakers_2:
            # Todos guardaron silencio
            partes_no = [nombre(idx) for idx in non_speakers_2]
            if len(partes_no) == 1:
                plantilla += (
                    f"<p align='justify'>Finalmente, al concederse la última palabra, "
                    f"{partes_no[0]} manifestó que no haría uso de ella.</p>"
                )
            else:
                no_str = ", ".join(partes_no[:-1]) + f" y {partes_no[-1]}"
                plantilla += (
                    f"<p align='justify'>Finalmente, al concederse la última palabra, "
                    f"{no_str} manifestaron que no harían uso de ella.</p>"
                )

        # Si AL MENOS UNO dijo algo
        else:
            total_speakers = len(speakers_2)
            # CASO A: Solo uno habló y ninguno guardó silencio
            if total_speakers == 1 and not non_speakers_2:
                idx_speaker, text_speaker = speakers_2[0]
                nm = nombre(idx_speaker)
                plantilla += (
                    f"<p align='justify'>Finalmente, al concederse la última palabra, "
                    f"{nm} dijo: “{text_speaker}”.</p>"
                )

            # CASO B: Más de uno habló, o hay alguno que no habló
            else:
                # Imprimimos ordenadamente a cada uno de los que sí hablaron
                for i, (idx_speaker, text_speaker) in enumerate(speakers_2):
                    nm = nombre(idx_speaker)
                    if i == 0:
                        # Primer orador
                        plantilla += (
                            f"<p align='justify'>Finalmente, al concederse la última palabra, "
                            f"{nm} dijo: “{text_speaker}”.</p>"
                        )
                    else:
                        # Siguientes oradores
                        plantilla += (
                            f"<p align='justify'>Seguidamente, {nm} dijo: “{text_speaker}”.</p>"
                        )

                # Ahora mencionamos a los que NO hablaron
                if non_speakers_2:
                    partes_no = [nombre(idx) for idx in non_speakers_2]
                    if len(partes_no) == 1:
                        plantilla += (
                            f"<p align='justify'>Por último, {partes_no[0]} "
                            "manifestó que no haría uso de la palabra.</p>"
                        )
                    else:
                        no_str = ", ".join(partes_no[:-1]) + f" y {partes_no[-1]}"
                        plantilla += (
                            f"<p align='justify'>Por último, {no_str} "
                            "manifestaron que no harían uso de la palabra.</p>"
                        )

        # Valoración de la prueba (corto y pego):
        plantilla += f"<p><b>5. Valoración de la prueba:</b> "
        caso_vf = self.var_caso_vf.currentText().strip()
        if n_hec == 1:
            el_los_hecho_s = "el hecho"
            ocurrio_eron = "ocurrió"
            han_sido_text = "ha sido"
        else:
            el_los_hecho_s = "los hechos"
            ocurrio_eron = "ocurrieron"
            han_sido_text = "han sido"
        
        if caso_vf == "No":
            if n_imp == 1:
                imputado_phrase = "del acusado" if sexos[0] == "M" else "de la acusada"
            else:
                if cant_fem == n_imp:
                    imputado_phrase = "de las acusadas"
                else:
                    imputado_phrase = "de los imputados"
            le_les = "le" if n_imp == 1 else "les"
            plantilla += f"los elementos de juicio enunciados y los argumentos desarrollados en la acusación base del juicio de la causa aquí juzgada, sumados a la argumentación del fiscal al momento emitir las conclusiones, en las que solicitó la condena –todo lo cual hago mío por razones de brevedad– satisfacen plenamente el estándar probatorio requerido para tener por acreditada la plataforma fáctica bajo análisis y la participación {imputado_phrase} tal como {le_les} ha sido atribuida.</p>"
        else:
            if caso_vf in ("violencia de género", "violencia familiar", "violencia de género doméstica"):
                if n_imp == 1:
                    imputado_phrase = "del acusado" if sexos[0] == "M" else "de la acusada"
                else:
                    imputado_phrase = "de las acusadas" if all(s == "F" for s in sexos) else "de los imputados"
                le_les = "le" if n_imp == 1 else "les"
                if caso_vf == "violencia de género doméstica":
                    plantilla += f"{el_los_hecho_s} motivo de juzgamiento configuran un caso de violencia de género doméstica. De acuerdo con ello, debe recordarse que el rasgo característico de la violencia de género es el posicionamiento del varón, respecto de la mujer, en una condición de superioridad, a través de cualquiera de los tipos de violencia (art. 5, ley 26485), y en desmedro de su derecho a contar con un ámbito de determinación para su personal proyecto de vida; de allí la demostración de poder, dominación o control por la violencia (TSJ, Sala Penal, S. nº 273, 23/06/2016, “Medina”, entre otros). Estos casos, a su vez, tienen “...particularidades que los diferencian de otros delitos pues aquí la víctima sufre reiterados comportamientos agresivos, una escalada de violencia cada día o semana más agravada y de mayor riesgo, caracterizada por su duración, multiplicidad y aumento de gravedad. Precisamente, el contexto de violencia, comprendido como un fenómeno de múltiples ofensas de gravedad progresiva que se extienden a través del tiempo, debe ser ponderado en su capacidad de suministrar indicios… Máxime, cuando estos hechos ocurren en un marco de vulnerabilidad, dado que raramente se realizan a la vista de terceros, porque una de las características de la dominación por violencia en sus múltiples manifestaciones es precisamente el aislamiento de la víctima. Las particulares características de los hechos de violencia doméstica y de género, hace que cobre especial relevancia, como también sucede con la violencia sexual, el relato de la víctima, el que adquiere un valor convictivo de preferente ponderación en la medida que resulte fiable y se encuentre corroborado por indicios, siempre que éstos tengan una confluencia de conjunto que conduzcan a dotar de razón suficiente la conclusión…” (TSJ, Sala Penal, S. n° 84, 04/05/2012, “Sánchez”, entre muchos otros). Y en función de tales circunstancias, es necesario abordar su investigación y juzgamiento bajo un criterio de amplitud probatoria (TSJ, Sala Penal, S. n° 266, 15/10/2010, “Agüero”; S. nº 28, 11/3/2014, “Sosa”; S. n° 182, 26/05/2017, “Oviedo”; entre muchos otros). Tales exigencias derivan de la obligación de debida diligencia que impone el conjunto de instrumentos internacionales ratificados por nuestro país para este tipo de casos (arts. 7 “b”, Convención Interamericana para Prevenir, Sancionar y Erradicar la Violencia contra la Mujer –Belém do Pará-, 2 “c”, CEDAW). A partir de dicho marco, considero que los elementos de juicio enunciados y los argumentos desarrollados en la acusación base del juicio, sumados a la argumentación del fiscal al momento de emitir las conclusiones, en las que solicitó la condena –todo lo cual hago mío por razones de brevedad– satisfacen plenamente el estándar probatorio requerido para tener por acreditada la plataforma fáctica bajo análisis y la intervención {imputado_phrase} tal como {le_les} ha sido atribuida.</p>"
                else:
                    plantilla += f"{el_los_hecho_s} motivo de juzgamiento configuran un caso de {caso_vf}. Los elementos de juicio enunciados y los argumentos desarrollados en la acusación base del juicio de la causa aquí juzgada, sumados a la argumentación del fiscal al momento emitir las conclusiones, en las que solicitó la condena –todo lo cual hago mío por razones de brevedad– satisfacen plenamente el estándar probatorio requerido para tener por acreditada la plataforma fáctica y la intervención {imputado_phrase} tal como {le_les} ha sido atribuida.</p>"


        pruebas_text = self.var_pruebas_importantes
        plantilla += (
            f"<p align='justify'>Al examinar el contenido de tales evidencias, las encuentro suficientes para dictar una condena, "
            f"pues –sin espacio para el principio según el cual la duda debe favorecer a la persona imputada– ponen de manifiesto que "
            f"{el_los_hecho_s} {ocurrio_eron} tal como {han_sido_text} en la acusación (TSJ, Sala Penal, “Bergamaschi”, S. n° 363, "
            f"26/0872021; “Moreira”, S. n° 361, 26/09/2022, entre otros). Tal confluencia es la que emerge, en especial a partir de "
            f"{pruebas_text}.</p>"
        )

        if n_imp == 1:
            defensa_phrase = "del imputado" if sexos[0] == "M" else "de la imputada"
            acusado_singular_plural = "el imputado" if sexos[0] == "M" else "la imputada"
            es_son = "es"
            debe_s = "debe"
            se_hallaran = "se hallara"
            responsable_s = "responsable"
            tal_es = "tal"
        else:
            defensa_phrase = "de las acusadas" if cant_fem == n_imp else "de los imputados"
            acusado_singular_plural = "las acusadas" if all(s == "F" for s in sexos) else "los imputados"
            es_son = "son"
            debe_s = "deben"
            se_hallaran = "se hallaran"
            responsable_s = "responsables"
            tal_es = "tales"

        plantilla += (
            f"<p align='justify'>Agrego que esta contundencia probatoria ha sido expresamente admitida por la defensa técnica "
            f"{defensa_phrase} durante la audiencia. Ello ocurrió, además, en un contexto en el que el tribunal se aseguró de corroborar "
            f"que {acusado_singular_plural} {se_hallaran} en plenas condiciones de libertad para reconocer su responsabilidad, "
            f"que comprendieran la naturaleza de lo que asentían y el alcance de los hechos que luego reconocieron y sus consecuencias jurídicas. "
            f"Más allá de lo ya expuesto, el contenido de la prueba y los fundamentos de la acusación constan en el expediente, "
            f"y las conclusiones de las partes han quedado en el registro fílmico de la audiencia. A todo ello me remito para su consulta si "
            f"fuere necesario, pues cualquier transcripción adicional de todo o parte de tal motivación de la premisa fáctica supondría un "
            f"desgaste innecesario e inútil que, incluso, contradiría los objetivos de economía y celeridad a los que se orienta la modalidad "
            f"abreviada de juicio elegida. Cabe recordar, en este sentido, que tanto el máximo tribunal de la Nación como el de la Provincia, "
            f"han sostenido de manera constante la validez de la argumentación por remisión en la medida en que esas razones sean asequibles, "
            f"tal como ocurre en el caso (cfme., CSJN \"Macasa S.A. v/ Caja Popular de Ahorro...\", Fallos 319:308; TSJ, Sala Penal, \"Rivero\", "
            f"S. n° 33, 9/11/1984; \"González\", S. n° 90, 16/10/2002; “Romero”, S. nº 50, 19/3/2008; entre otros). Aclaro, finalmente, "
            f"que no existen causales de inimputabilidad o de justificación (adviértase que ninguna de las partes ha hecho invocación alguna en "
            f"ese sentido), por lo que {acusado_singular_plural} {es_son} penalmente {responsable_s} y como {tal_es} {debe_s} responder.</p>"
        )

        nombres_imputados_conjunction = format_list_for_sentence(final_names_list)
        if n_hec == 1:
            el_los_hechos = "el hecho"
            dejarlo_s = "dejarlo"
            fijado_s = "fijado"
            ha_n_sido = "ha sido transcripto"
        else:
            el_los_hechos = "los hechos"
            dejarlo_s = "dejarlos"
            fijado_s = "fijados"
            ha_n_sido = "han sido transcriptos"

        texto_potenciales = ""
        if self.var_uso_terminos_potenciales.currentText() == "Sí":
            texto_potenciales = (
                ", debiendo entenderse que, con motivo de haberse arribado al grado de certeza exigido "
                "en esta instancia procesal, los términos potenciales allí utilizados deben ser comprendidos "
                "aquí de modo indicativo"
            )

        plantilla += (
            f"<p align='justify'><b>6. Conclusión:</b> en función de lo expuesto, corresponde dar por acreditada la responsabilidad "
            f"de {nombres_imputados_conjunction} en {el_los_hechos} motivo de juicio y {dejarlo_s} {fijado_s} tal como {ha_n_sido}"
            f"{texto_potenciales}. Dejo así satisfecha la exigencia impuesta en el artículo 408 inc. 3° del CPP y respondo afirmativamente "
            f"a esta primera cuestión.</p>"
        )

        frag_cal = []
        for i, imp in enumerate(self.imputados):
            nm = final_names_list[i]
            delit = imp["delitos"].text().strip()
            frag_cal.append(f"{nm} debe responder bajo el encuadre legal de {delit}")

        if len(frag_cal) == 1:
            final_calif_str = frag_cal[0]
        else:
            final_calif_str = "; ".join(frag_cal[:-1]) + "; y " + frag_cal[-1]

        salvedad = ""
        if self.var_calificacion_legal.currentText() == "Incorrecta":
            corr = self.var_correccion_calif.text().strip()
            if corr:
                salvedad = f", con la salvedad de que {corr}"

        plantilla += f"<p align='justify'><b>A LA SEGUNDA CUESTIÓN, {self.cargo_juez_en_mayusculas()} {juez_nombre.upper()} DIJO:</b></p>"
        calif_es_correcta = (self.var_calificacion_legal.currentText() == "Correcta")
        corr = self.var_correccion_calif.text().strip() if self.var_calificacion_legal.currentText() == "Incorrecta" else ""

        calif_list = []
        for delito, imput_names in delitos_dict.items():
            imput_str = format_list_for_sentence(imput_names)
            if len(imput_names) > 1:
                verbo = "deben responder"
            else:
                verbo = "debe responder"
            calif_list.append(f"{imput_str} {verbo} bajo el encuadre legal de {delito}")

        if len(calif_list) == 1:
            final_calif_str2 = calif_list[0]
        else:
            final_calif_str2 = "; ".join(calif_list[:-1]) + "; y " + calif_list[-1]

        if calif_es_correcta:
            subsuncion_line = (
                "La subsunción legal propuesta por la Fiscalía al emitir sus conclusiones resulta correcta. "
                "Dado que la subsunción legal propuesta por la Fiscalía coincide con la de la acusación base "
                "del juicio y no ha sido materia de controversia por las partes, me exime de mayores "
                "consideraciones, pues a los fines de la debida motivación jurídica de la sentencia, es "
                "suficiente la mención de la norma en la que se apoya la decisión (TSJ, Sala Penal, S. n° 190, "
                'del 11/8/2010, “Castillo”).'
            )
        else:
            salvedad_text = f", con la salvedad de que {corr}" if corr else ""
            subsuncion_line = (
                "La subsunción legal propuesta por la Fiscalía coincide con la de la acusación base "
                "del juicio y no ha sido materia de controversia por las partes, lo que me exime de "
                "mayores consideraciones, pues a los fines de la debida motivación jurídica de la "
                "sentencia, es suficiente la mención de la norma en la que se apoya la decisión "
                "(TSJ, Sala Penal, S. n° 190, del 11/8/2010, “Castillo”)"
                f"{salvedad_text}."
            )

        plantilla += (
            f"<p align='justify'>En función del modo en que se ha dado respuesta al primer interrogante, "
            f"{final_calif_str2}. {subsuncion_line}</p>"
            f"<p align='justify'>Así respondo a la presente cuestión.</p>"
        )

        plantilla += (
            f"<p align='justify'><b>A LA TERCERA CUESTIÓN, {self.cargo_juez_en_mayusculas()} {juez_nombre.upper()} DIJO:</b></p>"
        )

        if n_imp == 1:
            plantilla += (
                "<p align='justify'><b>1. Pena:</b> Para graduar la sanción a imponer, tengo en cuenta las pautas "
                "objetivas y subjetivas de mensuración de la pena establecidas en los arts. 40 y 41 del CP.</p>"
            )
        else:
            plantilla += (
                "<p align='justify'><b>1. Pena:</b> Para graduar las sanciones a imponer, tengo en cuenta las pautas "
                "objetivas y subjetivas de mensuración de la pena establecidas en los arts. 40 y 41 del CP.</p>"
            )

        introductions = ["respecto de", "en cuanto a", "con relación a", "en lo relativo a"]
        valuation_verbs = ["estimo", "valoro", "pondero", "considero"]

        for i, imp in enumerate(self.imputados):
            nm = final_names_list[i]
            pautas_str = imp["pautas"].text().strip()
            intro = introductions[i % len(introductions)]
            verb = valuation_verbs[i % len(valuation_verbs)]
            if i == 0:
                plantilla += f"<p align='justify'>Así, {intro} {nm}, {verb} {pautas_str}.</p>"
            else:
                capital_intro = intro[0].upper() + intro[1:]
                plantilla += f"<p align='justify'>{capital_intro} {nm}, {verb} {pautas_str}.</p>"

        introductions_2 = [
            "Asimismo,",
            "En el mismo sentido,",
            "De igual manera,",
            "Del mismo modo,"
        ]

        for i, imp in enumerate(self.imputados):
            nm = final_names_list[i]
            condena_str = imp["condena"].text().strip()
            if i == 0:
                plantilla += (
                    f"<p align='justify'>Por ello, teniendo en especial consideración el límite máximo que "
                    f"impone el art. 415 del CPP al Tribunal para la individualización judicial "
                    f"de la pena, al establecer que no se podrá aplicar una pena más grave que "
                    f"la pedida por el Representante del Ministerio Público Fiscal y acordada con "
                    f"el acusado y su defensor, ni modificar su forma de ejecución, corresponde "
                    f"imponerle a {nm}, para su tratamiento penitenciario, la pena de {condena_str}.</p>"
                )
            else:
                intro2 = introductions_2[(i - 1) % len(introductions_2)]
                plantilla += (
                    f"<p align='justify'>{intro2} corresponde imponerle a {nm} la pena de {condena_str}.</p>"
                )

        next_section = 2

        if self.var_decomiso_option.currentText() == "Sí":
            # 1) tomo el HTML completo (o el default)
            html_decomiso = self.var_decomiso_text.property("html") or self.TEXTO_DECOMISO_DEFECTO
            # 2) lo incrusto directamente en la plantilla (preserva <b>, <i>…)
            plantilla += (
                f"<p align='justify'><b>{numero_romano(next_section)}. Decomiso:</b> {html_decomiso}</p>"
            )
            next_section += 1

        # Honorarios
        imputados_publicos = []
        sexos_publicos = []
        defensores_publicos = set()
        imputados_privados = []
        for i, imp in enumerate(self.imputados):
            nm = final_names_list[i]
            tipo = imp["tipo_def"].currentText().strip().lower()
            def_name = imp["defensor"].text().strip()
            if tipo.startswith("púb"):
                imputados_publicos.append(nm)
                sexos_publicos.append("M" if imp["sexo_rb"][0].isChecked() else "F")
                if def_name:
                    defensores_publicos.add(def_name)
            else:
                imputados_privados.append(nm)

        if imputados_publicos:
            lista_def_pub = sorted(defensores_publicos)
            if lista_def_pub:
                nombres_defensa_publica = format_list_for_sentence(lista_def_pub)
            else:
                nombres_defensa_publica = "la Asesoría Letrada"
            cant_pub = len(imputados_publicos)
            if cant_pub == 1:
                if sexos_publicos[0] == "M":
                    phrase_al = "al imputado"
                    phrase_benef = "beneficiario"
                else:
                    phrase_al = "a la imputada"
                    phrase_benef = "beneficiaria"
            else:
                if all(s == "M" for s in sexos_publicos):
                    phrase_al = "a los imputados"
                    phrase_benef = "beneficiarios"
                elif all(s == "F" for s in sexos_publicos):
                    phrase_al = "a las imputadas"
                    phrase_benef = "beneficiarias"
                else:
                    phrase_al = "a los imputados"
                    phrase_benef = "beneficiarios"
            plantilla += (
                f"<p align='justify'><b>{numero_romano(next_section)}. Honorarios y eximición de tasa de justicia:</b> "
                f"por otra parte, debe retribuirse la labor prestada por la defensa pública a cargo de "
                f"{nombres_defensa_publica}, la que, conforme las reglas cualitativas del artículo 39 de la ley arancelaria, "
                f"estimo adecuado fijar en la suma de 30 jus (arts. 24, 36, 39, 89, 90 y cc. Ley 9459), y a la vez eximir {phrase_al} "
                f"del pago de la tasa de justicia por ser {phrase_benef} de la asistencia jurídica gratuita (art. 31 ley 7982).</p>"
            )
            next_section += 1

        if imputados_privados:
            cant_priv = len(imputados_privados)
            verbo_abonar = "abone" if cant_priv == 1 else "abonen"
            nombres_privados_str = format_list_for_sentence(imputados_privados)
            plantilla += (
                f"<p align='justify'><b>{numero_romano(next_section)}. Tasa de justicia:</b> corresponde emplazar a {nombres_privados_str} "
                f"para que, en el plazo de quince días desde que quede firme la presente sentencia, {verbo_abonar} la suma equivalente a 1,5 "
                f"jus en concepto de Tasa de Justicia, bajo apercibimiento de certificarse su existencia y librarse título para su remisión "
                f"a la Oficina de Tasa de Justicia del Área Administración del Poder Judicial a los fines de su ejecución (arts. 295 y cc "
                f"del Código Tributario Provincial, ley 6006 y sus modificatorias).</p>"
            )
            next_section += 1

        if self.var_restriccion_option.currentText() == "Sí":
            html_restriccion = self.var_restriccion_text.property("html") or self.TEXTO_RESTRICCION_DEFECTO
            plantilla += (
                f"<p align='justify'><b>{numero_romano(next_section)}. Restricción de contacto y acercamiento:</b> {html_restriccion}</p>"
            )
            next_section += 1

        caso_vf_lower = self.var_caso_vf.currentText().lower()
        if caso_vf_lower in ("violencia de género", "violencia de género doméstica", "violencia familiar"):
            extra_ley = " y por el art. 28 de la Ley provincial 9283"
        else:
            extra_ley = ""

        victima_text = self.var_victima.text().strip()
        if not victima_text:
            victims_pronoun = "la persona damnificada"
            require_phrase = "requerírsele"
            volunt_phrase = "manifieste su voluntad"
        else:
            splitted = [v.strip() for v in victima_text.split(",") if v.strip()]
            if len(splitted) <= 1:
                victims_pronoun = "la persona damnificada"
                require_phrase = "requerírsele"
                volunt_phrase = "manifieste su voluntad"
            else:
                victims_pronoun = "las personas damnificadas"
                require_phrase = "requerírseles"
                volunt_phrase = "manifiesten su voluntad"

        plantilla += (
            f"<p align='justify'><b>{numero_romano(next_section)}. Comunicaciones:</b> finalmente, de conformidad a lo dispuesto "
            f"por el art. 11 bis –penúltimo párrafo– de la Ley 24660{extra_ley}, así como por el art. 96 del CPP, debe informarse "
            f"lo resuelto a {victims_pronoun} y {require_phrase} que {volunt_phrase} en relación a las facultades que les corresponde "
            f"a partir del dictado de esta sentencia. También se deberá efectuar el cómputo de pena y formar el legajo de ejecución "
            f"(art. 4 del Acuerdo Reglamentario nº 896, Serie A, del Excmo. Tribunal Superior de Justicia) y, una vez que quede firme "
            f"la presente sentencia, oficiar al Registro Nacional de Reincidencia a los fines del art. 2° de la Ley 22117.</p>"
        )
        next_section += 1

        # Mantener el cierre y título de “RESUELVO”
        plantilla += (
            "<p align='justify'>Así respondo a la presente cuestión.</p>"
            "<p align='justify'>Por todo lo expuesto, y normas legales citadas, <b>RESUELVO:</b></p>"
        )

        # ── Resuelvo ───────────────────────────────────────────
        html_resuelvo = self.var_resuelvo.property("html") or ""
        if html_resuelvo:
            plantilla += f"<p align='justify'>{html_resuelvo}</p>"


        plantilla = f'<div style="text-align: justify;">{plantilla}</div>'

        old_plain = self._prev_plain
        self.texto_plantilla.setHtml(plantilla)
        self.texto_plantilla.setAlignment(Qt.AlignJustify)

        new_plain = self.texto_plantilla.toPlainText()
        if old_plain:
            self._highlight_diff(old_plain, new_plain)
        self._prev_plain = new_plain

        QTimer.singleShot(0, lambda: self.texto_plantilla.verticalScrollBar().setValue(pos))
    
    def _sync_imp(self, idx: int, key: str, value: str):
        while len(self.data.imputados) <= idx:
            self.data.imputados.append({})
        self.data.imputados[idx][key] = value.strip()


from PySide6.QtWidgets import QApplication, QMessageBox

def confirm_and_quit(widget) -> None:
    """Muestra un QMessageBox; si el usuario acepta, cierra TODA la app."""
    ans = QMessageBox.question(
        widget,
        "Cerrar la aplicación",
        "¿Está seguro de que desea salir?\nSe cerrarán todas las ventanas.",
        QMessageBox.Yes | QMessageBox.No,
        QMessageBox.No
    )
    if ans == QMessageBox.Yes:
        QApplication.quit()


if __name__ == "__main__":
    from PySide6.QtWidgets import QApplication
    import sys
    app = QApplication(sys.argv)
    data = CausaData()                # instancia propia de prueba
    w = SentenciaWidget(data)
    w.resize(1300, 700)
    w.show()
    sys.exit(app.exec()) 