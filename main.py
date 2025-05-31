# -*- coding: utf-8 -*-
"""
Generador de documentos judiciales – PySide6
Interfaz: datos generales + pestañas de imputados (sin colores forzados)
"""
import sys, os, json
from datetime import datetime
from pathlib import Path

from PySide6.QtCore    import Qt
from PySide6.QtGui     import QIcon, QClipboard, QAction
from html import unescape
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QLabel, QLineEdit, QTextEdit,
    QComboBox, QPushButton, QGridLayout, QVBoxLayout, QTabWidget,
    QFileDialog, QMessageBox, QSplitter, QCheckBox, QScrollArea, QDialog,
    QDialogButtonBox, QRadioButton, QButtonGroup
)
from PySide6.QtGui import QFont
from PySide6.QtGui import QTextBlockFormat
from PySide6.QtGui import QTextCharFormat
from PySide6.QtGui import QTextCursor 
from datetime import timedelta
from core_data import CausaData
from tramsent import SentenciaWidget
from docx import Document
from PySide6.QtCore import QSignalBlocker
from sentencia_window import SentenciaWindow
from PySide6.QtGui import QTextDocument
import re
from PySide6.QtCore import QMimeData
from PySide6.QtWidgets import QHBoxLayout
from widgets import NoWheelComboBox
from widgets import NoWheelComboBox, NoWheelSpinBox
def _DEBUG_unicode(tag: str, txt: str, n: int = 120):
    # imprime los primeros “n” caracteres con su code-point
    print(f"\n{tag}:")
    for c in txt[:n]:
        print(f"{repr(c)} U+{ord(c):04X}", end="  ")
    print("…\n")


def resource_path(rel: str) -> str:
    if getattr(sys, "frozen", False):
        return os.path.join(sys._MEIPASS, rel)          # type: ignore
    return os.path.join(os.path.abspath("."), rel)

CAUSAS_DIR = Path("causas_guardadas")
CAUSAS_DIR.mkdir(exist_ok=True)

_UNIDADES = (
    '', 'uno', 'dos', 'tres', 'cuatro', 'cinco', 'seis',
    'siete', 'ocho', 'nueve', 'diez', 'once', 'doce',
    'trece', 'catorce', 'quince', 'dieciséis', 'diecisiete',
    'dieciocho', 'diecinueve', 'veinte', 'veintiuno',
    'veintidós', 'veintitrés', 'veinticuatro', 'veinticinco',
    'veintiséis', 'veintisiete', 'veintiocho', 'veintinueve'
)
_DECENAS  = ('treinta', 'cuarenta', 'cincuenta', 'sesenta',
             'setenta', 'ochenta', 'noventa')
_CENTENAS = ('ciento', 'doscientos', 'trescientos', 'cuatrocientos',
             'quinientos', 'seiscientos', 'setecientos', 'ochocientos',
             'novecientos')
# -- helper -------------------------------------------------------------
_rx_bold      = re.compile(r'<span[^>]*font-weight:600[^>]*>(.*?)</span>', re.S)
_rx_italic    = re.compile(r'<span[^>]*font-style:italic[^>]*>(.*?)</span>', re.S)
_rx_bold_it   = re.compile(r'<span[^>]*font-weight:600[^>]*font-style:italic[^>]*>(.*?)</span>', re.S)
_rx_spans     = re.compile(r'<span[^>]*>(.*?)</span>', re.S)
_rx_p_cleanup = re.compile(r'<p style="[^"]*text-align:([^";]+)[^"]*">')
    # -------------------------------------------------------------------------
    #  helpers internos de RTF – pégalos encima de copy_to_clipboard
_rx_tag      = re.compile(r'<(/?)(b|strong|i|em|u|p)(?:\s+[^>]*)?>', re.I)
_rx_p_align  = re.compile(r'text-align\s*:\s*(left|right|center|justify)', re.I)
def _qt_to_simple_html(qt_html: str) -> str:
    """Convierte el HTML verboso de Qt en etiquetas básicas (<b>, <i>, …)."""
    h = qt_html
    # negrita + cursiva (hay que hacerla antes que las otra-s)
    h = _rx_bold_it.sub(r'<b><i>\1</i></b>', h)
    # solo negrita
    h = _rx_bold.sub(r'<b>\1</b>', h)
    # solo cursiva
    h = _rx_italic.sub(r'<i>\1</i>', h)
    # cualquier otro <span …> -> quitar <span>
    h = _rx_spans.sub(r'\1', h)
    # <p style=" … text-align:justify …">  ->  <p style="text-align:justify;">
    h = _rx_p_cleanup.sub(lambda m: f'<p style="text-align:{m.group(1)};">', h)
    return h
def _html_to_rtf_fragment(html: str) -> str:
    """
    Convierte un HTML muy sencillo (p, b/strong, i/em, u)
    a la secuencia RTF equivalente.
    """
    rtf = []
    stack = []              # para llevar el estado <b>, <i>, <u>

    pos = 0
    for m in _rx_tag.finditer(html):
        text = html[pos:m.start()]
            # --- texto normal (escapar) --------------------------------------
        text = (text.replace('\\', r'\\')
                    .replace('{',  r'\{')
                    .replace('}',  r'\}')
                    .replace('&nbsp;', ' '))
        rtf.append(text)
        pos = m.end()
        closing, tag = m.group(1), m.group(2).lower()
        if tag == 'p':                                     # <p …>
            if closing:                    # </p>
                rtf.append(r'\par ')
            else:                          # <p …>
                style = m.group(0)
                alg   = 'justify'          # por defecto
                ma = _rx_p_align.search(style)
                if ma:
                    alg = ma.group(1).lower()
                rtf.append(r'\pard')
                rtf.append({'left': r'\ql ',
                            'right': r'\qr ',
                            'center': r'\qc ',
                            'justify': r'\qj '}[alg])
        elif tag in ('b', 'strong'):
            rtf.append(r'\b0 ' if closing else r'\b ')
        elif tag in ('i', 'em'):
            rtf.append(r'\i0 ' if closing else r'\i ')
        elif tag == 'u':
            rtf.append(r'\ulnone ' if closing else r'\ul ')

        # resto del texto tras la última etiqueta
    tail = html[pos:]
    tail = (tail.replace('\\', r'\\')
                .replace('{',  r'\{')
                .replace('}',  r'\}')
                .replace('&nbsp;', ' '))
    rtf.append(tail)

    return ''.join(rtf)
    # -------------------------------------------------------------------------
def numero_a_letras(n: int) -> str:
    if n == 0: return 'cero'
    if n == 100: return 'cien'
    if n < 30: return _UNIDADES[n]
    if n < 100:
        return _DECENAS[n//10 - 3] if n%10==0 else f"{_DECENAS[n//10 - 3]} y {_UNIDADES[n%10]}"
    if n < 1000:
        return (_CENTENAS[n//100 -1] if n%100 else _CENTENAS[n//100-1]) \
               + ('' if n%100==0 else f" {numero_a_letras(n%100)}")
    if n == 1000: return 'mil'
    if n < 2000:  return f"mil {numero_a_letras(n%1000)}"
    if n < 1_000_000:
        miles=n//1000; resto=n%1000
        txt=f"{numero_a_letras(miles)} mil"
        return txt if resto==0 else f"{txt} {numero_a_letras(resto)}"
    return str(n)

_MESES = ['','enero','febrero','marzo','abril','mayo','junio',
          'julio','agosto','septiembre','octubre','noviembre','diciembre']
def fecha_a_letras(dt: datetime) -> str:
    return f"{numero_a_letras(dt.day)} de {_MESES[dt.month]} de {numero_a_letras(dt.year)}"
def num_letras(n:int)->str:
    if n==0: return 'cero'
    if n==100: return 'cien'
    if n<30: return _UNIDADES[n]
    if n<100:
        return _DECENAS[n//10-3] if n%10==0 else f"{_DECENAS[n//10-3]} y {_UNIDADES[n%10]}"
    if n<1000:
        return (_CENTENAS[n//100-1] if n%100 else _CENTENAS[n//100-1]) + ('' if n%100==0 else f" {num_letras(n%100)}")
    if n==1000: return 'mil'
    if n<2000: return f"mil {num_letras(n%1000)}"
    if n<1_000_000:
        miles,resto = divmod(n,1000)
        txt=f"{num_letras(miles)} mil"
        return txt if resto==0 else f"{txt} {num_letras(resto)}"
    return str(n)
def fecha_letras(dt:datetime)->str:
    return f"{num_letras(dt.day)} de {_MESES[dt.month]} de {num_letras(dt.year)}"

def _sanitize_html(html_raw: str) -> str:
    """
    Devuelve SOLO el fragmento que estaba dentro de <body>,
    manteniendo <b>, <i>, <u> y quitando todo estilo / saltos raros.
    """
    import re, html

    # A)  ───── EXTRAEMOS SOLO <body> … </body> ─────
    m = re.search(r'<body[^>]*>(.*?)</body>', html_raw, flags=re.I | re.S)
    if m:
        html_raw = m.group(1)
    # (si por algún motivo no hay <body>, seguimos con lo que venga)

    # B)  ───── A partir de aquí van los pasos que ya tenías ─────
    # a) <strong>/<em> → <b>/<i>
    html_raw = re.sub(r'</?strong>', lambda m: '<b>' if m.group(0)[1] != '/' else '</b>', html_raw, flags=re.I)
    html_raw = re.sub(r'</?em>',     lambda m: '<i>' if m.group(0)[1] != '/' else '</i>', html_raw, flags=re.I)

    # b) <span style="font-weight:...">…</span> → <b>…</b>
    html_raw = re.sub(
        r'<span[^>]*style="[^"]*font-weight\s*:\s*(?:bold|700)[^"]*"[^>]*>(.*?)</span>',
        r'<b>\1</b>',
        html_raw,
        flags=re.I | re.S
    )

    # c) quitamos atributos style, class, dir, lang…
    html_raw = re.sub(r'\s*(style|class|dir|lang)="[^"]*"', '', html_raw, flags=re.I)

    # d) quitamos cualquier <span> remanente
    html_raw = re.sub(r'</?span[^>]*>', '', html_raw, flags=re.I)

    # d-bis) fuera <br>
    html_raw = re.sub(r'(?i)<br\s*/?>', ' ', html_raw)

    # d-ter) fuera párrafos vacíos de Qt
    html_raw = re.sub(
        r'<p[^>]*-qt-paragraph-type:empty[^>]*>\s*(<br\s*/?>)?\s*</p>',
        ' ',
        html_raw,
        flags=re.I
    )

    # e) limpia saltos y nbsp
    html_raw = re.sub(
        r'(\r\n|\r|\n|&#10;|&#13;|\u2028|\u2029|&nbsp;)',
        ' ',
        html_raw
    )

    # f) compacta espacios
    html_raw = re.sub(r'\s+', ' ', html_raw).strip()

    # g) si quedó un único <p> que envuelve todo, lo quitamos
    if html_raw.lower().startswith('<p') and html_raw.lower().endswith('</p>'):
        html_raw = re.sub(r'^<p[^>]*>|</p>$', '', html_raw, flags=re.I).strip()

    return html.unescape(html_raw)

class MainWindow(QMainWindow):
    def __init__(self, data: CausaData, parent=None): 
        super().__init__()
        self.data = data
        self.setWindowTitle("Tramites del 415")
        self.resize(1100, 610)
        self.setWindowIcon(QIcon(resource_path("icono5.ico")))
        self.imputados_widgets: list[dict[str, object]] = []
        self.tabs_imp: QTabWidget | None = None

        # ---------- splitter (izq. datos | der. plantillas) -----------------
        splitter = QSplitter(Qt.Horizontal, self)
        self.setCentralWidget(splitter)
        splitter.setSizes([400, 700])
        # ---------- panel de texto (derecha) -------------------------------
        right_panel  = QWidget()
        right_layout = QVBoxLayout(right_panel)

        self.selector_imp = NoWheelComboBox()
        self.selector_imp.currentIndexChanged.connect(self.update_for_imp)
        right_layout.addWidget(self.selector_imp)

        self.tabs_txt = QTabWidget()              # plantillas generadas
        right_layout.addWidget(self.tabs_txt, 1)  # “1” => ocupa todo el resto
        splitter.addWidget(right_panel)
        splitter.setStretchFactor(1, 1)

        self.text_edits: dict[str, QTextEdit] = {}
        for name in (
            "Pedido de audiencia", "Oficio OGA", "Decreto audiencia",
            "Oficio notificación", "Acta renuncia", "Constancia grabación",
            "Certificado víctimas", "Oficio Neuro", "Oficio CIV",
            "Oficio libertad", "Oficio Policía", "Oficio Reincidencia",
            "Oficio cómputo", "Oficio SPC", "Oficio comunicación",
            "Legajo", "Puesta a disposición"
        ):
            te = QTextEdit(); te.setReadOnly(True)
            te.setFontFamily("Times New Roman"); te.setFontPointSize(12)
            cont = QWidget(); lay = QVBoxLayout(cont)
            lay.addWidget(te)
            btn = QPushButton("Copiar al portapapeles")
            btn.clicked.connect(lambda _=False, t=te: self.copy_to_clipboard(t))
            lay.addWidget(btn)
            if name == "Oficio OGA":
                btn_gen = QPushButton("Generar planilla para OGA")
                btn_gen.clicked.connect(self.generate_planilla_oga)
                lay.addWidget(btn_gen)
            self.tabs_txt.addTab(cont, name)
            self.text_edits[name] = te

        left_scroll = QScrollArea(); left_scroll.setWidgetResizable(True)
        left_inner  = QWidget();     left_scroll.setWidget(left_inner)
        splitter.insertWidget(0, left_scroll)
        splitter.setStretchFactor(0, 0)

        self.form = QGridLayout(left_inner)
        self.form.setAlignment(Qt.AlignTop)
        left_inner.setMinimumWidth(380)
        self._row = 0

        def label(text: str): self.form.addWidget(QLabel(text), self._row, 0)

        def add_line(attr: str, text: str) -> QLineEdit:
            label(text)
            le = QLineEdit(); le.textChanged.connect(self.update)
            self.form.addWidget(le, self._row, 1); self._row += 1
            setattr(self, attr, le); return le

        def add_combo(attr: str, text: str, items: list[str], editable=False) -> QComboBox:
            label(text)
            cb = NoWheelComboBox(); cb.addItems(items); cb.setEditable(editable)
            cb.currentIndexChanged.connect(self.update_template)
            cb.editTextChanged.connect(self.update_template)
            self.form.addWidget(cb, self._row, 1); self._row += 1
            setattr(self, attr, cb); return cb

        label("Número de imputados:")
        self.combo_n = NoWheelComboBox(); self.combo_n.addItems([str(i) for i in range(1, 21)]); self.combo_n.currentIndexChanged.connect(self.update_template)

        self.form.addWidget(self.combo_n, self._row, 1); self._row += 1


        self.entry_caratula   = add_line('entry_caratula',   "Carátula:")
        self.combo_articulo   = add_combo('combo_articulo',  "Cámara o juzgado:",
                                        ["Cámara en lo Criminal y Correccional", "Juzgado de Control"])
        opciones_tribunal = [
            "Cámara en lo Criminal y Correccional de Primera Nominación", 
            "Cámara en lo Criminal y Correccional de Segunda Nominación", 
            "Cámara en lo Criminal y Correccional de Tercera Nominación", 
            "Cámara en lo Criminal y Correccional de Cuarta Nominación", 
            "Cámara en lo Criminal y Correccional de Quinta Nominación", 
            "Cámara en lo Criminal y Correccional de Sexta Nominación", 
            "Cámara en lo Criminal y Correccional de Séptima Nominación", 
            "Cámara en lo Criminal y Correccional de Octava Nominación", 
            "Cámara en lo Criminal y Correccional de Novena Nominación", 
            "Cámara en lo Criminal y Correccional de Décima Nominación", 
            "Cámara en lo Criminal y Correccional de Onceava Nominación", 
            "Cámara en lo Criminal y Correccional de Doceava Nominación", 
            "Juzgado de Control en lo Penal Económico", 
            "Juzgado de Control y Faltas N° 2", 
            "Juzgado de Control y Faltas N° 3", 
            "Juzgado de Control y Faltas N° 4", 
            "Juzgado de Control y Faltas N° 5", 
            "Juzgado de Control en Violencia de Género y Familiar N° 1", 
            "Juzgado de Control en Violencia de Género y Familiar N° 2", 
            "Juzgado de Control y Faltas N° 7", 
            "Juzgado de Control y Faltas N° 8", 
            "Juzgado de Control y Faltas N° 9", 
            "Juzgado de Control y Faltas N° 10",
            "Juzgado de Control y Faltas N° 11", 
            "Juzgado de Control de Lucha contra el Narcotráfico"
        ]

        # Reemplazamos entry_tribunal por un combo editable
        self.entry_tribunal = add_combo(
            'entry_tribunal',            # atributo
            "Tribunal:",                 # etiqueta
            opciones_tribunal,           # ítems del combo
            editable=True                # lo hacemos editable
        )
        self.entry_secretaria = add_line('entry_secretaria', "Secretaría a cargo de:")
        self.entry_fiscal     = add_line('entry_fiscal',     "Fiscal (con su fiscalía):")
        self.entry_fecha      = add_line('entry_fecha',      "Fecha de audiencia:")
        horas = [f"{h:02d}:{m:02d}" for h in range(24) for m in (0, 30)]
        self.combo_hora       = add_combo('combo_hora',      "Hora de audiencia:", horas)
        salas = [f"Sala OGA {i} del MOPLO" for i in range(1, 11)] + \
                ["Sala de audiencias de la Cámara en lo Criminal y Correccional"]
        self.combo_sala       = add_combo('combo_sala',      "Sala de audiencia:", salas, editable=True)
        self.entry_funcionario= add_line('entry_funcionario',"Funcionario que firma:")

        self.entry_sentencia  = add_line('entry_sentencia',  "Sentencia (número y fecha):")

        # 1) preview
        self.var_resuelvo    = QTextEdit()
        self.var_resuelvo.setReadOnly(True)
        self.var_resuelvo.setFixedHeight(80)
        base_font = QFont("Times New Roman", 12)
        self.var_resuelvo.setFont(base_font)
        self.var_resuelvo.document().setDefaultFont(base_font)
        # alias para compatibilidad con el resto
        self.entry_resuelvo  = self.var_resuelvo

        # 2) botón
        btn_edit_res = QPushButton("Editar Resuelvo…")
        btn_edit_res.clicked.connect(self.abrir_ventana_resuelvo)

        # 3) lo metemos en el form
        row = self._row
        self.form.addWidget(QLabel("Resuelvo:"), row, 0)
        container = QWidget()
        hlay = QHBoxLayout(container)
        hlay.addWidget(self.var_resuelvo, 1)
        self.var_resuelvo.hide()
        hlay.addWidget(btn_edit_res)
        self.form.addWidget(container, row, 1)
        self._row += 1

        self.entry_firmantes  = add_line('entry_firmantes',  "Firmantes de la sentencia:")
        self.combo_renuncia   = add_combo('combo_renuncia',  "Renuncia a casación:", ["Sí", "No"])

        label("Número de hechos:")
        self.spin_hechos = NoWheelSpinBox(); self.spin_hechos.setRange(1, 15)
        self.spin_hechos.valueChanged.connect(self.rebuild_hechos)
        self.form.addWidget(self.spin_hechos, self._row, 1); self._row += 1

        self.tabs_hechos = QTabWidget()
        self.form.addWidget(self.tabs_hechos, self._row, 0, 1, 2)
        self._row += 1
        
        # === Cargar datos existentes desde self.data ===
        self.entry_caratula.setText(self.data.caratula)
        self.entry_tribunal.setCurrentText(self.data.tribunal)
        self.entry_secretaria.setText(getattr(self.data, "secretaria", ""))
        self.entry_fecha.setText(self.data.fecha_audiencia)
        self.combo_hora      .setCurrentText(getattr(self.data, "hora_audiencia", ""))
        self.combo_sala.setCurrentText(self.data.sala)
        self.entry_funcionario.setText(getattr(self.data, 'funcionario', ''))
        self.entry_fiscal.setText(self.data.fiscal_nombre)
        self.entry_sentencia .setText(getattr(self.data, "sentencia_num", ""))

        self.entry_firmantes .setText(getattr(self.data, "firmantes",   ""))
        self.combo_renuncia  .setCurrentText("Sí" if getattr(self.data, "renuncia", False) else "No")

        self.spin_hechos.setValue(getattr(self.data, "num_hechos", len(self.data.hechos) or 1))
        
        # Número de imputados (dispara rebuild_imputados con la cantidad correcta)
        self.combo_n.setCurrentText(str(self.data.n_imputados or 1))
                # ───── helpers *dentro* de __init__ ───────────────────────────
        
        # Construcción de pestañas de hechos
        self.hechos_widgets = []
        self.rebuild_hechos()

        # Construcción de pestañas de imputados
        self.tabs_imp = QTabWidget()
        self.form.addWidget(self.tabs_imp, self._row, 0, 1, 2)
        self._row += 1
        
        self.combo_n.currentIndexChanged.connect(self.rebuild_imputados)
        
        # Reconstruir dinámicamente a partir de self.data.imputados
        self.imputados_widgets = []
        self.rebuild_imputados()

        for txt, slot in (("Guardar causa", self.guardar_causa),
                        ("Abrir causa",  self.cargar_causa),
                        ("Eliminar causa",self.eliminar_causa)):
            btn = QPushButton(txt); btn.clicked.connect(slot)
            self.form.addWidget(btn, self._row, 0, 1, 2); self._row += 1
        
        btn_sentencia = QPushButton("▶ Ver sentencia")
        btn_sentencia.clicked.connect(self.abrir_sentencia)
        self.form.addWidget(btn_sentencia, self._row, 0, 1, 2)
        self._row += 1

        self.data.apply_to_main(self)
        splitter.setSizes([400, 700])
        self.update_template()


    def abrir_sentencia(self) -> None:
        """Salta a la pantalla de ‘Sentencia’."""

        # 1) Guardar los cambios hechos en Trámites
        self.data.from_main(self)

        if getattr(self, "_sent_win", None) is None:
            # instanciamos sin parent para que tenga su propia entrada en la barra de tareas
            self._sent_win = SentenciaWindow(self.data, parent=None)
            # nos guardamos el main para luego re-show() cuando cierren la sentencia
            self._sent_win.main_win = self
            self._sent_win.destroyed.connect(
                lambda _=None: setattr(self, "_sent_win", None)
            )

        # 3) Cada vez que la muestro, me aseguro de que pregunte al cerrar
        self._sent_win.skip_confirm = False          # ← diálogo de confirmación habilitado
        self._sent_win.show()

        # 4) Oculto la ventana de Trámites para que no queden dos en pantalla
        self.hide()

    def showEvent(self, ev):
        super().showEvent(ev)
        self.data.apply_to_main(self)
    def closeEvent(self, event):
        """Intercepta el cierre de la pantalla de trámites."""
        confirm_and_quit(self)
        event.ignore()      

    def show_tramites(self):
        """
        Muestra la ventana de 'Trámites'.
        Evitamos volver a llamar a show_tramites desde SentenciaWidget
        para no entrar en recursión.
        """
        # si ya hay un diálogo de sentencia abierto lo reaprovechamos
        if getattr(self, "_sent_widget", None):
            self._sent_widget.abrir_tramites()        # <<-- ¡ya no volvemos aquí!
            return

        # si no lo hay, lo creamos y, cuando aparezca, le pedimos que abra trámites
        self.abrir_sentencia(open_tramites=False)     # ← abre la sentencia “normal”
        if getattr(self, "_sent_widget", None):
            self._sent_widget.abrir_tramites()
    
    def rebuild_imputados(self):
        """Reconstruye las pestañas de imputados SIN perder lo escrito."""
            # ←– bloquea re-entradas
        if getattr(self, "_building", False):
            return
        self._building = True       
        # 1) Guardar lo ya cargado
        prev_data: list[dict[str, object]] = []
        for w in getattr(self, "imputados_widgets", []):
            dato = {}
            for k, widget in w.items():
                if isinstance(widget, QLineEdit):
                    dato[k] = widget.text()
                elif isinstance(widget, QComboBox):
                    dato[k] = widget.currentText()
                elif isinstance(widget, QCheckBox):
                    dato[k] = widget.isChecked()
            prev_data.append(dato)

        self.tabs_imp.clear()
        self.imputados_widgets = []
        n = int(self.combo_n.currentText())

        for i in range(1, n + 1):
            tab = QWidget()
            grid = QGridLayout(tab)
            row = 0

            def add_pair(text, widget):
                nonlocal row
                grid.addWidget(QLabel(text), row, 0)
                grid.addWidget(widget, row, 1)
                row += 1

            def mk_line():
                le = QLineEdit()
                le.textChanged.connect(self.update)
                return le

            def mk_combo(items, editable=False):
                cb = NoWheelComboBox()
                cb.addItems(items)
                cb.setEditable(editable)
                cb.currentIndexChanged.connect(self.update)
                cb.editTextChanged.connect(self.update)
                return cb

            w: dict[str, object] = {}
            # Crear widgets
            w['tipo'] = mk_combo(['efectiva', 'condicional'])
            add_pair("Tipo de pena:", w['tipo'])

            w['nombre'] = mk_line()
            w['nombre'].textChanged.connect(self._refresh_imp_names_in_selector)
            add_pair("Nombre y apellido:", w['nombre'])

            w['dni'] = mk_line()
            add_pair("DNI:", w['dni'])

            w['estable'] = mk_combo([
                "Complejo Carcelario n.° 1 (Bouwer)",
                "Establecimiento Penitenciario n.° 9 (UCA)",
                "Establecimiento Penitenciario n.° 3 (para mujeres)",
                "Complejo Carcelario n.° 2 (Cruz del Eje)",
                "Establecimiento Penitenciario n.° 4 (Colonia Abierta Monte Cristo)",
                "Establecimiento Penitenciario n.° 5 (Villa María)",
                "Establecimiento Penitenciario n.° 6 (Río Cuarto)",
                "Establecimiento Penitenciario n.° 7 (San Francisco)",
                "Establecimiento Penitenciario n.° 8 (Villa Dolores)"
            ], editable=True)

            add_pair("Establecimiento:", w['estable'])


            w['defensa'] = mk_line()
            add_pair("Defensa:", w['defensa'])

            w['detenc'] = mk_line()
            add_pair("Duración detención:", w['detenc'])

            w['delitos'] = mk_line()
            add_pair("Delitos atribuidos:", w['delitos'])

            w['victimas'] = mk_line()
            add_pair("Víctimas:", w['victimas'])

            w['condena'] = mk_line()
            add_pair("Condena:", w['condena'])

            w['hechos_n'] = mk_combo(['uno', 'más'])
            add_pair("Hechos (uno/más):", w['hechos_n'])

            w['fechas'] = mk_line()
            add_pair("Fechas de los hechos:", w['fechas'])

            w['decreto'] = mk_line()
            add_pair("Decreto cómputo:", w['decreto'])

            w['firm_dec'] = mk_line()
            add_pair("Firmantes cómputo:", w['firm_dec'])

            w['trat'] = mk_combo(['se le brinde un tratamiento interdisciplinario acorde a la problemática de adicción a sustancias estupefacientes que padece'], editable=True)
            add_pair("Tratamiento SPC:", w['trat'])

            w['punto'] = mk_line()
            add_pair("Punto que ordena tratamientos:", w['punto'])

            w['datos'] = mk_line()
            add_pair("Datos personales:", w['datos'])

            w['cumpl'] = mk_line()
            add_pair("Fecha cumplimiento total:", w['cumpl'])

            w['neuro'] = QCheckBox("Incluir Oficio al Neuropsiquiátrico")
            w['neuro'].stateChanged.connect(self.update)
            grid.addWidget(w['neuro'], row, 0, 1, 2)
            row += 1

            w['civ'] = QCheckBox("Incluir Oficio al Centro Integral de Varones")
            w['civ'].stateChanged.connect(self.update)
            grid.addWidget(w['civ'], row, 0, 1, 2)

            # Recuperar datos guardados en prev_data
            if i-1 < len(prev_data):
                old = prev_data[i-1]
                for k, val in old.items():
                    if k in w:
                        if isinstance(w[k], QLineEdit):
                            w[k].setText(val)
                        elif isinstance(w[k], QComboBox):
                            w[k].setCurrentText(val)
                        elif isinstance(w[k], QCheckBox):
                            w[k].setChecked(bool(val))

            if i-1 < len(self.data.imputados):
                dato = self.data.imputados[i-1]
                with QSignalBlocker(w['nombre']):
                    w['nombre'].setText(dato.get('nombre', ''))
                with QSignalBlocker(w['datos']):
                    w['datos'].setText(dato.get('datos', ''))
                with QSignalBlocker(w['defensa']):
                    w['defensa'].setText(dato.get('defensa', ''))
                with QSignalBlocker(w['delitos']):
                    w['delitos'].setText(dato.get('delitos', ''))
                with QSignalBlocker(w['condena']):    
                    w['condena'].setText(dato.get('condena', ''))
                # Campos opcionales:
                w['dni'].setText(dato.get('dni', ''))
                w['estable'].setCurrentText(dato.get('estable', ''))
                with QSignalBlocker(w['victimas']):
                    w['victimas'].setText(dato.get('victimas', ''))
                w['decreto'].setText(dato.get('decreto', ''))
                w['firm_dec'].setText(dato.get('firm_dec', ''))
                w['trat'].setCurrentText(dato.get('trat', ''))
                w['punto'].setText(dato.get('punto', ''))
                with QSignalBlocker(w['detenc']):
                    w['detenc'].setText(dato.get('detenc', ''))
                with QSignalBlocker(w['fechas']):
                    w['fechas'].setText(dato.get('fechas', ''))
                w['cumpl'].setText(dato.get('cumpl', ''))
                w['hechos_n'].setCurrentText(dato.get('hechos_n', ''))
                w['tipo'].setCurrentText(dato.get('tipo', ''))
                w['neuro'].setChecked(bool(dato.get('neuro', False)))
                w['civ'].setChecked(bool(dato.get('civ', False)))

            self.tabs_imp.addTab(tab, f"Imputado {i}")
            self.imputados_widgets.append(w)
            print(f"[rebuild_imputados] pestaña {i}: claves creadas →", list(w.keys()))
        # Actualizar selector y disparar actualización general
        self.selector_imp.clear()
        self.selector_imp.addItems([f"Imputado {i}" for i in range(1, n+1)])
        self.imp_index = 0
        self.selector_imp.setCurrentIndex(0)
        self._refresh_imp_names_in_selector()
        self._building = False

    def rebuild_hechos(self):
        if getattr(self, "_building_hechos", False):
            return
        self._building_hechos = True

        prev_data = []
        for w in getattr(self, "hechos_widgets", []):
            dato = {}
            for k, widget in w.items():
                if isinstance(widget, QLineEdit):
                    dato[k] = widget.text()
                elif isinstance(widget, QRadioButton):
                    dato[k] = widget.isChecked()
            prev_data.append(dato)

        self.tabs_hechos.clear()
        self.hechos_widgets = []
        n = self.spin_hechos.value()

        for i in range(1, n + 1):
            tab = QWidget()
            grid = QGridLayout(tab)
            row = 0

            def add_pair(text, widget):
                nonlocal row
                grid.addWidget(QLabel(text), row, 0)
                grid.addWidget(widget, row, 1)
                row += 1

            le_desc = QLineEdit(); le_desc.setReadOnly(True)
            btn_desc = QPushButton("Redactar el hecho")
            btn_desc.clicked.connect(lambda _=False, idx=i-1: self.abrir_ventana_hecho_desc(idx))
            grid.addWidget(QLabel(f"Descripción suceso #{i}"), row, 0)
            grid.addWidget(btn_desc, row, 1); row += 1

            le_aclar = QLineEdit(); add_pair("Aclaraciones:", le_aclar)
            le_ofi = QLineEdit(); add_pair("Oficina que elevó:", le_ofi)
            rb_j = QRadioButton("Juzgado"); rb_f = QRadioButton("Fiscalía"); rb_j.setChecked(True)
            grp = QButtonGroup(tab); grp.addButton(rb_j); grp.addButton(rb_f)
            grid.addWidget(rb_j, row, 0); grid.addWidget(rb_f, row, 1); row += 1
            le_auto = QLineEdit(); add_pair("N° del auto:", le_auto)
            le_fec = QLineEdit(); add_pair("Fecha de elevación:", le_fec)

            for w in [le_desc, le_aclar, le_ofi, le_auto, le_fec]:
                w.textChanged.connect(self.update_template)
            for w in [rb_j, rb_f]:
                w.toggled.connect(self.update_template)

            self.tabs_hechos.addTab(tab, f"Hecho {i}")
            self.hechos_widgets.append({
                "descripcion": le_desc,
                "aclaraciones": le_aclar,
                "oficina": le_ofi,
                "rb_j": rb_j,
                "rb_f": rb_f,
                "num_auto": le_auto,
                "fecha_elev": le_fec,
            })

            if i-1 < len(prev_data):
                old = prev_data[i-1]
                le_desc.setText(old.get("descripcion", ""))
                le_aclar.setText(old.get("aclaraciones", ""))
                le_ofi.setText(old.get("oficina", ""))
                rb_j.setChecked(old.get("rb_j", True))
                rb_f.setChecked(old.get("rb_f", False))
                le_auto.setText(old.get("num_auto", ""))
                le_fec.setText(old.get("fecha_elev", ""))

            if i-1 < len(self.data.hechos):
                dato = self.data.hechos[i-1]
                le_desc.setProperty("html", dato.get("descripcion", ""))
                doc = QTextDocument(); doc.setHtml(dato.get("descripcion", ""))
                le_desc.setText(doc.toPlainText()[:200])
                le_aclar.setText(dato.get("aclaraciones", ""))
                le_ofi.setText(dato.get("oficina", ""))
                rb_j.setChecked(dato.get("juzgado", True))
                rb_f.setChecked(not dato.get("juzgado", True))
                le_auto.setText(dato.get("num_auto", ""))
                le_fec.setText(dato.get("fecha_elev", ""))

        self._building_hechos = False
        self.update_template()

    def abrir_ventana_resuelvo(self):
        # recupero el HTML completo o, si no existe, la representación
        html_actual = self.entry_resuelvo.property("html") or ""
        self._rich_text_dialog(
            "Editar texto de Resuelvo",
            html_actual,
            self._guardar_resuelvo_html
        )

    def _guardar_resuelvo_html(self, html_limpio: str) -> None:
        clean = html_limpio.strip()
        # 1) guardo el HTML completo
        self.entry_resuelvo.setProperty("html", clean)
        self.entry_resuelvo.setHtml(clean)
        # 2) texto plano sin límite
        doc = QTextDocument()
        doc.setHtml(clean)
        preview = doc.toPlainText().replace("\n", " ")
        # 3) actualizo tu modelo con HTML y texto completo
        self.data.resuelvo_html = clean
        self.data.resuelvo      = preview
        # 4) refresco la plantilla
        self.update_template()

    def _guardar_html_lineedit(self, qlineedit, html):
        clean = html.strip()
        qlineedit.setProperty("html", clean)
        doc = QTextDocument(); doc.setHtml(clean)
        qlineedit.setText(doc.toPlainText()[:200])
        self.update_template()

    def abrir_ventana_hecho_desc(self, idx: int):
        qle = self.hechos_widgets[idx]["descripcion"]
        html_actual = qle.property("html") or qle.text()
        self._rich_text_dialog(
            f"Editar descripción del suceso #{idx+1}",
            html_actual,
            lambda h: self._guardar_html_lineedit(qle, h),
        )

    def _toggle_bold(self, editor: QTextEdit):
        cursor = editor.textCursor()
        if not cursor.hasSelection():
            return
        fmt = QTextCharFormat()
        bold_now = cursor.charFormat().fontWeight() > QFont.Normal
        fmt.setFontWeight(QFont.Normal if bold_now else QFont.Bold)
        cursor.mergeCharFormat(fmt)


    def _rich_text_dialog(self, title: str, initial_html: str, on_accept):
        dlg = QDialog(self)
        dlg.setWindowTitle(title)
        dlg.resize(650, 420)

        lay_top = QVBoxLayout(dlg)

        # — Toolbar de negrita —
        from PySide6.QtWidgets import QHBoxLayout
        toolbar = QHBoxLayout()
        btn_bold = QPushButton("B")
        btn_bold.setCheckable(True)
        btn_bold.setFixedWidth(32)
        btn_bold.setStyleSheet("font-weight:bold;")
        toolbar.addWidget(btn_bold)
        toolbar.addStretch()
        lay_top.addLayout(toolbar)

        # — Editor —
        editor = QTextEdit()
        editor.setAcceptRichText(True)
        base_font = QFont("Times New Roman", 12)
        editor.setFont(base_font)
        editor.document().setDefaultFont(base_font)
        editor.setHtml(initial_html or "")
        lay_top.addWidget(editor)

        # negrita con botón
        btn_bold.clicked.connect(lambda: self._toggle_bold(editor))

        # negrita con Ctrl+B
        editor.addAction(
            QAction(self, shortcut="Ctrl+B", triggered=lambda: self._toggle_bold(editor))
        )

        # — OK / Cancel —
        btn_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        lay_top.addWidget(btn_box)

        def _on_ok():
            raw = editor.toHtml()
            clean = _sanitize_html(raw)
            clean = unescape(clean)
            on_accept(clean)
            dlg.accept()

        btn_box.accepted.connect(_on_ok)
        btn_box.rejected.connect(dlg.reject)

        dlg.exec()

    @staticmethod
    def html_a_plano(html: str, mantener_saltos: bool = True) -> str:
        if not html:
            return ""

        doc = QTextDocument(); doc.setHtml(html)
        texto = doc.toPlainText()

        # → equivale a &nbsp; y &nbsp; finos (202F)
        texto = texto.replace("\u00A0", " ").replace("\u202F", " ")

        if not mantener_saltos:
            texto = texto.replace("\n", " ")

        return texto.strip()

    def _refresh_imp_names_in_selector(self):
        """Muestra el nombre si está cargado (“Imputado 1 – Pérez”)."""
        for i,w in enumerate(self.imputados_widgets):
            nom = w['nombre'].text().strip()
            txt = f"Imputado {i+1}" + (f" – {nom}" if nom else "")
            self.selector_imp.setItemText(i, txt)
    def update_for_imp(self, idx: int):
        """Se llama cuando el usuario elige otro imputado."""
        self.imp_index = min(idx, len(self.imputados_widgets) - 1)
        self._plantilla_oficio_notificacion()
        self._plantilla_oficio_neuro()        # ← idem
        self._plantilla_oficio_civ()
        self._plantilla_oficio_libertad()
        self._plantilla_oficio_policia()
        self._plantilla_oficio_reincidencia()
        self._plantilla_oficio_computo()
        self._plantilla_oficio_spc()
        self._plantilla_oficio_comunicacion()
        self._plantilla_legajo()
        self._plantilla_puesta_disposicion()

    def _imp(self):
        """Devuelve el dict del imputado activo o {} si el índice está fuera de rango."""
        if 0 <= getattr(self, 'imp_index', 0) < len(self.imputados_widgets):
            return self.imputados_widgets[self.imp_index]
        return {}


    def copy_to_clipboard(self, te: QTextEdit) -> None:
        """
        Copia **texto plano**, **RTF** (Times 12 + estilos) y **HTML**.
        Word elegirá el HTML (con negritas/alineaciones), pero el RTF
        queda por si lo necesita otro programa.
        """
        from PySide6.QtCore    import QMimeData
        from PySide6.QtWidgets import QApplication
        from PySide6.QtGui     import QClipboard

        # ---------- 1) texto sin formato --------------------------------------
        plain_text = te.toPlainText().strip()


# ---------- 2) HTML limpio + CSS Times 12 ---------------------------------
        basic_html = te.toHtml()

        # 1) limpiar tamaños de fuente
        basic_html = re.sub(r'font-size\s*:[^;"]+;?', '', basic_html, flags=re.I)

        # 2) limpiar "text-align:left" y align="left"
        basic_html = re.sub(r'text-align\s*:\s*left\s*;?', '', basic_html, flags=re.I)
        basic_html = re.sub(r'align="left"\s*', '', basic_html, flags=re.I)

        # 3) a cada <p …> que NO traiga center/right le agregamos justify inline
        def _ensure_justify(m):
            tag = m.group(0)
            # si ya tiene center/right, lo dejamos
            if re.search(r'(text-align\s*:\s*(center|right))|(align="(center|right)")',
                        tag, flags=re.I):
                return tag
            # si ya tiene style="", insertar el justify antes del cierre
            if 'style="' in tag:
                return re.sub(r'style="([^"]*)"', lambda s:
                            f'style="{s.group(1)}text-align:justify;"', tag)
            # no tenía style → agregamos uno
            return tag[:-1] + ' style="text-align:justify;">'

        basic_html = re.sub(r'<p[^>]*>', _ensure_justify, basic_html)

        # 4) borrar style="" vacíos que hayan quedado
        basic_html = re.sub(r'style="\s*"', '', basic_html)

        # 5) armar el fragmento completo (¡sin regla CSS que fuerce justify!)
        html_full = (
            "<!DOCTYPE html><html><head><meta charset='UTF-8'>"
            "<style>"
            "body{font-family:'Times New Roman',serif;"
            "font-size:12pt;line-height:1.0;margin:0;}"
            "</style></head><body><!--StartFragment-->"
            + basic_html +
            "<!--EndFragment--></body></html>"
        )

        # ---------- 3) RTF (Times 12, alineación + b/i/u) ----------------------
        rtf_paragraphs = []
        for para_html in re.findall(r'<p[^>]*>.*?</p>', basic_html, flags=re.S|re.I):
            rtf_paragraphs.append(_html_to_rtf_fragment(para_html))

        rtf_content = (
            r"{\rtf1\ansi\deff0"
            r"{\fonttbl{\f0 Times New Roman;}}"
            r"\fs24 " + ''.join(rtf_paragraphs) + "}"
        )

        # ---------- 4) al portapapeles (HTML lo dejamos el último) ------------
        mime = QMimeData()
        mime.setText(plain_text)
        mime.setData("text/rtf", rtf_content.encode("utf-8"))
        mime.setHtml(html_full)
        QApplication.clipboard().setMimeData(mime, QClipboard.Clipboard)


    def update(self):
        if getattr(self, "_building", False):
            return            # todavía estamos construyendo pestañas
        # ------------ plantillas ------------
        self._plantilla_pedido()
        self._plantilla_oficio_oga()
        self._plantilla_decreto_audiencia()
        self._plantilla_oficio_notificacion()
        self._plantilla_acta_renuncia()
        self._plantilla_constancia_grabacion()
        self._plantilla_certificado_victimas()
        self._plantilla_oficio_neuro()
        self._plantilla_oficio_civ()
        self._plantilla_oficio_libertad()
        self._plantilla_oficio_policia()
        self._plantilla_oficio_reincidencia()
        self._plantilla_oficio_computo()
        self._plantilla_oficio_spc()
        self._plantilla_oficio_comunicacion()
        self._plantilla_legajo()
        self._plantilla_puesta_disposicion()

        # ------------ demo para pestañas sin implementar ------------
        demo = "(Plantilla no implementada todavía)"
        for k, te in self.text_edits.items():
            if k not in (
                "Pedido de audiencia", "Oficio OGA", "Decreto audiencia",
                "Oficio notificación", "Acta renuncia", "Constancia grabación",
                "Certificado víctimas", "Oficio Neuro", "Oficio CIV",
                "Oficio libertad", "Oficio Policía", "Oficio Reincidencia",
                "Oficio cómputo", "Oficio SPC", "Oficio comunicación",
                "Legajo", "Puesta a disposición"
            ):
                te.setPlainText(demo)

        self.data.from_main(self)

    def update_template(self):
        self.data.from_main(self)
        self.update()
        # sólo reflejo el modelo en la UI si ya construí las pestañas
        if hasattr(self, 'tabs_imp') and self.tabs_imp is not None:
            self.data.apply_to_main(self)

    def _plantilla_pedido(self):
        fecha=fecha_letras(datetime.now())
        texto=(f"Córdoba, {fecha}.\n"
               "Atento al requerimiento de audiencia oral de juicio abreviado inicial, "
               "admítase la solicitud y requiérase vía e‑oficio a la Oficina de Gestión de Audiencias (OGA) "
               "que fije día y hora de realización de la audiencia presencial y asigne la sala para su desarrollo "
               "(art. 336 del CPP y Anexo II del AR n.º 1747 Serie “A” de fecha 1/4/2022).")
        te=self.text_edits["Pedido de audiencia"]; te.clear()
        self._insert_paragraph(te,texto,Qt.AlignJustify)

    def _plantilla_oficio_oga(self):
        te=self.text_edits["Oficio OGA"]; te.clear()
        cur=te.textCursor()

        blk_right=QTextBlockFormat(); blk_right.setAlignment(Qt.AlignRight)
        blk_left =QTextBlockFormat(); blk_left .setAlignment(Qt.AlignJustify)
        blk_just =QTextBlockFormat(); blk_just.setAlignment(Qt.AlignJustify)
        blk_center=QTextBlockFormat(); blk_center.setAlignment(Qt.AlignCenter)

        fmt=QTextCharFormat(); fmt.setFontFamily("Times New Roman"); fmt.setFontPointSize(12)
        fmt_bold=QTextCharFormat(fmt); fmt_bold.setFontWeight(QFont.Bold) # Bold

        def insert(text: str, blk: QTextBlockFormat, cfmt: QTextCharFormat = fmt):
            cur.insertBlock(blk)
            cur.setCharFormat(cfmt)         # 1º formato …
            cur.insertText(text) 

        hoy=datetime.now(); fecha=f"Córdoba, {hoy.day} de {_MESES[hoy.month]} de {hoy.year}."
        insert(fecha, blk_right)

        cur.insertBlock(blk_left)

        insert("Sr. Director de", blk_left, fmt_bold)
        insert("OGA Penal",        blk_left, fmt_bold)
        insert("S ____________/______________D", blk_left, fmt_bold)

        cur.insertBlock(blk_left)

        car=self.entry_caratula.text()
        camara=self.combo_articulo.currentText().split()[0].lower()  # “esta” / “este”
        articulo = "esta" if camara=="cámara" else "este"
        sec=self.entry_secretaria.text()

        cuerpo=(f"En los autos caratulados {car}, que se tramitan en {articulo} "
                f"{self.entry_tribunal.currentText()}, secretaría a cargo de {sec}, "
                "se ha resuelto librar a Ud. el presente oficio a fin de solicitar fecha y "
                "hora de audiencia de juicio abreviado inicial, conforme la información que se "
                "suministra por archivo adjunto.")
        insert(cuerpo, blk_just, fmt)
        cur.insertBlock(blk_left)
        self._bold_occurrences(te, [car])
        insert("Saludo a Ud. atentamente.", blk_center, fmt_bold)
    
    def _plantilla_decreto_audiencia(self):
        te = self.text_edits["Decreto audiencia"]
        te.clear()
        cur = te.textCursor()

        blk_left   = QTextBlockFormat(); blk_left.setAlignment(Qt.AlignJustify)
        blk_just   = QTextBlockFormat(); blk_just.setAlignment(Qt.AlignJustify)

        fmt_norm   = QTextCharFormat()
        fmt_norm.setFontFamily("Times New Roman"); fmt_norm.setFontPointSize(12)

        fmt_bu     = QTextCharFormat(fmt_norm)           # bold & underline
        fmt_bu.setFontWeight(QFont.Bold)
        fmt_bu.setFontUnderline(True)

        fmt_it = QTextCharFormat(fmt_norm)          # ***nuevo*** – cursiva
        fmt_it.setFontItalic(True)

        cur.insertBlock(blk_left); cur.setCharFormat(fmt_norm)
        cur.insertText(f"Córdoba, {fecha_letras(datetime.now())}.")

        cur.insertBlock(blk_left)   # línea en blanco

        nombres = [w['nombre'].text().strip() for w in self.imputados_widgets
                if w['nombre'].text().strip()]
        def lista(l):
            if not l:                return ""
            if len(l)==1:            return l[0]
            if len(l)==2:            return f"{l[0]} y {l[1]}"
            return ", ".join(l[:-1])+f" y {l[-1]}"

        if   len(nombres)==0: sit = "del imputado"
        elif len(nombres)==1: sit = f"del imputado {nombres[0]}"
        else:                 sit = f"de los imputados {lista(nombres)}"

        texto_ini = ("Atento a lo informado por la Oficina de Gestión de Audiencias (OGA) "
                    "mediante oficio electrónico, notifíquese a las partes que se ha fijado "
                    "audiencia a los fines de resolver la situación procesal "
                    f"{sit} para el ")

        cur.insertBlock(blk_just); cur.setCharFormat(fmt_norm)
        cur.insertText(texto_ini)

        fecha_aud = self.entry_fecha.text()
        hora_aud  = self.combo_hora.currentText()
        sala_aud  = self.combo_sala.currentText()

        cur.setCharFormat(fmt_bu)
        cur.insertText(f"día {fecha_aud} a las {hora_aud} h en la {sala_aud} de Tribunales II")

        cur.setCharFormat(fmt_norm)
        cur.insertText(" (art. 336 del CPP).")

    def _plantilla_oficio_notificacion(self):
        te = self.text_edits["Oficio notificación"]
        te.clear()
        cur = te.textCursor()

        blk_right  = QTextBlockFormat(); blk_right.setAlignment(Qt.AlignRight)
        blk_left   = QTextBlockFormat(); blk_left .setAlignment(Qt.AlignJustify)
        blk_just   = QTextBlockFormat(); blk_just .setAlignment(Qt.AlignJustify)
        blk_center = QTextBlockFormat(); blk_center.setAlignment(Qt.AlignCenter)

        fmt_norm = QTextCharFormat()
        fmt_norm.setFontFamily("Times New Roman"); fmt_norm.setFontPointSize(12)

        fmt_bold = QTextCharFormat(fmt_norm); fmt_bold.setFontWeight(QFont.Bold)
        fmt_bu   = QTextCharFormat(fmt_bold); fmt_bu.setFontUnderline(True)

        def insert(text, blk, fmt=fmt_norm):
            cur.insertBlock(blk); cur.setCharFormat(fmt); cur.insertText(text)

        car  = self.entry_caratula.text()
        art  = "esta" if self.combo_articulo.currentText().startswith("Cámara") else "este"
        trib = self.entry_tribunal.currentText()
        sec  = self.entry_secretaria.text()
        func = self.entry_funcionario.text()

        imp = self._imp()
        if not imp:                 # ← puede ocurrir en el milisegundo inicial
            return                  #   (todavía no existen los tabs)

        nombre_w  = imp.get('nombre')
        dni_w     = imp.get('dni')
        estable_w = imp.get('estable')

        nombre  = nombre_w.text()              if isinstance(nombre_w,  QLineEdit) else ""
        dni     = dni_w.text()                 if isinstance(dni_w,     QLineEdit) else ""
        estable = estable_w.currentText()      if isinstance(estable_w, QComboBox) else ""

        map_est = {
            "CC1 (Bouwer)": "Complejo Carcelario n.° 1 (Bouwer)",
            "EP9 (UCA)"   : "Establecimiento Penitenciario n.° 9 (UCA)",
            "EP3 (para mujeres)": "Establecimiento Penitenciario n.° 3 (para mujeres)",
            "CC2 (Cruz del Eje)": "Complejo Carcelario n.° 2 (Cruz del Eje)",
            "EP4 (Monte Cristo)": "Establecimiento Penitenciario n.° 4 (Colonia Abierta Monte Cristo)",
            "EP5 (Villa María)": "Establecimiento Penitenciario n.° 5 (Villa María)",
            "EP6 (Río Cuarto)":  "Establecimiento Penitenciario n.° 6 (Río Cuarto)",
            "EP7 (San Francisco)":"Establecimiento Penitenciario n.° 7 (San Francisco)",
            "EP8 (Villa Dolores)": "Establecimiento Penitenciario n.° 8 (Villa Dolores)",
        }
        estab = map_est.get(estable, estable)

        hoy = datetime.now()
        fecha_num = f"Córdoba, {hoy.day} de {_MESES[hoy.month]} de {hoy.year}."

        insert(fecha_num, blk_right)

        cur.insertBlock(blk_left)          # línea en blanco

        for linea in (
            "Sra. Jefa del Servicio Penitenciario",
            "de la Provincia de Córdoba",
            "S ____________/______________D",
        ):
            insert(linea, blk_left, fmt_bold)

        cur.insertBlock(blk_left)          # línea en blanco

        texto = (
            f"En los autos caratulados {car}, que se tramitan en {art} {trib} "
            f"se ha resuelto enviar el presente oficio a fin de solicitarle quiera tener a bien "
            "notificar la siguiente cédula al imputado "
            f"{nombre}, DNI n.° {dni}, cuya constancia de diligenciamiento deberá ser remitida "
            "a esta dependencia judicial:\n"
        )
        # parte con negritas (carátula y nombre)
        insert(texto, blk_just, fmt_norm)
        # resaltar carátula y penado
        self._bold_occurrences(te, [car, nombre])

        insert("CÉDULA DE NOTIFICACIÓN", blk_center, fmt_bu)

        for linea in (
            f"TRIBUNAL: {trib}.",
            f"SECRETARÍA: {sec}.",
            f"SEÑOR/A: {nombre}.",
            f"DOMICILIO: {estab}.",
        ):
            insert(linea, blk_left, fmt_norm)

        cur.insertBlock(blk_left)          # línea en blanco
        # ---- 6) párrafo con resolución (traemos texto del decreto) -------------
        res = self.text_edits["Decreto audiencia"].toPlainText().replace("\n", " ")

        texto_res = (f"Se le hace saber a Ud. que en los autos caratulados {car}, "
                    f"que se tramitan en {art} {trib} se ha dictado la siguiente resolución: “")

        cur.insertBlock(blk_just)      # ← 1 solo bloque para todo el párrafo
        cur.setCharFormat(fmt_norm)
        cur.insertText(texto_res)
        self._bold_occurrences(te, [car])          # carátula en negrita opcional
        cur.insertText(res.strip())                # decreto sin salto extra
        cur.insertText(f"” Fdo.: {func}.")         # cierra comilla + firma

        cur.insertBlock(blk_left)                  # línea en blanco real
        self._bu_between(te, "día ", "Tribunales II")  # negrita+subrayado al tramo
        insert("QUEDA UD. DEBIDAMENTE NOTIFICADO.", blk_center, fmt_bold)
        insert("Sin otro particular, saludo a Ud. atte.", blk_center, fmt_bold)

    def _bold_occurrences(self, te: QTextEdit, words: list[str]):
        """Pone en negrita cada aparición literal de las palabras dadas."""
        doc = te.document()

        # formato “Times 12 pt” en negrita
        fmt_b = QTextCharFormat()
        fmt_b.setFontFamily("Times New Roman")
        fmt_b.setFontPointSize(12)
        fmt_b.setFontWeight(QFont.Bold)

        for w in words:
            cursor = QTextCursor(doc)       # empieza al principio
            while True:
                cursor = doc.find(w, cursor)
                if cursor.isNull():
                    break                   # ya no hay más coincidencias
                cursor.mergeCharFormat(fmt_b)
    def _bu_between(self, te: QTextEdit, start_word: str, end_word: str):
        """Pone Bold+Underline desde la 1.ª aparición de start_word
        hasta la 1.ª aparición de end_word (inclusive)."""
        doc = te.document()
        cur_ini = doc.find(start_word)
        cur_end = doc.find(end_word, cur_ini)
        if cur_ini.isNull() or cur_end.isNull():
            return
        cur_ini.setPosition(cur_ini.selectionEnd())   # salto al final de start
        cur_end.setPosition(cur_end.selectionEnd())   # fin incluido
        cursor = QTextCursor(doc)
        cursor.setPosition(cur_ini.position())
        cursor.setPosition(cur_end.position(),
                        QTextCursor.KeepAnchor)

        fmt = QTextCharFormat()
        fmt.setFontFamily("Times New Roman")
        fmt.setFontPointSize(12)
        fmt.setFontWeight(QFont.Bold)
        fmt.setFontUnderline(True)
        cursor.mergeCharFormat(fmt)

    def _plantilla_acta_renuncia(self):
        te = self.text_edits["Acta renuncia"]
        te.clear()
        cur = te.textCursor()

        blk_just = QTextBlockFormat(); blk_just.setAlignment(Qt.AlignJustify)
        fmt_norm = QTextCharFormat()
        fmt_norm.setFontFamily("Times New Roman"); fmt_norm.setFontPointSize(12)

        def insert(text, blk=blk_just, fmt=fmt_norm):
            cur.insertBlock(blk); cur.setCharFormat(fmt); cur.insertText(text)

        if self.combo_renuncia.currentText() != "Sí":
            insert("No hubo renuncia a los plazos para interponer recurso de casación.")
            return                                         # ← nada más que hacer

        hora_aud = self.combo_hora.currentText()
        try:
            hora_ren = (datetime.strptime(hora_aud, "%H:%M") +
                        timedelta(hours=1)).strftime("%H:%M")
        except ValueError:
            hora_ren = "Hora inválida" if hora_aud else "Hora no especificada"

        hoy = datetime.now()
        fecha_letras = f"{hoy.day} de {_MESES[hoy.month]} de {hoy.year}"

        car = self.entry_caratula.text()
        fiscal = self.entry_fiscal.text() or "Sin datos de fiscal"

        imps  = self.imputados_widgets
        nombres = [w['nombre'].text().strip() for w in imps if w['nombre'].text().strip()]
        defensas = list({w['defensa'].text().strip() for w in imps if w['defensa'].text().strip()})

        imp_txt = ("el imputado" if len(nombres)==1 else "los imputados")
        nombres_txt = "Sin datos del penado" if not nombres else \
                    nombres[0] if len(nombres)==1 else ", ".join(nombres[:-1]) + f" y {nombres[-1]}"
        def_txt = "Sin datos de defensa" if not defensas else \
                defensas[0] if len(defensas)==1 else ", ".join(defensas[:-1]) + f" y {defensas[-1]}"

        plantilla1 = (
            f"En la ciudad de Córdoba, el {fecha_letras}, siendo las {hora_ren} horas, "
            f"en los presentes autos caratulados {car}, luego de haberse impuesto los "
            "fundamentos y el veredicto del día de la fecha, "
            f"{fiscal}; {def_txt}; y {imp_txt} {nombres_txt} "
            "manifestaron su voluntad de renunciar al plazo para interponer el recurso "
            "establecido en los arts. 468 y 469 del CPP, conforme lo estipulado por el "
            "art. 474 del CPP."
        )
        plantilla2 = (
            "Con lo que dio por terminado el acto, el que previa lectura dada en alta voz "
            "y ratificación de su contenido, firman las partes, todo por "
            "ante mí, de lo que doy fe."
        )

        insert(plantilla1)
        insert("")                # línea en blanco
        insert(plantilla2)

    def _plantilla_constancia_grabacion(self):
        """Genera la constancia con el enlace a la grabación."""
        te = self.text_edits["Constancia grabación"]
        te.clear()
        cur = te.textCursor()

        blk = QTextBlockFormat(); blk.setAlignment(Qt.AlignJustify)
        fmt = QTextCharFormat()
        fmt.setFontFamily("Times New Roman"); fmt.setFontPointSize(12)

        def insert(text):
            cur.insertBlock(blk); cur.setCharFormat(fmt); cur.insertText(text)

        fecha_aud   = self.entry_fecha.text().strip()
        hoy         = datetime.now().strftime("%d/%m/%Y")

        nombres = [w['nombre'].text().strip()
                for w in self.imputados_widgets if w['nombre'].text().strip()]

        if len(nombres) == 0:
            nombres_txt = ""
        elif len(nombres) == 1:
            nombres_txt = nombres[0]
        else:
            nombres_txt = ", ".join(nombres[:-1]) + f" y {nombres[-1]}"

        plantilla = (
            "Por medio de la presente, adjunto el archivo PDF que contiene el enlace "
            "de la grabación de la audiencia de juicio abreviado inicial celebrada "
            f"con fecha {fecha_aud}, en la que se resolvió la situación procesal "
            f"de {nombres_txt}. Of., {hoy}."
        )
        insert(plantilla)

    def _plantilla_certificado_victimas(self):
        """Completa la pestaña “Certificado víctimas”."""
        te = self.text_edits["Certificado víctimas"]
        te.clear()
        cur = te.textCursor()

        blk = QTextBlockFormat(); blk.setAlignment(Qt.AlignJustify)
        fmt = QTextCharFormat()
        fmt.setFontFamily("Times New Roman"); fmt.setFontPointSize(12)

        def insert(text):
            cur.insertBlock(blk); cur.setCharFormat(fmt); cur.insertText(text)

        victimas = [w['victimas'].text().strip()
                    for w in self.imputados_widgets
                    if w['victimas'].text().strip()]

        seen = set(); victimas = [v for v in victimas if not (v in seen or seen.add(v))]

        if len(victimas) == 0:
            vict_txt = ""
        elif len(victimas) == 1:
            vict_txt = victimas[0]
        else:
            vict_txt = ", ".join(victimas[:-1]) + f" y {victimas[-1]}"

        hoy = datetime.now().strftime("%d/%m/%Y")

        plantilla = (
            f"Certifico: que en el día de la fecha logré entablar comunicación con "
            f"{vict_txt}, damnificado/s en la presente causa, a fin de hacerle/s conocer "
            "la sentencia recaída en autos y conocer su voluntad respecto de las facultades "
            "que le/s confiere el art. 11 bis de la Ley 24.660. En dicha ocasión, "
            f"{vict_txt} manifestó/aron su voluntad de SER / NO SER anoticiado/s de los "
            f"eventuales beneficios de libertad. Of., {hoy}."
        )
        insert(plantilla)

    def _plantilla_oficio_neuro(self):
        te = self.text_edits["Oficio Neuro"]
        te.clear()
        cur = te.textCursor()
        
        imp = self._imp()
        if not imp or not imp['neuro'].isChecked() \
        or imp['tipo'].currentText() != "condicional":
            te.setPlainText("No aplica para penas efectivas o no seleccionado.")
            return

        blk_left   = QTextBlockFormat(); blk_left.setAlignment(Qt.AlignJustify)
        blk_right  = QTextBlockFormat(); blk_right.setAlignment(Qt.AlignRight)
        blk_center = QTextBlockFormat(); blk_center.setAlignment(Qt.AlignCenter)
        blk_just   = QTextBlockFormat(); blk_just.setAlignment(Qt.AlignJustify)

        fmt_norm = QTextCharFormat()
        fmt_norm.setFontFamily("Times New Roman")
        fmt_norm.setFontPointSize(12)

        fmt_b  = QTextCharFormat(fmt_norm)      # negrita
        fmt_b.setFontWeight(QFont.Bold)

        fmt_bu = QTextCharFormat(fmt_b)         # negrita + subrayado
        fmt_bu.setFontUnderline(True)

        fmt_it = QTextCharFormat(fmt_norm)      # cursiva
        fmt_it.setFontItalic(True)

        def ins(text, blk, fmt=fmt_norm):
            cur.insertBlock(blk)
            cur.setCharFormat(fmt)
            cur.insertText(text)

        hoy = datetime.now()
        ins(f"Córdoba, {hoy.day} de {_MESES[hoy.month]} de {hoy.year}.", blk_right)

        cur.insertBlock(blk_left)            # línea en blanco real

        for i, linea in enumerate((
            "AL SR. DIRECTOR",
            "DEL HOSPITAL",
            "NEUROPSIQUIÁTRICO",
            "PROVINCIAL",
            "(Rector León Morra 160)",       # ← solo esta va subrayada
            "S___________/___________D",
        )):
            fmt = fmt_bu if i == 4 else fmt_b
            ins(linea, blk_left, fmt)

        cur.insertBlock(blk_left)            # línea en blanco

        car  = self.entry_caratula.text()
        art  = ("esta" if self.combo_articulo.currentText()
                        .startswith("Cámara") else "este")
        trib = self.entry_tribunal.currentText()
        imp  = self._imp()
        nom  = imp.get("nombre").text()
        dni  = imp.get("dni").text()

        texto1 = (f"En los autos caratulados {car}, que se tramitan ante {art} "
                f"{trib}, se ha resuelto librar a Ud. el presente a fin de "
                f"solicitarle que arbitre los medios necesarios para que "
                f"{nom}, DNI n.° {dni}, reciba en la institución a su cargo un "
                "tratamiento interdisciplinario acorde con la problemática de "
                "adicción a sustancias estupefacientes que padece. Fundamenta el "
                "presente lo resuelto por veredicto dictado por este tribunal en "
                "el día de la fecha, en el que se impuso a la persona nombrada la "
                "pena bajo una serie de condiciones, entre ellas: ")
        ins(texto1, blk_just)
        self._bold_occurrences(te, [car, nom])
        cur.setCharFormat(fmt_it)
        cur.insertText(
            "“Iniciar un tratamiento interdisciplinario acorde a la problemática de "
            "adicción a sustancias estupefacientes que padece, debiendo presentar "
            "constancia del inicio del mismo en el término de 15 días ante el "
            "tribunal de ejecución interviniente”. "
        )
        cur.setCharFormat(fmt_norm)

        texto2 = ("En consecuencia, se solicita a Ud. la elaboración de un informe "
                "periódico dirigido a este tribunal, en el que comente la "
                "asistencia al tratamiento, así como su avance, y todo otro dato "
                "de interés.")
        cur.insertText(texto2)

        cur.insertBlock(blk_left) 
        # ---------- 4) Cierre --------------------------------------------
        ins("Saluda a Ud. atte.", blk_center, fmt_b)

    def _plantilla_oficio_civ(self):
        te = self.text_edits["Oficio CIV"]
        te.clear()
        cur = te.textCursor()

        imp = self._imp()
        if not imp or not imp['civ'].isChecked() \
        or imp['tipo'].currentText() != "condicional":
            te.setPlainText("No aplica para penas efectivas o no seleccionado.")
            return

        blk_r = QTextBlockFormat(); blk_r.setAlignment(Qt.AlignRight)
        blk_l = QTextBlockFormat(); blk_l.setAlignment(Qt.AlignJustify)
        blk_j = QTextBlockFormat(); blk_j.setAlignment(Qt.AlignJustify)
        blk_c = QTextBlockFormat(); blk_c.setAlignment(Qt.AlignCenter)

        fmt   = QTextCharFormat()
        fmt.setFontFamily("Times New Roman"); fmt.setFontPointSize(12)
        fmt_b = QTextCharFormat(fmt); fmt_b.setFontWeight(QFont.Bold)
        fmt_bu= QTextCharFormat(fmt_b); fmt_bu.setFontUnderline(True)

        def ins(t,b,f=fmt): cur.insertBlock(b); cur.setCharFormat(f); cur.insertText(t)

        hoy     = datetime.now()
        fecha_r = f"Córdoba, {hoy.day} de {_MESES[hoy.month]} de {hoy.year}."
        car     = self.entry_caratula.text()
        art     = "esta" if self.combo_articulo.currentText().startswith("Cámara") else "este"
        trib    = self.entry_tribunal.currentText()
        sec     = self.entry_secretaria.text()
        penado  = imp['nombre'].text()
        dni     = imp['dni'].text()

        ins(fecha_r, blk_r)

        cur.insertBlock(blk_l)

        for i,linea in enumerate((
            "AL SR. DIRECTOR DEL",
            "CENTRO INTEGRAL DE VARONES",   # ← subrayado
            "(Rondeau 258, Nueva Córdoba)",
            "S______________/______________D",
        )):
            ins(linea, blk_l, fmt_bu if i==1 else fmt_b)

        cur.insertBlock(blk_l)

        cuerpo = (
            f"En los presentes autos caratulados {car}, que se tramitan por ante {art} "
            f"{trib}, secretaría a cargo de {sec}, por disposición de S.S. se dirige a Ud. "
            f"el presente oficio a fin de solicitarle disponga los medios necesarios para "
            f"brindar asistencia psicoterapéutica a {penado}, DNI n.° {dni}, con relación a "
            "su problemática de violencia de género. Tal petición encuentra razón en que "
            "este Tribunal dispuso como condición de su libertad la realización de dicho "
            "tratamiento."
        )
        ins(cuerpo, blk_j)
        self._bold_occurrences(te, [car, penado])
        cur.insertBlock(blk_l)

        ins("Sin otro particular, saluda a Ud. atte.", blk_c, fmt_b)
 
    def _recopila_datos_imp(self) -> dict[str, str]:
        imp = self._imp() or {}

        def _txt(w):
            if isinstance(w, QLineEdit):
                html = w.property("html")
                if html:
                    doc = QTextDocument(); doc.setHtml(html)
                    return doc.toPlainText()
                return w.text()
            if isinstance(w, QComboBox):
                return w.currentText()
            return ""
        raw_html = self.entry_resuelvo.property("html") or ""

        # ① texto plano SIN saltos (\n → espacio)
        resuelvo_plano = self.html_a_plano(raw_html, mantener_saltos=False)
        return {
            'caratula'     : self.entry_caratula.text(),
            'tribunal'     : self.entry_tribunal.currentText(),
            'articulo'     : self.combo_articulo.currentText(),
            'sentencia'    : self.entry_sentencia.text(),
            'resuelvo'     : resuelvo_plano,

            'firmantes'    : self.entry_firmantes.text(),
            'penado'       : _txt(imp.get('nombre')),
            'dni'          : _txt(imp.get('dni')),
            'estable'      : _txt(imp.get('estable')),

            'decreto_computo' : _txt(imp.get('decreto')),  # ← alias para plantillas viejas
            # —— Delitos / hechos ——
            'delitos'      : _txt(imp.get('delitos')),   # nombre “nuevo”
            'hechos'       : "un hecho" if _txt(imp.get('hechos_n')) == "uno" else "hechos",

            'fechas_hechos': _txt(imp.get('fechas')),
            'condena'      : _txt(imp.get('condena')),
            'detencion'    : _txt(imp.get('detenc')),
            'computo_pena' : _txt(imp.get('decreto')),
            'defensa'      : _txt(imp.get('defensa')),
            'victimas'     : _txt(imp.get('victimas')),
            'datos'        : _txt(imp.get('datos')),    
            'renuncia'     : self.combo_renuncia.currentText(),
            'fecha_aud'    : self.entry_fecha.text(),
            'tipo_pena'    : _txt(imp.get('tipo')),
            'tratamiento_ordenado'      : _txt(imp.get('trat')),
            'parte_resuelvo_tratamiento': _txt(imp.get('punto')),
            'firmantes_decreto'         : _txt(imp.get('firm_dec')),
            'cumpl' : _txt(imp.get('cumpl')),
        }

    def _fecha_num(self) -> str:
        h = datetime.now()
        return f"Córdoba, {h.day:02d}/{h.month:02d}/{h.year}"

    def _insert_encabezado(self, cursor, lineas, blk_left, fmt_b, idx_sub=None):
        """Escribe las líneas de ‘encabezado_lineas’ (negrita; la posición
        idx_sub –si se da– irá negrita + subrayado)."""
        fmt_bu = QTextCharFormat(fmt_b)
        fmt_bu.setFontUnderline(True)
        for i, ln in enumerate(lineas):
            cursor.insertBlock(blk_left)
            cursor.setCharFormat(fmt_bu if i == idx_sub else fmt_b)
            cursor.insertText(ln)

    def _plantilla_oficio_libertad(self):
        te = self.text_edits["Oficio libertad"]
        te.clear()

        imp = self._imp()                                    # imputado activo
        tipo_w = imp.get("tipo") if imp else None
        if not imp or not isinstance(tipo_w, QComboBox):
            te.setPlainText("Aún no hay datos del imputado.")
            return
        if tipo_w.currentText() != "condicional":            # sólo para penas condicionales
            te.setPlainText("No aplica para penas efectivas.")
            return

        cur = te.textCursor()
        blk_r, blk_l, blk_c, blk_j = (QTextBlockFormat() for _ in range(4))
        blk_r.setAlignment(Qt.AlignRight)
        blk_l.setAlignment(Qt.AlignJustify)
        blk_c.setAlignment(Qt.AlignCenter)
        blk_j.setAlignment(Qt.AlignJustify)

        fmt     = QTextCharFormat(); fmt.setFontFamily("Times New Roman"); fmt.setFontPointSize(12)
        fmt_b   = QTextCharFormat(fmt);  fmt_b.setFontWeight(QFont.Bold)
        fmt_it  = QTextCharFormat(fmt);  fmt_it.setFontItalic(True)

        def ins(text, blk, f=fmt):
            cur.insertBlock(blk); cur.setCharFormat(f); cur.insertText(text)

        hoy = datetime.now()
        ins(f"Córdoba, {hoy.day} de {_MESES[hoy.month]} de {hoy.year}.", blk_r)

        self._insert_encabezado(
            cur,
            (
                "A LA SRA. JEFA DEL",
                "SERVICIO PENITENCIARIO",
                "DE LA PROVINCIA DE CÓRDOBA",
                "S___________/___________D",
            ),
            blk_l, fmt_b
        )
        cur.insertBlock(blk_l)          # línea en blanco real

        car   = self.entry_caratula.text()
        trib  = self.entry_tribunal.currentText()
        art   = "esta" if self.combo_articulo.currentText().startswith("Cámara") else "este"

        nom   = imp['nombre'].text()
        dni   = imp['dni'].text()
        con   = imp['condena'].text()

        texto = (
            f"En los autos caratulados {car} que se tramitan ante {art} {trib}, "
            "se ha dispuesto dirigir a Ud. el presente a fin de que disponga lo necesario "
            "para que se ponga inmediatamente en libertad, desde la Alcaidía de Tribunales II, "
            f"a {nom}, DNI n.° {dni}, en virtud de que por veredicto de este tribunal dictado "
            f"en el día de la fecha se le impuso la pena de {con}, disponiéndose su inmediata libertad. "
            "Deberá labrarse el acta respectiva y deberá requerírsele a la persona condenada que fije "
            "domicilio, el que deberá quedar consignado en el acta de libertad. La libertad se deberá "
            "disponer previa constatación de que el nombrado no se encuentre a disposición de otro tribunal."
        )
        ins(texto, blk_j)
        self._bold_occurrences(te, [car, nom])
        ins("", blk_l)                  # línea en blanco
        ins("Sin otro particular, saludo a Ud. atte.", blk_c, fmt_b)

    def _plantilla_oficio_policia(self):
        te = self.text_edits["Oficio Policía"]
        te.clear()
        cur = te.textCursor()
        blk_left  = QTextBlockFormat(); blk_left .setAlignment(Qt.AlignJustify)
        blk_right = QTextBlockFormat(); blk_right.setAlignment(Qt.AlignRight)
        blk_center= QTextBlockFormat(); blk_center.setAlignment(Qt.AlignCenter)
        blk_just  = QTextBlockFormat(); blk_just .setAlignment(Qt.AlignJustify)
        imp = self._imp()                       # ← ① recupero el imputado activo
        tipo_widget = imp.get("tipo") if imp else None  

        fmt = QTextCharFormat(); fmt.setFontFamily("Times New Roman"); fmt.setFontPointSize(12)
        fmt_b = QTextCharFormat(fmt); fmt_b.setFontWeight(QFont.Bold)
        # ---------- FORMATOS DE CARÁCTER -------------------------------
        fmt_norm = QTextCharFormat()
        fmt_norm.setFontFamily("Times New Roman")
        fmt_norm.setFontPointSize(12)

        fmt_b  = QTextCharFormat(fmt_norm)      # negrita
        fmt_b.setFontWeight(QFont.Bold)

        fmt_bu = QTextCharFormat(fmt_b)         # negrita + subrayado
        fmt_bu.setFontUnderline(True)

        fmt_it = QTextCharFormat(fmt_norm)      # cursiva
        fmt_it.setFontItalic(True)
        # helper local
        def ins(text, blk, fmt=fmt_norm):
            cur.insertBlock(blk)
            cur.setCharFormat(fmt)
            cur.insertText(text)

        hoy = datetime.now()
        ins(f"Córdoba, {hoy.day} de {_MESES[hoy.month]} de {hoy.year}.", blk_right)

        self._insert_encabezado(
            cur,
            (
                "AL SEÑOR DIRECTOR DE LA",
                "DIVISIÓN DOCUMENTACIÓN PERSONAL",
                "POLICÍA DE LA PROVINCIA DE CÓRDOBA",
                "S______________/______________D",
            ),
            blk_left, fmt_b
        )
        cur.insertBlock(blk_left)        # línea en blanco

        d = self._recopila_datos_imp()
        camara=self.combo_articulo.currentText().split()[0].lower()  # “esta” / “este”
        articulo = "esta" if camara=="cámara" else "este"

        texto1 = (
            f"En los autos caratulados {d['caratula']}, tramitados por ante "
            f"{articulo} {d['tribunal']}, se ha resuelto librar a Ud. el presente "
            f"a fin de que proceda a la anotación correspondiente de la sentencia n.° {d['sentencia']} en "
            f"los presentes autos, con relación a {d['penado']}, DNI n.° {d['dni']}, por {d['hechos']} "
            f"de fecha {d['fechas_hechos']}, que RESUELVE: "
        )

        cur.insertBlock(blk_just); cur.setCharFormat(fmt); cur.insertText(texto1)
        cur.setCharFormat(fmt_it)
   # ------------------------------------------------------------------------
        cur.insertText(f"“{d['resuelvo']}”")

        cur.setCharFormat(fmt)
        cur.insertText(f". Fdo.: {d['firmantes']}.")

        self._bold_occurrences(
            te,
            [d['caratula'], "sentencia n.°", d['sentencia'],
            d['penado'], "RESUELVE:", "Fdo.:", d['firmantes']]
        )
        cur.insertBlock(blk_left) # línea en blanco real
        # ── 4) Saludo centrado en negrita ────────────────────────────────
        cur.insertBlock(blk_center); cur.setCharFormat(fmt_b)
        cur.insertText("Saludo a Ud. atte.")

    def _plantilla_oficio_reincidencia(self):
        """Genera el oficio para el Registro Nacional de Reincidencia según el nuevo modelo."""
        import re
        from datetime import datetime

        te = self.text_edits["Oficio Reincidencia"]
        te.clear()
        cur = te.textCursor()

        # — formatos de bloque y carácter —
        blk_center = QTextBlockFormat(); blk_center.setAlignment(Qt.AlignCenter)
        blk_left   = QTextBlockFormat(); blk_left .setAlignment(Qt.AlignJustify)

        fmt_norm = QTextCharFormat()
        fmt_norm.setFontFamily("Times New Roman"); fmt_norm.setFontPointSize(12)

        fmt_bold  = QTextCharFormat(fmt_norm); fmt_bold.setFontWeight(QFont.Bold)
        fmt_bu    = QTextCharFormat(fmt_bold);  fmt_bu.setFontUnderline(True)
        fmt_italic= QTextCharFormat(fmt_norm);  fmt_italic.setFontItalic(True)

        def ins(text: str, blk: QTextBlockFormat, fmt: QTextCharFormat):
            cur.insertBlock(blk); cur.setCharFormat(fmt); cur.insertText(text)

        def ins_pair(title: str, value: str):
            cur.insertBlock(blk_left)
            cur.setCharFormat(fmt_bold)
            cur.insertText(f"{title}: ")
            cur.setCharFormat(fmt_norm)
            cur.insertText(value)

        # — encabezado centrado —
        ins("MINISTERIO DE JUSTICIA, SEGURIDAD Y DERECHOS HUMANOS", blk_center, fmt_norm)
        ins("REGISTRO NACIONAL DE REINCIDENCIA",                  blk_center, fmt_norm)
        cur.insertBlock(blk_center)  # espacio antes del título
        ins("TESTIMONIO DE SENTENCIA CONDENATORIA",               blk_center, fmt_bu)
        cur.insertBlock(blk_center)  # espacio antes del título
        # — datos del expediente —
        ins_pair("Sentencia", f"N° {self.entry_sentencia.text().strip()}")
        ins_pair("Tribunal interviniente", 
                 f"{self.entry_tribunal.currentText()}, Secretaría n.° {self.entry_secretaria.text().strip()}")
        ins_pair("Otros juzgados o tribunales intervinientes en la causa con anterioridad", "")
        ins_pair("Expediente",           self.entry_caratula.text().strip())

        imp = self._imp()
        ins_pair("Datos personales",            imp.get("datos").text().strip()    if imp.get("datos")    else "")
        ins_pair("Fecha de comisión del delito", imp.get("fechas").text().strip()   if imp.get("fechas")   else "")
        ins_pair("Localidad de comisión del delito", "Córdoba")
        ins_pair("Damnificado",                  imp.get("victimas").text().strip() if imp.get("victimas") else "")
        ins_pair("Descripción de la pena",       f"prisión de ejecución {imp.get('tipo').currentText()}" if imp.get("tipo") else "")

        # — pena con viñeta —
        cur.insertBlock(blk_left)
        cur.setCharFormat(fmt_bold)
        cur.insertText("•    Pena: ")
        cur.setCharFormat(fmt_norm)
        cur.insertText(imp.get("condena").text().strip() if imp.get("condena") else "")

        # — extraer sólo el punto con 'Declarar' —
        # 1) obtener el HTML original y aplanarlo sin saltos
        raw = self.entry_resuelvo.property("html") or ""
        full = self.html_a_plano(raw, mantener_saltos=False)
        pattern = r'\b([IVX]+|\d+)\.\s+([\s\S]*?)(?=(?:[IVX]+|\d+)\.\s+|$)'



        declarar = []
        for m in re.finditer(pattern, full, re.DOTALL|re.IGNORECASE):
            num, txt = m.group(1), m.group(2).strip()
            if re.search(r'\bdeclar', txt, re.IGNORECASE):
                declarar.append(f"{num}. {txt}")
        punto = ' '.join(declarar)

        # — TESTIMONIO: en cursiva, entre comillas y con puntos suspensivos —
        cur.insertBlock(blk_left)
        cur.setCharFormat(fmt_bu)
        cur.insertText("TESTIMONIO")
        # continuar en negrita sin subrayado para ":" y espacio
        cur.setCharFormat(fmt_bold)
        cur.insertText(": ")
        cur.setCharFormat(fmt_italic)
        cur.insertText(f'"(...) {punto} (...)"')
        # restaurar formato normal
        cur.setCharFormat(fmt_norm)

        # — siguientes campos —
        ins_pair("Fecha de cumplimiento total de la pena",
                 imp.get("cumpl").text().strip() if imp.get("cumpl") else "")

        if self.combo_renuncia.currentText() == "Sí":
            fecha_firme = self.entry_fecha.text().strip()
        else:
            fecha_firme = ""
        ins_pair("Fecha en que la sentencia quedó firme", fecha_firme)

        ins_pair("Fecha de envío del testimonio", datetime.now().strftime("%d/%m/%Y"))
        ins_pair("Organismo remitente", "Poder Judicial de la Provincia de Córdoba")

    def _plantilla_oficio_computo(self):
        te = self.text_edits["Oficio cómputo"]
        te.clear()

        imp = self._imp()                                    # imputado activo
        tipo_w = imp.get("tipo") if imp else None            # widget ‘tipo de pena’
        if not imp or not isinstance(tipo_w, QComboBox):
            te.setPlainText("Aún no hay datos del imputado.")
            return
        if tipo_w.currentText() != "efectiva":               # sólo para penas efectivas
            te.setPlainText("No aplica para penas condicionales.")
            return

        cur = te.textCursor()

        blk_r, blk_l, blk_c, blk_j = (QTextBlockFormat() for _ in range(4))
        blk_r.setAlignment(Qt.AlignRight)
        blk_l.setAlignment(Qt.AlignJustify)
        blk_c.setAlignment(Qt.AlignCenter)
        blk_j.setAlignment(Qt.AlignJustify)
        imp = self._imp()                       # ← ① recupero el imputado activo
        tipo_widget = imp.get("tipo") if imp else None 
        # ── datos que ya tenemos centralizados ─────────────────────────────
        d    = self._recopila_datos_imp()
        sec  = self.entry_secretaria.text()

        mapa_est = {
            "CC1 (Bouwer)": "Complejo Carcelario n.° 1 (Bouwer)",
            "EP9 (UCA)"   : "Establecimiento Penitenciario n.° 9 (UCA)",
            "EP3 (para mujeres)": "Establecimiento Penitenciario n.° 3 (para mujeres)",
            "CC2 (Cruz del Eje)" : "Complejo Carcelario n.° 2 (Cruz del Eje)",
            "EP4 (Colonia Abierta Monte Cristo)": "Establecimiento Penitenciario n.° 4 (Colonia Abierta Monte Cristo)",
            "EP5 (Villa María)": "Establecimiento Penitenciario n.° 5 (Villa María)",
            "EP6 (Río Cuarto)" : "Establecimiento Penitenciario n.° 6 (Río Cuarto)",
            "EP7 (San Francisco)": "Establecimiento Penitenciario n.° 7 (San Francisco)",
            "EP8 (Villa Dolores)": "Establecimiento Penitenciario n.° 8 (Villa Dolores)",
        }
        estab = mapa_est.get(self._imp()['estable'].currentText(),
                            self._imp()['estable'].currentText())

        fmt = QTextCharFormat(); fmt.setFontFamily("Times New Roman"); fmt.setFontPointSize(12)
        fmt_b = QTextCharFormat(fmt); fmt_b.setFontWeight(QFont.Bold)
        # ---------- FORMATOS DE CARÁCTER -------------------------------
        fmt_norm = QTextCharFormat()
        fmt_norm.setFontFamily("Times New Roman")
        fmt_norm.setFontPointSize(12)

        fmt_b  = QTextCharFormat(fmt_norm)      # negrita
        fmt_b.setFontWeight(QFont.Bold)

        fmt_bu = QTextCharFormat(fmt_b)         # negrita + subrayado
        fmt_bu.setFontUnderline(True)

        fmt_it = QTextCharFormat(fmt_norm)      # cursiva
        fmt_it.setFontItalic(True)
        # helper local
        def ins(text, blk, fmt=fmt_norm):
            cur.insertBlock(blk)
            cur.setCharFormat(fmt)
            cur.insertText(text)

        hoy = datetime.now()
        ins(f"Córdoba, {hoy.day} de {_MESES[hoy.month]} de {hoy.year}.", blk_r)

        self._insert_encabezado(
            cur,
            (
                "SRA. JEFA DEL SERVICIO",
                "PENITENCIARIO DE LA",
                "PROVINCIA DE CÓRDOBA",
                "S______________/______________D",
            ),
            blk_l, fmt_b
        )
        cur.insertBlock(blk_l)    # línea en blanco
        camara=self.combo_articulo.currentText().split()[0].lower()  # “esta” / “este”
        articulo = "esta" if camara=="cámara" else "este"
        # ── 3) cuerpo 1 ────────────────────────────────────────────────────
        txt1 = (
            f"En los autos caratulados {d['caratula']}, que se tramitan por ante "
            f"{articulo} {d['tribunal']}, se ha resuelto enviar el presente oficio "
            f"a fin de solicitarle quiera tener a bien notificar la siguiente cédula a "
            f"{d['penado']}, DNI n.° {d['dni']}, cuya constancia de diligenciamiento deberá "
            f"ser remitida a esta dependencia judicial:\n"
        )
        cur.insertBlock(blk_j); cur.setCharFormat(fmt); cur.insertText(txt1)

        cur.insertBlock(blk_l)
        cur.insertBlock(blk_c); cur.setCharFormat(fmt_bu)
        cur.insertText("CÉDULA DE NOTIFICACIÓN")
        cur.insertBlock(blk_l)

        datos_ced = (
            f"TRIBUNAL: {d['tribunal']}, Fructuoso Rivera n.° 720, Palacio de Tribunales II.",
            f"SECRETARÍA: {sec}.",
            f"SEÑOR: {d['penado']}.",
            f"DOMICILIO: {estab}.",
        )
        for ln in datos_ced:
            cur.insertBlock(blk_l); cur.setCharFormat(fmt); cur.insertText(ln)
        cur.insertBlock(blk_l)

        txt2 = (
            f"Se hace saber a Ud. que en los autos caratulados {d['caratula']}, que se "
            f"tramitan por ante {articulo} {d['tribunal']}, se ha dictado la siguiente resolución: "
        )
        cur.insertBlock(blk_j); cur.setCharFormat(fmt); cur.insertText(txt2)

        cur.setCharFormat(fmt_it)
        cur.insertText(f"“{d['decreto_computo']}”")
        cur.setCharFormat(fmt)
        cur.insertText(f". Fdo.: {d['firmantes_decreto']}.")

        cur.insertBlock(blk_r); cur.setCharFormat(fmt)
        cur.insertText(f"Of. {self._fecha_num()}.")
        cur.insertBlock(blk_c); cur.setCharFormat(fmt_b)
        cur.insertText("Saludo a Ud. atte.")

        self._bold_occurrences(
            te,
            [d['caratula'], d['penado'], "Fdo.:", d['firmantes_decreto']]
        )

    def _plantilla_oficio_spc(self):
        te = self.text_edits["Oficio SPC"]
        te.clear()

        imp = self._imp()                                    # imputado activo
        tipo_w = imp.get("tipo") if imp else None            # widget ‘tipo de pena’
        if not imp or not isinstance(tipo_w, QComboBox):
            te.setPlainText("Aún no hay datos del imputado.")
            return
        if tipo_w.currentText() != "efectiva":               # sólo para penas efectivas
            te.setPlainText("No aplica para penas condicionales.")
            return
        cur = te.textCursor()

        blk_r, blk_l, blk_c, blk_j = (QTextBlockFormat() for _ in range(4))
        blk_r.setAlignment(Qt.AlignRight);  blk_l.setAlignment(Qt.AlignJustify)
        blk_c.setAlignment(Qt.AlignCenter); blk_j.setAlignment(Qt.AlignJustify)

        fmt   = QTextCharFormat(); fmt.setFontFamily("Times New Roman"); fmt.setFontPointSize(12)
        fmt_b = QTextCharFormat(fmt); fmt_b.setFontWeight(QFont.Bold)
        fmt_it= QTextCharFormat(fmt); fmt_it.setFontItalic(True)
        # ── datos que ya tenemos centralizados ─────────────────────────────
        d    = self._recopila_datos_imp()
        sec  = self.entry_secretaria.text()
        # Helper local
        def ins(text, blk, f=fmt):
            cur.insertBlock(blk); cur.setCharFormat(f); cur.insertText(text)

        hoy = datetime.now()
        ins(f"Córdoba, {hoy.day} de {_MESES[hoy.month]} de {hoy.year}.", blk_r)

        self._insert_encabezado(
            cur,
            (
                "SR. DIRECTOR DEL",
                "ESTABLECIMIENTO PENITENCIARIO",
                "PBRO. LUCHESSE –BOWER–",
                "S__________________/__________________D",
            ),
            blk_l, fmt_b
        )
        cur.insertBlock(blk_l)        # línea en blanco

        articulo = "esta" if self.combo_articulo.currentText().startswith("Cámara") else "este"

        txt1 = (
            f"En los autos caratulados {d['caratula']}, que se tramitan por ante "
            f"{articulo} {d['tribunal']}, se ha dispuesto librar a Ud. el presente, "
            f"a fin de que cumplimente con lo resuelto por este tribunal en la sentencia n.° "
            f"{d['sentencia']}, con relación a {d['penado']}, DNI n.° {d['dni']}, a los efectos de que "
            f"arbitre los medios necesarios para que {d['tratamiento_ordenado']}."
        )
        ins(txt1, blk_j)

        ins("", blk_l)                         # salto real
        ins("Para mayor recaudo se transcribe la parte resolutiva que así lo dispone: ", blk_j)
        cur.setCharFormat(fmt_it)
        cur.insertText(f"“{d['parte_resuelvo_tratamiento']}”")
        cur.setCharFormat(fmt)
        cur.insertText(f". Fdo.: {d['firmantes']}.")

        cur.insertBlock(blk_l)                # línea en blanco real
        # ── 5) saludo ──────────────────────────────────────────────────────
        cur.insertBlock(blk_c); cur.setCharFormat(fmt_b)
        cur.insertText("Saludo a Ud. atte.")

        self._bold_occurrences(
            te,
            [d['caratula'], "sentencia n.°", d['sentencia'],
            d['penado'], "Fdo.:", d['firmantes']]
        )

    def _plantilla_oficio_comunicacion(self):
        te = self.text_edits["Oficio comunicación"]
        te.clear()

        imp = self._imp()                               # imputado activo
        tipo_w = imp.get("tipo") if imp else None       # widget “tipo de pena”
        if not imp or not isinstance(tipo_w, QComboBox):
            te.setPlainText("Aún no hay datos del imputado.")
            return
        if tipo_w.currentText() != "efectiva":          # sólo aplica a penas efectivas
            te.setPlainText("No es necesario en penas de ejecución condicional.")
            return

        cur = te.textCursor()
        blk_r, blk_l, blk_c, blk_j = (QTextBlockFormat() for _ in range(4))
        blk_r.setAlignment(Qt.AlignRight);  blk_l.setAlignment(Qt.AlignJustify)
        blk_c.setAlignment(Qt.AlignCenter); blk_j.setAlignment(Qt.AlignJustify)

        fmt   = QTextCharFormat(); fmt.setFontFamily("Times New Roman"); fmt.setFontPointSize(12)
        fmt_b = QTextCharFormat(fmt); fmt_b.setFontWeight(QFont.Bold)
        fmt_it= QTextCharFormat(fmt); fmt_it.setFontItalic(True)

        def ins(text, blk, f=fmt):
            cur.insertBlock(blk); cur.setCharFormat(f); cur.insertText(text)

        hoy = datetime.now()
        ins(f"Córdoba, {hoy.day} de {_MESES[hoy.month]} de {hoy.year}.", blk_r)

        self._insert_encabezado(
            cur,
            (
                "SRA. JEFA DEL SERVICIO",
                "PENITENCIARIO DE LA",
                "PROVINCIA DE CÓRDOBA",
                "S______________/______________D",
            ),
            blk_l, fmt_b
        )
        cur.insertBlock(blk_l)

        d = self._recopila_datos_imp()
        articulo = "esta" if self.combo_articulo.currentText().startswith("Cámara") else "este"

        txt1 = (
            f"En los autos caratulados {d['caratula']}, tramitados por ante "
            f"{articulo} {d['tribunal']}, se ha resuelto librar a Ud. el presente "
            f"a fin de informarle que el imputado {d['penado']}, DNI n.° {d['dni']}, "
            f"ha sido condenado a la pena de {d['condena']}. Ello en virtud de que se "
            f"ha llevado a cabo un juicio abreviado inicial y mediante sentencia n.° "
            f"{d['sentencia']}, se resolvió: "
        )
        ins(txt1, blk_j)

        cur.setCharFormat(fmt_it)
        cur.insertText(f"“{d['resuelvo']}”")
        cur.setCharFormat(fmt)
        cur.insertText(f". Fdo.: {d['firmantes']}.")
        self._bold_occurrences(
            te,
            [d['caratula'], d['penado'], "sentencia n.°", d['sentencia'],
            d['condena'], "Fdo.:", d['firmantes']]
        )

        cur.insertBlock(blk_j)
        add = "Asimismo, se hace saber que dicha sentencia quedó firme con fecha"
        if d['renuncia'] == "Sí":
            add += f" {d['fecha_aud']} por renuncia expresa de las partes a los plazos para interponer recurso de casación"
        add += ". A continuación, se transcribe el decreto que establece el cómputo definitivo de la pena impuesta: "
        ins(add, blk_j)

        cur.setCharFormat(fmt_it)
        cur.insertText(f"“{d['decreto_computo']}”")
        cur.setCharFormat(fmt)
        cur.insertText(f". Fdo.: {d['firmantes_decreto']}.")
        self._bold_occurrences(te, ["Fdo.:", d['firmantes_decreto']])
        cur.insertBlock(blk_l)                # línea en blanco real
        # ── 7) saludo ─────────────────────────────────────────────────────
        cur.insertBlock(blk_c); cur.setCharFormat(fmt_b)
        cur.insertText("Saludo a Ud. atte.")

    def _plantilla_legajo(self):
        te = self.text_edits["Legajo"]
        te.clear()

        d   = self._recopila_datos_imp()
        imp = self._imp()
        if not imp or not isinstance(imp.get("tipo"), QComboBox):
            te.setPlainText("Aún no hay datos del imputado.")
            return

        cur = te.textCursor()
        blk_l  = QTextBlockFormat(); blk_l.setAlignment(Qt.AlignJustify)
        blk_c  = QTextBlockFormat(); blk_c.setAlignment(Qt.AlignCenter)
        blk_j  = QTextBlockFormat(); blk_j.setAlignment(Qt.AlignJustify)

        fmt    = QTextCharFormat(); fmt.setFontFamily("Times New Roman"); fmt.setFontPointSize(12)
        fmt_b  = QTextCharFormat(fmt); fmt_b.setFontWeight(QFont.Bold)
        fmt_bu = QTextCharFormat(fmt_b); fmt_bu.setFontUnderline(True)
        fmt_u  = QTextCharFormat(fmt);  fmt_u.setFontUnderline(True)

        def ins(text, blk, f=fmt):
            cur.insertBlock(blk); cur.setCharFormat(f); cur.insertText(text)

        ins("LEGAJO DE REMISIÓN AL JUZGADO DE EJECUCIÓN PENAL", blk_c, fmt_bu)
        subt = "Pena privativa de la libertad de ejecución condicional" \
            if d['tipo_pena'] == "condicional" else "Pena privativa de la libertad"
        ins(subt, blk_c, fmt_b)

        campos = (
            ("Causa caratulada", d['caratula']),
            ("Tribunal",         d['tribunal']),
            ("Penado",           f"{d['penado']}, {d['datos']}"),
            ("Detención",        d['detencion']),
            ("Sentencia",        f"n.° {d['sentencia']}"),
            ("Delitos",          d['delitos']),
            ("Condena",          d['condena']),
            ("Cómputo de pena",   d.get('cumpl', "")),
            ("Defensa",          d['defensa']),
            ("Víctimas",         d['victimas']),
        )

        for titulo, contenido in campos:
            ins("", blk_l)                         # línea en blanco
            ins(titulo, blk_l, fmt_u)              # sub‑raya el título
            cur.insertText(f": {contenido}", fmt)  # contenido normal

        ins("", blk_l)
        ins("-La sentencia recaída en autos se encuentra firme y el cómputo de pena es definitivo.", blk_j)
        ins("-Se hace saber a Ud. que ya se han remitido los correspondientes oficios al Servicio Penitenciario en cumplimiento del art. 505 del Código Procesal Penal de la Provincia de Córdoba, a la Policía de la Provincia de Córdoba y al Registro Nacional de Reincidencia, comunicando la sentencia dictada en autos y el cómputo de pena.", blk_j)

    def _plantilla_puesta_disposicion(self):
        te = self.text_edits["Puesta a disposición"]
        te.clear()

        imp  = self._imp() or {}
        tipo = imp.get("tipo").currentText() if isinstance(imp.get("tipo"), QComboBox) else ""

        if tipo != "efectiva":
            te.setPlainText("No es necesario en penas de ejecución condicional.")
            return

        cur = te.textCursor()
        blk_r = QTextBlockFormat(); blk_r.setAlignment(Qt.AlignRight)
        blk_l = QTextBlockFormat(); blk_l.setAlignment(Qt.AlignJustify)
        blk_c = QTextBlockFormat(); blk_c.setAlignment(Qt.AlignCenter)
        blk_j = QTextBlockFormat(); blk_j.setAlignment(Qt.AlignJustify)

        fmt   = QTextCharFormat(); fmt.setFontFamily("Times New Roman"); fmt.setFontPointSize(12)
        fmt_b = QTextCharFormat(fmt); fmt_b.setFontWeight(QFont.Bold)
        fmt_bu= QTextCharFormat(fmt_b); fmt_bu.setFontUnderline(True)

        def ins(text, blk, f=fmt):
            cur.insertBlock(blk); cur.setCharFormat(f); cur.insertText(text)

        d        = self._recopila_datos_imp()
        articulo = "esta" if self.combo_articulo.currentText().startswith("Cámara") else "este"
        hoy      = datetime.now()
        fecha    = f"Córdoba, {hoy.day} de {_MESES[hoy.month]} de {hoy.year}."

        ins(fecha, blk_r)

        self._insert_encabezado(
            cur,
            (
                "SRA. JEFA DEL SERVICIO",
                "PENITENCIARIO DE LA",
                "PROVINCIA DE CÓRDOBA",
                "S______________/______________D",
            ),
            blk_l, fmt_b
        )
        cur.insertBlock(blk_l)   # línea en blanco

        texto = (
            f"En los autos caratulados {d['caratula']}, que se tramitan ante "
            f"{articulo} {d['tribunal']}, se le hace saber que el condenado "
            f"{d['penado']}, DNI n.° {d['dni']}, queda a exclusiva disposición del "
            "Juzgado de Ejecución Penal n.° ……, bajo las actuaciones del Cuerpo de "
            f"Ejecución de Pena Privativa de Libertad de {d['penado']} (SAC n.º ……), "
            "siempre que no se encuentre a disposición de otro tribunal."
        )
        ins(texto, blk_j)

        self._bold_occurrences(
            te,
            [d['caratula'], d['penado'],
            "Cuerpo de Ejecución", "de Pena Privativa de Libertad"]
        )

        cur.insertBlock(blk_l)
        ins("Sin otro particular, saludo a Ud. atte.", blk_c, fmt_b)

    def _insert_paragraph(self, te:QTextEdit, text:str, align:int):
        cur=te.textCursor()
        blk=QTextBlockFormat(); blk.setAlignment(align)
        fmt=QTextCharFormat(); fmt.setFontFamily("Times New Roman"); fmt.setFontPointSize(12)
        for p in text.split("\n"):
            cur.insertBlock(blk); cur.setCharFormat(fmt); cur.insertText(p)

    def generate_planilla_oga(self):
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # 1) Crear documento
        doc = Document()

        # 2) Cabecera principal (tabla 1 fila x 2 cols, celdas fusionadas)
        tbl_hdr = doc.add_table(rows=1, cols=2, style="Table Grid")
        hdr_cells = tbl_hdr.rows[0].cells
        hdr_cells[0].text = "SOLICITUD DE AUDIENCIA DE JUICIO ABREVIADO"
        hdr_cells[0].merge(hdr_cells[1])
        # centrar texto y poner negrita
        p = hdr_cells[0].paragraphs[0]
        run = p.runs[0]
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 3) Datos generales (tabla 4 filas x 2 cols)
        tbl_gen = doc.add_table(rows=4, cols=2, style="Table Grid")
        filas = [
            ("Fecha posible para fijación de audiencia", self.entry_fecha.text()),
            ("Hora",                                   self.combo_hora.currentText()),
            ("Despacho/Tribunal solicitante",          self.entry_tribunal.currentText()),
            ("EXPEDIENTE",                             self.entry_caratula.text()),
        ]
        for row, (etq, val) in zip(tbl_gen.rows, filas):
            row.cells[0].text = etq
            row.cells[1].text = val

        # 4) Datos Ministerio Público Fiscal (tabla 1x2 con encabezado fusionado)
        doc.add_paragraph()  # salto de línea
        tbl_mpf = doc.add_table(rows=1, cols=2, style="Table Grid")
        mpf_hdr = tbl_mpf.rows[0].cells
        mpf_hdr[0].text = "DATOS MINISTERIO PÚBLICO FISCAL"
        mpf_hdr[0].merge(mpf_hdr[1])
        p = mpf_hdr[0].paragraphs[0]; p.runs[0].font.bold = True; p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tbl_mpf2 = doc.add_table(rows=1, cols=2, style="Table Grid")
        tbl_mpf2.rows[0].cells[0].text = "Fiscalía"
        tbl_mpf2.rows[0].cells[1].text = self.entry_fiscal.text()

        # 5) Actividad a desarrollar (fija)
        doc.add_paragraph()
        tbl_act = doc.add_table(rows=1, cols=2, style="Table Grid")
        tbl_act.rows[0].cells[0].text = "Tipo de Audiencia/Uso de Sala"
        tbl_act.rows[0].cells[1].text = "Audiencia Oral de Juicio Abreviado"

        # 6) Para cada imputado: bloque de 5 filas
        for i, w in enumerate(self.imputados_widgets, start=1):
            doc.add_paragraph()  # separación
            doc.add_paragraph(f"Imputado {i}", style="Intense Quote").runs[0].font.bold = True

            tbl_i = doc.add_table(rows=7, cols=2, style="Table Grid")
            datos = [
                ("Apellido y nombre",                         w['nombre'].text()),
                ("DNI",                                       w['dni'].text()),
                ("Delitos",                                   w['delitos'].text()),
                ("Fecha de detención",                        w['detenc'].text()),
                ("Defensa (y defensoria si corresponde)",     w['defensa'].text()),
                ("Número de teléfono (si es privada)",        ""),
                ("Correo electrónico (si es privada)",        ""),
            ]
            for row, (etq, val) in zip(tbl_i.rows, datos):
                row.cells[0].text = etq
                row.cells[1].text = val

        # 7) Indisponibilidad de agenda
        doc.add_paragraph()
        tbl_ind = doc.add_table(rows=1, cols=2, style="Table Grid")
        ind_hdr = tbl_ind.rows[0].cells
        ind_hdr[0].text = "Indisponibilidad de agenda"
        ind_hdr[0].merge(ind_hdr[1])
        p = ind_hdr[0].paragraphs[0]; p.runs[0].font.bold = True; p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tbl_ind2 = doc.add_table(rows=2, cols=2, style="Table Grid")
        tbl_ind2.rows[0].cells[0].text = "días"
        tbl_ind2.rows[1].cells[0].text = "horas"
        # las celdas [*][1] las dejamos en blanco para que el usuario complete si hace falta

        # 8) Guardar
        path, _ = QFileDialog.getSaveFileName(
            self, "Guardar planilla para OGA", "", "Word (*.docx)"
        )
        if path:
            doc.save(path)
            QMessageBox.information(self, "OK", "Planilla para OGA generada correctamente.")

    def guardar_causa(self):
        path, _ = QFileDialog.getSaveFileName(self, "Guardar causa",
                                              str(CAUSAS_DIR), "JSON (*.json)")
        if not path: return
        gen = {
            'caratula': self.entry_caratula.text(),
            'articulo': self.combo_articulo.currentText(),
            'tribunal': self.entry_tribunal.currentText(),
            'secretaria': self.entry_secretaria.text(),
            'fecha': self.entry_fecha.text(),
            'hora': self.combo_hora.currentText(),
            'sala': self.combo_sala.currentText(),
            'funcionario': self.entry_funcionario.text(),
            'fiscal': self.entry_fiscal.text(),
            'sentencia': self.entry_sentencia.text(),
            'resuelvo': self.entry_resuelvo.property("html") or "",
            'firmantes': self.entry_firmantes.text(),
            'renuncia': self.combo_renuncia.currentText()
        }
        imps = []
        for w in self.imputados_widgets:
            imps.append({k: (
                w[k].text() if isinstance(w[k], QLineEdit) else
                w[k].currentText() if isinstance(w[k], QComboBox) else
                w[k].isChecked()
            ) for k in w})

        with open(path, "w", encoding="utf-8") as f:
            json.dump({'generales': gen, 'imputados': imps},
                      f, ensure_ascii=False, indent=2)
        QMessageBox.information(self, "OK", "Causa guardada.")

    def cargar_causa(self):
        path, _ = QFileDialog.getOpenFileName(self, "Cargar causa",
                                              str(CAUSAS_DIR), "JSON (*.json)")
        if not path: return
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            g = data.get("generales", {})
            self.entry_caratula.setText(g.get("caratula", ""))
            self.combo_articulo.setCurrentText(g.get("articulo", ""))
            self.entry_tribunal.setCurrentText(g.get("tribunal", ""))
            self.entry_secretaria.setText(g.get("secretaria", ""))
            self.entry_fecha.setText(g.get("fecha", ""))
            self.combo_hora.setCurrentText(g.get("hora", ""))
            self.combo_sala.setCurrentText(g.get("sala", ""))
            self.entry_funcionario.setText(g.get("funcionario", ""))
            self.entry_fiscal.setText(g.get("fiscal", ""))
            self.entry_sentencia.setText(g.get("sentencia", ""))
            txt = g.get("resuelvo", "")
            if "<" in txt:
                self.entry_resuelvo.setHtml(txt)
            else:
                self.entry_resuelvo.setPlainText(txt)
            self.entry_resuelvo.setProperty("html", txt)
            self.entry_firmantes.setText(g.get("firmantes", ""))
            self.combo_renuncia.setCurrentText(g.get("renuncia", ""))

            imps = data.get("imputados", [])
            self.combo_n.setCurrentText(str(max(1, len(imps))))
            for idx, imp in enumerate(imps):
                w = self.imputados_widgets[idx]
                for k, v in imp.items():
                    if isinstance(w[k], QLineEdit):   w[k].setText(v)
                    elif isinstance(w[k], QComboBox): w[k].setCurrentText(v)
                    elif isinstance(w[k], QCheckBox): w[k].setChecked(bool(v))
            self.update_template()
        except Exception as e:
            QMessageBox.critical(self, "Error", str(e))

    def eliminar_causa(self):
        path, _ = QFileDialog.getOpenFileName(self, "Eliminar causa",
                                              str(CAUSAS_DIR), "JSON (*.json)")
        if path and QMessageBox.question(
            self, "Confirmar", f"¿Eliminar {Path(path).name}?"
        ) == QMessageBox.Yes:
            Path(path).unlink(missing_ok=True)

# utils.py  ─────────────────────────────────────────────────────────────
from PySide6.QtWidgets import QApplication, QMessageBox

def confirm_and_quit(widget) -> None:
    """Muestra un QMessageBox; si el usuario acepta, cierra TODA la app."""
    ans = QMessageBox.question(
        widget,
        "Cerrar la aplicación",
        "¿Está seguro de que desea salir?\nSe cerrarán todas las ventanas.",
        QMessageBox.Yes | QMessageBox.No,
        QMessageBox.No
    )
    if ans == QMessageBox.Yes:
        QApplication.quit()

def main():
    app  = QApplication(sys.argv)
    data = CausaData()          # ① instancia compartida
    win  = MainWindow(data)     # ② pásala a la ventana principal
    win.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()
